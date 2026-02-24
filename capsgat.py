import sys
import re
import json
import os
import math
import tempfile
import webbrowser
import logging
from pathlib import Path
from datetime import datetime
from collections import deque
import queue
import threading
import time
from PyQt5.QtWidgets import (QApplication, QMainWindow, QVBoxLayout, QHBoxLayout,
                             QTextEdit, QListWidget, QPushButton, QWidget, QLabel,
                             QFileDialog, QMessageBox, QSpinBox, QShortcut, QFrame,
                             QInputDialog, QLineEdit, QDialog, QDialogButtonBox,
                             QGridLayout, QPlainTextEdit, QCheckBox, QTabWidget, QRadioButton,
                             QSlider, QProgressBar, QMenuBar, QMenu, QAction, QFontDialog,
                             QGroupBox, QScrollArea, QSizePolicy, QComboBox, QStackedWidget, QStyle, QSplashScreen)
from PyQt5.QtCore import Qt, QTimer, QUrl, pyqtSignal, QPoint, QRect, QElapsedTimer, QThread, QSize
from PyQt5.QtGui import QFont, QKeySequence, QColor, QTextCharFormat, QSyntaxHighlighter, QIcon, QPixmap, QPainter, QPen, QBrush, QPainterPath

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import pyaudio
    import soundfile as sf
    HAS_PYAUDIO = True
except ImportError:
    HAS_PYAUDIO = False
    logger.warning("PyAudio not installed. Fallback audio will not work.")
    
def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class SimpleAudioPlayer(QThread):
    """Simple fallback audio player using PyAudio (no speed control)"""
    playback_started = pyqtSignal()
    playback_stopped = pyqtSignal()
    position_changed = pyqtSignal(float)  # Position in seconds
    
    def __init__(self):
        super().__init__()
        self.audio_path = None
        self.is_playing = False
        self.is_paused = False
        self.stop_flag = False
        self.current_position = 0.0
        self.duration = 0.0
        self.sample_rate = 44100
        self.channels = 1
        self.pyaudio = None
        self.stream = None
        self.audio_data = None
        self.lock = threading.Lock()
        
    def load_file(self, audio_path):
        """Load audio file"""
        try:
            import numpy as np
            
            self.audio_path = audio_path
            
            # Load audio data
            audio_data, self.sample_rate = sf.read(audio_path)
            
            # Convert to mono if needed
            if len(audio_data.shape) > 1:
                audio_data = audio_data.mean(axis=1)
            
            # Convert to int16
            if audio_data.dtype != np.int16:
                audio_data = (audio_data * 32767).astype(np.int16)
            
            self.audio_data = audio_data.tobytes()
            self.duration = len(audio_data) / self.sample_rate
            self.channels = 1
            
            logger.info(f"Simple audio loaded: {self.duration:.2f}s")
            return True
            
        except Exception as e:
            logger.error(f"Failed to load audio for fallback: {e}")
            return False
    
    def play(self):
        """Start playback"""
        with self.lock:
            if not HAS_PYAUDIO:
                logger.error("PyAudio not available for fallback")
                return
            
            # Stop any existing playback
            self._stop_playback()
            
            self.is_playing = True
            self.is_paused = False
            self.stop_flag = False
            
            if not self.isRunning():
                self.start()
            
            self.playback_started.emit()
    
    def pause(self):
        """Pause playback"""
        with self.lock:
            self.is_playing = False
            self.is_paused = True
    
    def stop(self):
        """Stop playback"""
        with self.lock:
            self.stop_flag = True
            self.is_playing = False
            self.is_paused = False
            self.current_position = 0.0
            
            self._stop_playback()
            
            self.playback_stopped.emit()
    
    def _stop_playback(self):
        """Safely stop audio playback"""
        try:
            if self.stream:
                if self.stream.is_active():
                    self.stream.stop_stream()
                self.stream.close()
                self.stream = None
        except Exception as e:
            logger.debug(f"Error stopping stream: {e}")
            self.stream = None
        
        # Give time for stream to close
        time.sleep(0.05)
    
    def seek(self, position_seconds):
        """Seek to position"""
        with self.lock:
            new_position = max(0, min(self.duration, position_seconds))
            self.current_position = new_position
            
            # If currently playing, restart from new position
            if self.is_playing and not self.is_paused:
                # Signal the thread to restart
                self.stop_flag = True
                
                # Wait a bit for thread to respond
                time.sleep(0.05)
                
                # Reset and restart
                self.stop_flag = False
                self.is_paused = False
                
                # Clear stream
                self._stop_playback()
    
    def get_position(self):
        """Get current position"""
        return self.current_position
    
    def run(self):
        """Main playback thread"""
        if not HAS_PYAUDIO or not self.audio_data:
            logger.error("Fallback player: No PyAudio or audio data")
            return
        
        try:
            import numpy as np
            import pyaudio
            
            with self.lock:
                self.pyaudio = pyaudio.PyAudio()
            
            while not self.stop_flag:
                # Check if we should play
                if not self.is_playing or self.is_paused:
                    time.sleep(0.01)
                    continue
                
                # Calculate start position in bytes
                start_byte = int(self.current_position * self.sample_rate * 2)
                
                with self.lock:
                    # Create new stream
                    self.stream = self.pyaudio.open(
                        format=pyaudio.paInt16,
                        channels=self.channels,
                        rate=self.sample_rate,
                        output=True,
                        frames_per_buffer=2048
                    )
                
                # Play from current position
                chunk_size = 4096
                data_bytes = len(self.audio_data)
                
                for i in range(start_byte, data_bytes, chunk_size):
                    if self.stop_flag:
                        break
                    
                    if not self.is_playing or self.is_paused:
                        break
                    
                    chunk = self.audio_data[i:i+chunk_size]
                    if chunk:
                        try:
                            if self.stream:
                                self.stream.write(chunk)
                        except Exception as e:
                            logger.warning(f"Error writing to stream: {e}")
                            break
                        
                        # Update position
                        self.current_position = i / (self.sample_rate * 2)
                        self.position_changed.emit(self.current_position)
                        
                        # Check if at end
                        if i + chunk_size >= data_bytes:
                            self.stop()
                            break
                    
                    # Small sleep to prevent CPU overuse
                    time.sleep(chunk_size / (self.sample_rate * 2) * 0.9)
                
                # Clean up stream
                with self.lock:
                    self._stop_playback()
                
                # If we broke out of loop but still supposed to be playing,
                # we were seeking or paused
                if self.is_playing and not self.is_paused and not self.stop_flag:
                    # Continue from current position
                    continue
                else:
                    break
            
        except Exception as e:
            logger.error(f"Error in fallback player: {e}")
        finally:
            with self.lock:
                self._stop_playback()
                if self.pyaudio:
                    try:
                        self.pyaudio.terminate()
                    except:
                        pass
                    self.pyaudio = None
    
    def cleanup(self):
        """Clean up resources"""
        self.stop()
        if self.isRunning():
            self.quit()
            self.wait(1000)  # Wait up to 1 second
            
class VlcAudioPlayer(QThread):
    """Audio player using VLC media player"""
    playback_started = pyqtSignal()
    playback_paused = pyqtSignal()
    playback_stopped = pyqtSignal()
    position_changed = pyqtSignal(float)  # Position in seconds
    end_reached = pyqtSignal()
    
    def __init__(self):
        super().__init__()
        self.vlc_available = False
        self.instance = None
        self.player = None
        self.media = None
        self.is_playing = False
        self.duration = 0.0
        self.current_position = 0.0
        self.playback_speed = 1.0
        self.audio_file_path = None
        
        # Try to import VLC and create instance
        try:
            import vlc
            self.vlc = vlc
            self.instance = self.vlc.Instance()
            self.player = self.instance.media_player_new()
            self.vlc_available = True
            
            # Timer for position updates
            self.position_timer = QTimer()
            self.position_timer.timeout.connect(self.update_position)
            self.position_timer.start(100)  # Update every 100ms
            
            # Event manager for VLC events
            self.event_manager = self.player.event_manager()
            self.event_manager.event_attach(self.vlc.EventType.MediaPlayerEndReached, self._on_end_reached)
            
        except ImportError:
            logger.warning("VLC not available")
            self.vlc_available = False
        
    def load_file(self, audio_path):
        """Load audio file"""
        if not self.vlc_available:
            logger.error("VLC not available, cannot load file")
            return False
            
        try:
            self.audio_file_path = audio_path
            self.media = self.instance.media_new(audio_path)
            self.player.set_media(self.media)
            
            # Get duration
            self.media.parse()
            time.sleep(0.1)  # Give VLC time to parse
            self.duration = self.media.get_duration() / 1000.0  # Convert ms to seconds
            
            # Set initial speed
            self.player.set_rate(1.0)
            self.playback_speed = 1.0
            
            logger.info(f"Audio loaded: {audio_path}, duration: {self.duration:.2f}s")
            return True
            
        except Exception as e:
            logger.error(f"Failed to load audio with VLC: {e}")
            return False
    
    def play(self):
        """Start playback"""
        if not self.vlc_available or not self.player:
            return
        if self.player.play() == 0:
            self.is_playing = True
            self.playback_started.emit()
            logger.info("Playback started")
        else:
            logger.error("Failed to start playback")
    
    def pause(self):
        """Pause playback"""
        if not self.vlc_available or not self.player:
            return
        self.player.pause()
        self.is_playing = False
        self.playback_paused.emit()
        logger.info("Playback paused")
    
    def stop(self):
        """Stop playback"""
        if not self.vlc_available or not self.player:
            return
        self.player.stop()
        self.is_playing = False
        self.current_position = 0.0
        self.playback_stopped.emit()
        logger.info("Playback stopped")
    
    def seek(self, position_seconds):
        """Seek to position in seconds"""
        if not self.vlc_available or not self.player:
            return
        try:
            position_ms = int(position_seconds * 1000)
            self.player.set_time(position_ms)
            self.current_position = position_seconds
            self.position_changed.emit(position_seconds)
            logger.debug(f"Seeked to {position_seconds:.2f}s")
        except Exception as e:
            logger.error(f"Error seeking: {e}")
    
    def set_speed(self, speed):
        """Set playback speed (0.5 to 2.0)"""
        if not self.vlc_available or not self.player:
            return False
        try:
            speed = max(0.25, min(4.0, speed))
            self.player.set_rate(speed)
            self.playback_speed = speed
            logger.info(f"Playback speed set to {speed:.1f}x")
            return True
        except Exception as e:
            logger.error(f"Error setting playback speed: {e}")
            return False
    
    def get_position(self):
        """Get current position in seconds"""
        if not self.vlc_available or not self.player:
            return self.current_position
        try:
            time_ms = self.player.get_time()
            if time_ms >= 0:
                self.current_position = time_ms / 1000.0
            return self.current_position
        except:
            return self.current_position
    
    def get_state(self):
        """Get current player state"""
        if not self.vlc_available or not self.player:
            return None
        return self.player.get_state()
    
    def update_position(self):
        """Update and emit current position"""
        if self.is_playing:
            pos = self.get_position()
            if pos >= 0:
                self.position_changed.emit(pos)
                
                # Check if we've reached the end
                if self.duration > 0 and pos >= self.duration - 0.1:
                    self.stop()
    
    def _on_end_reached(self, event):
        """Handle end of media"""
        self.is_playing = False
        self.end_reached.emit()
        logger.info("End of media reached")
    
    def cleanup(self):
        """Clean up resources"""
        if hasattr(self, 'position_timer'):
            self.position_timer.stop()
        self.stop()
        if self.player:
            self.player.release()
        if self.instance:
            self.instance.release()

class TextSelectionDialog(QDialog):
    def __init__(self, block_text, parent=None):
        super().__init__(parent)
        self.block_text = block_text
        self.start_pos = 0
        self.end_pos = 0
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Select Text")
        self.setGeometry(300, 300, 700, 300)
        
        layout = QVBoxLayout(self)
        
        instructions = QLabel("Use ← → arrows to adjust selection, Shift+arrows to extend, then press Enter:")
        instructions.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions)
        
        self.text_display = QTextEdit()
        self.text_display.setReadOnly(True)
        self.text_display.setMaximumHeight(100)
        self.text_display.setStyleSheet("font-family: monospace; font-size: 14px;")
        layout.addWidget(self.text_display)
        
        self.selection_label = QLabel("Selection: (none)")
        layout.addWidget(self.selection_label)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        for button in button_box.buttons():
            button.setFocusPolicy(Qt.NoFocus)
            
        layout.addWidget(button_box)
        
        self.update_display()
        self.setFocusPolicy(Qt.StrongFocus)
        self.setFocus()
        
    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Left and not event.modifiers() & Qt.ShiftModifier:
            if self.start_pos > 0:
                self.start_pos -= 1
                self.end_pos = self.start_pos
            self.update_display()
            event.accept()
        elif event.key() == Qt.Key_Right and not event.modifiers() & Qt.ShiftModifier:
            if self.end_pos < len(self.block_text):
                self.start_pos += 1
                self.end_pos = self.start_pos
            self.update_display()
            event.accept()
        elif event.key() == Qt.Key_Left and event.modifiers() & Qt.ShiftModifier:
            if self.start_pos > 0:
                self.start_pos -= 1
            self.update_display()
            event.accept()
        elif event.key() == Qt.Key_Right and event.modifiers() & Qt.ShiftModifier:
            if self.end_pos < len(self.block_text):
                self.end_pos += 1
            self.update_display()
            event.accept()
        elif event.key() in (Qt.Key_Return, Qt.Key_Enter):
            self.accept()
            event.accept()
        else:
            super().keyPressEvent(event)
    
    def update_display(self):
        before_text = self.block_text[:self.start_pos]
        selected_text = self.block_text[self.start_pos:self.end_pos]
        after_text = self.block_text[self.end_pos:]
        
        html_content = f"""
        <div style="font-family: monospace; font-size: 14px; padding: 10px;">
            <span style="background-color: #e0e0e0; padding: 5px; border-radius: 3px;">{before_text}</span>
            <span style="background-color: #ffcc00; padding: 5px; border-radius: 3px;">{selected_text}</span>
            <span style="background-color: #e0e0e0; padding: 5px; border-radius: 3px;">{after_text}</span>
        </div>
        """
        
        self.text_display.setHtml(html_content)
        
        if self.start_pos == self.end_pos:
            self.selection_label.setText(f"Selection: (none) - Position: {self.start_pos}")
        else:
            self.selection_label.setText(f"Selection: '{selected_text}' (positions {self.start_pos}-{self.end_pos})")
    
    def get_selection(self):
        return self.start_pos, self.end_pos, self.block_text[self.start_pos:self.end_pos]

class BlockSplitDialog(QDialog):
    def __init__(self, block_text, parent=None):
        super().__init__(parent)
        self.block_text = block_text
        self.split_position = 0
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Split Block")
        self.setGeometry(300, 300, 700, 250)
        
        layout = QVBoxLayout(self)
        
        instructions = QLabel("Use ← → arrows to position split, then press Enter to confirm:")
        instructions.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions)
        
        self.text_display = QTextEdit()
        self.text_display.setReadOnly(True)
        self.text_display.setMaximumHeight(100)
        self.text_display.setStyleSheet("font-family: monospace; font-size: 14px;")
        layout.addWidget(self.text_display)
        
        self.cursor_label = QLabel("Split position: 0")
        layout.addWidget(self.cursor_label)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        for button in button_box.buttons():
            button.setFocusPolicy(Qt.NoFocus)
            
        layout.addWidget(button_box)
        
        self.update_display()
        self.setFocusPolicy(Qt.StrongFocus)
        self.setFocus()
        
    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Left:
            self.split_position = max(0, self.split_position - 1)
            self.update_display()
            event.accept()
        elif event.key() == Qt.Key_Right:
            self.split_position = min(len(self.block_text), self.split_position + 1)
            self.update_display()
            event.accept()
        elif event.key() in (Qt.Key_Return, Qt.Key_Enter, Qt.Key_Space):
            self.accept()
            event.accept()
        else:
            super().keyPressEvent(event)
    
    def update_display(self):
        before_text = self.block_text[:self.split_position]
        after_text = self.block_text[self.split_position:]
        
        html_content = f"""
        <div style="font-family: monospace; font-size: 14px; padding: 10px;">
            <span style="background-color: #c8f7c8; padding: 5px; border-radius: 3px;">{before_text}</span>
            <span style="background-color: #f7c8c8; padding: 5px; border-radius: 3px;">{after_text}</span>
        </div>
        """
        
        self.text_display.setHtml(html_content)
        self.cursor_label.setText(f"Split position: {self.split_position} (text will be split after character {self.split_position})")

class SymbolCategory:
    def __init__(self, name, symbols, descriptions=None):
        self.name = name
        self.symbols = symbols
        self.descriptions = descriptions or [""] * len(symbols)
        self.selected_index = 0

class AddCustomSymbolDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add Custom Symbol")
        self.setGeometry(300, 300, 500, 250)
        layout = QVBoxLayout(self)

        # Symbol type selector
        type_layout = QHBoxLayout()
        type_layout.addWidget(QLabel("Symbol type:"))
        self.type_combo = QComboBox()
        self.type_combo.addItems(["Simple", "Segment Wrapper", "Comment Wrapper", "Comment with Reach"])
        self.type_combo.currentTextChanged.connect(self.update_fields)
        type_layout.addWidget(self.type_combo)
        layout.addLayout(type_layout)

        # Stacked widget for different input fields
        self.stacked = QStackedWidget()
        layout.addWidget(self.stacked)

        # ---- Page 0: Simple ----
        simple_page = QWidget()
        simple_layout = QVBoxLayout(simple_page)
        self.simple_symbol_edit = QLineEdit()
        self.simple_symbol_edit.setPlaceholderText("Enter symbol text")
        simple_layout.addWidget(QLabel("Symbol:"))
        simple_layout.addWidget(self.simple_symbol_edit)
        self.simple_desc_edit = QLineEdit()
        self.simple_desc_edit.setPlaceholderText("Optional description")
        simple_layout.addWidget(QLabel("Description (optional):"))
        simple_layout.addWidget(self.simple_desc_edit)
        self.stacked.addWidget(simple_page)

        # ---- Page 1: Segment Wrapper ----
        wrapper_page = QWidget()
        wrapper_layout = QVBoxLayout(wrapper_page)
        self.wrapper_left_edit = QLineEdit()
        self.wrapper_left_edit.setPlaceholderText("e.g., <<")
        wrapper_layout.addWidget(QLabel("Left side:"))
        wrapper_layout.addWidget(self.wrapper_left_edit)
        self.wrapper_right_edit = QLineEdit()
        self.wrapper_right_edit.setPlaceholderText("e.g., >>")
        wrapper_layout.addWidget(QLabel("Right side:"))
        wrapper_layout.addWidget(self.wrapper_right_edit)
        self.wrapper_desc_edit = QLineEdit()
        self.wrapper_desc_edit.setPlaceholderText("Optional description")
        wrapper_layout.addWidget(QLabel("Description (optional):"))
        wrapper_layout.addWidget(self.wrapper_desc_edit)
        self.stacked.addWidget(wrapper_page)

        # ---- Page 2: Comment Wrapper ----
        comment_page = QWidget()
        comment_layout = QVBoxLayout(comment_page)
        self.comment_left_edit = QLineEdit()
        self.comment_left_edit.setPlaceholderText("e.g., ((")
        comment_layout.addWidget(QLabel("Left side:"))
        comment_layout.addWidget(self.comment_left_edit)
        self.comment_right_edit = QLineEdit()
        self.comment_right_edit.setPlaceholderText("e.g., ))")
        comment_layout.addWidget(QLabel("Right side:"))
        comment_layout.addWidget(self.comment_right_edit)
        self.comment_desc_edit = QLineEdit()
        self.comment_desc_edit.setPlaceholderText("Optional description")
        comment_layout.addWidget(QLabel("Description (optional):"))
        comment_layout.addWidget(self.comment_desc_edit)
        self.stacked.addWidget(comment_page)

        # ---- Page 3: Comment with Reach ----
        reach_page = QWidget()
        reach_layout = QVBoxLayout(reach_page)
        self.reach_left_edit = QLineEdit()
        self.reach_left_edit.setPlaceholderText("e.g., [at^=")
        reach_layout.addWidget(QLabel("Action left side:"))
        reach_layout.addWidget(self.reach_left_edit)
        self.reach_right_action_edit = QLineEdit()
        self.reach_right_action_edit.setPlaceholderText("e.g., ]")
        reach_layout.addWidget(QLabel("Action right side:"))
        reach_layout.addWidget(self.reach_right_action_edit)
        self.reach_right_segment_edit = QLineEdit()
        self.reach_right_segment_edit.setPlaceholderText("e.g., ]")
        reach_layout.addWidget(QLabel("Segment right side:"))
        reach_layout.addWidget(self.reach_right_segment_edit)
        self.reach_desc_edit = QLineEdit()
        self.reach_desc_edit.setPlaceholderText("Optional description")
        reach_layout.addWidget(QLabel("Description (optional):"))
        reach_layout.addWidget(self.reach_desc_edit)
        self.stacked.addWidget(reach_page)

        # OK / Cancel buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

        self.update_fields()

    def update_fields(self):
        """Switch stacked page based on selected type"""
        index = self.type_combo.currentIndex()
        self.stacked.setCurrentIndex(index)

    def get_symbol_data(self):
        """Return dictionary with symbol info, or None if validation fails"""
        data = {}
        index = self.type_combo.currentIndex()
        if index == 0:  # Simple
            symbol = self.simple_symbol_edit.text().strip()
            if not symbol:
                QMessageBox.warning(self, "Missing Data", "Symbol text cannot be empty.")
                return None
            data = {
                'type': 'simple',
                'display': symbol,
                'value': symbol,
                'description': self.simple_desc_edit.text().strip() or symbol
            }
        elif index == 1:  # Segment Wrapper
            left = self.wrapper_left_edit.text().strip()
            right = self.wrapper_right_edit.text().strip()
            if not left or not right:
                QMessageBox.warning(self, "Missing Data", "Both left and right sides are required.")
                return None
            display = f"{left}text{right}"
            data = {
                'type': 'wrapper',
                'display': display,
                'left': left,
                'right': right,
                'description': self.wrapper_desc_edit.text().strip() or display
            }
        elif index == 2:  # Comment Wrapper
            left = self.comment_left_edit.text().strip()
            right = self.comment_right_edit.text().strip()
            if not left or not right:
                QMessageBox.warning(self, "Missing Data", "Both left and right sides are required.")
                return None
            display = f"{left}comment{right}"
            data = {
                'type': 'comment',
                'display': display,
                'left': left,
                'right': right,
                'description': self.comment_desc_edit.text().strip() or display
            }
        else:  # Comment with Reach
            left = self.reach_left_edit.text().strip()
            right_action = self.reach_right_action_edit.text().strip()
            right_segment = self.reach_right_segment_edit.text().strip()
            if not left or not right_action or not right_segment:
                QMessageBox.warning(self, "Missing Data", "All three sides are required.")
                return None
            display = f"{left}attitude{right_action}text{right_segment}"
            data = {
                'type': 'comment_reach',
                'display': display,
                'left': left,
                'right': right_action,
                'segment_right': right_segment,
                'description': self.reach_desc_edit.text().strip() or display
            }
        return data

class EnhancedSymbolDialog(QDialog):
    custom_symbols_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "custom_symbols.json")
    custom_symbols = []

    def __init__(self, parent=None, initial_category=None):
        super().__init__(parent)
        self.parent = parent
        self.initial_category = initial_category
        self.categories = []
        self.current_category_index = 0
        self.selected_option = 0
        self.load_custom_symbols()
        self.init_categories()
        self.init_ui()

    def init_categories(self):
        # GAT2
        self.categories.append(SymbolCategory(
            "GAT2",
            ["(.)", "(-)", "(--)", "(---)", "(_._)", "(())", "<<>>", "[ ]",
             "°h", "°hh", "°hhh", "h°", "hh°", "hhh°"],
            ["micropause", "short pause", "medium pause", "long pause",
             "measured pause", "comment", "action", "overlap",
             "short inhale", "medium inhale", "long inhale",
             "short exhale", "medium exhale", "long exhale"]
        ))
        # Dresing & Pehl
        self.categories.append(SymbolCategory(
            "Dresing && Pehl",
            ["(.)", "(..)", "(...)", "(_)", "//", "(   )", "⏱️"],
            ["short pause", "medium pause", "long pause",
             "measured pause", "overlap", "comment", "insert timestamp"]
        ))
        # TiQ
        self.categories.append(SymbolCategory(
            "TiQ",
            ["(.)", "(_)", "(())", "└", "@(.)@", "@(_)@", "@(   )@", "°   °", "//   //"],
            ["short pause", "measured pause", "comment",
             "overlap marker", "short laughter", "laughing seconds",
             "laughing speech", "quiet speech", "listener's signal"]
        ))
        # Custom
        custom_symbols_list = [s['display'] for s in self.custom_symbols]
        custom_descriptions = [s.get('description', s['type']) for s in self.custom_symbols]
        self.categories.append(SymbolCategory(
            "Custom",
            custom_symbols_list,
            custom_descriptions
        ))

    def load_custom_symbols(self):
        try:
            if os.path.exists(EnhancedSymbolDialog.custom_symbols_file):
                with open(EnhancedSymbolDialog.custom_symbols_file, 'r', encoding='utf-8') as f:
                    self.custom_symbols = json.load(f)
            else:
                self.custom_symbols = []
        except Exception as e:
            logger.error(f"Failed to load custom symbols: {e}")
            self.custom_symbols = []

    def save_custom_symbols(self):
        try:
            with open(EnhancedSymbolDialog.custom_symbols_file, 'w', encoding='utf-8') as f:
                json.dump(self.custom_symbols, f, indent=2)
        except Exception as e:
            logger.error(f"Failed to save custom symbols: {e}")
            QMessageBox.warning(self, "Error", f"Could not save custom symbols: {e}")

    def init_ui(self):
        self.setWindowTitle("Insert Symbol")
        self.setGeometry(300, 300, 650, 500)

        layout = QVBoxLayout(self)

        # Category tabs
        tab_layout = QHBoxLayout()
        self.category_buttons = []
        for i, category in enumerate(self.categories):
            btn = QPushButton(category.name)
            btn.setCheckable(True)
            btn.setChecked(i == 0)
            btn.setFocusPolicy(Qt.TabFocus)
            btn.setStyleSheet("""
                QPushButton {
                    padding: 8px 15px;
                    font-weight: bold;
                    border: 2px solid #ccc;
                    border-radius: 5px 5px 0 0;
                    background-color: #f0f0f0;
                }
                QPushButton:checked {
                    background-color: #4a90e2;
                    color: white;
                    border-bottom-color: #4a90e2;
                }
                QPushButton:hover:!checked {
                    background-color: #e0e0e0;
                }
                QPushButton:focus {
                    outline: 2px solid #ff6600;
                }
            """)
            btn.clicked.connect(lambda checked, idx=i: self.switch_category(idx))
            tab_layout.addWidget(btn)
            self.category_buttons.append(btn)

        tab_instruction = QLabel("<i>(Tab to switch)</i>")
        tab_instruction.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
        tab_instruction.setStyleSheet("color: #555; font-size: 11px;")
        tab_layout.addWidget(tab_instruction)
        tab_layout.addSpacing(23)
        layout.addLayout(tab_layout)
        
        # Management buttons for custom category
        self.management_layout = QHBoxLayout()
               
        self.btn_add_symbol = QPushButton("Add New Symbol")
        self.btn_add_symbol.clicked.connect(self.add_new_symbol)
        self.btn_add_symbol.setVisible(False)
        
        self.btn_delete_custom = QPushButton("Delete Selected")
        self.btn_delete_custom.clicked.connect(self.delete_selected_symbol)
        self.btn_delete_custom.setVisible(False)
        self.btn_delete_custom.setStyleSheet("""
            QPushButton:hover {
                background-color: #fdd2d2;
                border: 1px solid #cc0000;
                border-radius: 3px;
            }
            QPushButton:disabled {
                background-color: #cccccc;
                border-color: #999999;
                color: #666666;
            }
        """)

        self.btn_export_custom = QPushButton("Export")
        self.btn_export_custom.clicked.connect(self.export_custom_symbols)
        self.btn_export_custom.setVisible(False)

        self.btn_import_custom = QPushButton("Import")
        self.btn_import_custom.clicked.connect(self.import_custom_symbols)
        self.btn_import_custom.setVisible(False)

        self.management_layout.addWidget(self.btn_add_symbol)
        self.management_layout.addWidget(self.btn_delete_custom)
        self.management_layout.addStretch()
        self.management_layout.addWidget(self.btn_export_custom)
        self.management_layout.addWidget(self.btn_import_custom)

        layout.addLayout(self.management_layout)

        # Scrollable grid area
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setFrameShape(QFrame.NoFrame)
        self.category_widget = QWidget()
        self.category_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.category_layout = QGridLayout(self.category_widget)
        self.category_layout.setSpacing(10)
        self.category_layout.setAlignment(Qt.AlignCenter)
        self.scroll_area.setWidget(self.category_widget)
        layout.addWidget(self.scroll_area)

        # Selected description label
        self.selected_label = QLabel("")
        self.selected_label.setAlignment(Qt.AlignCenter)
        self.selected_label.setStyleSheet("""
            QLabel {
                background-color: #f0f0f0;
                border: 1px solid #ccc;
                border-radius: 4px;
                padding: 6px;
                font-size: 12px;
                color: #333;
                margin-top: 5px;
                margin-bottom: 5px;
            }
        """)
        layout.addWidget(self.selected_label)

        # OK / Cancel buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

        # Shortcuts for Tab / Shift+Tab to switch categories
        self.tab_shortcut = QShortcut(QKeySequence(Qt.Key_Tab), self)
        self.tab_shortcut.activated.connect(self.next_category)
        self.shift_tab_shortcut = QShortcut(QKeySequence(Qt.Key_Backtab), self)
        self.shift_tab_shortcut.activated.connect(self.prev_category)
        self.tab_shortcut.setContext(Qt.WidgetWithChildrenShortcut)
        self.shift_tab_shortcut.setContext(Qt.WidgetWithChildrenShortcut)

        self.update_category_display()
        if self.initial_category is not None:
            self.switch_category(self.initial_category)

        self.setFocusPolicy(Qt.StrongFocus)
        self.setFocus()

    def next_category(self):
        new_index = (self.current_category_index + 1) % len(self.categories)
        self.switch_category(new_index)

    def prev_category(self):
        new_index = (self.current_category_index - 1) % len(self.categories)
        self.switch_category(new_index)

    def switch_category(self, index):
        if 0 <= index < len(self.categories):
            for i, btn in enumerate(self.category_buttons):
                btn.setChecked(i == index)

            self.current_category_index = index
            self.selected_option = 0
            self.categories[index].selected_index = 0
            self.update_category_display()
            self.setFocus()

            # Show/hide custom management buttons
            is_custom = (index == len(self.categories) - 1)
            self.btn_add_symbol.setVisible(is_custom)
            self.btn_export_custom.setVisible(is_custom)
            self.btn_import_custom.setVisible(is_custom)
            self.btn_delete_custom.setVisible(is_custom)

    def update_category_display(self):
        # Clear all widgets from the grid layout
        for i in reversed(range(self.category_layout.count())):
            widget = self.category_layout.itemAt(i).widget()
            if widget:
                widget.setParent(None)

        category = self.categories[self.current_category_index]

        self.symbol_labels = []
        for i, (symbol, desc) in enumerate(zip(category.symbols, category.descriptions)):
            escaped_symbol = (symbol.replace('&', '&amp;')
                                   .replace('<', '&lt;')
                                   .replace('>', '&gt;'))

            label = QLabel(f"<b>{escaped_symbol}</b>")
            label.setAlignment(Qt.AlignCenter)
            label.setToolTip(desc)
            label.setCursor(Qt.PointingHandCursor)
            label.setFocusPolicy(Qt.ClickFocus)

            # Set size limits
            label.setMinimumSize(110, 60)
            label.setMaximumSize(180, 100)
            label.setSizePolicy(QSizePolicy.Maximum, QSizePolicy.Maximum)

            if i == self.selected_option:
                label.setStyleSheet("""
                    QLabel {
                        border: 3px solid #ff6600;
                        border-radius: 8px;
                        padding: 15px;
                        background-color: #fff0cc;
                        font-size: 14px;
                    }
                """)
            else:
                label.setStyleSheet("""
                    QLabel {
                        border: 2px solid #ccc;
                        border-radius: 8px;
                        padding: 15px;
                        background-color: #f9f9f9;
                        font-size: 14px;
                    }
                    QLabel:hover {
                        background-color: #e0e0e0;
                        border-color: #999;
                    }
                """)

            label.mousePressEvent = lambda event, idx=i: self.symbol_clicked(idx)

            self.category_layout.addWidget(label, i // 4, i % 4)
            self.symbol_labels.append(label)

        # Ensure the selected symbol is visible
        if self.symbol_labels:
            self.scroll_area.ensureWidgetVisible(self.symbol_labels[self.selected_option])

        # Update description label
        desc = category.descriptions[self.selected_option] if category.descriptions else ""
        self.selected_label.setText(desc)

    def symbol_clicked(self, index):
        self.selected_option = index
        self.categories[self.current_category_index].selected_index = index
        self.update_category_display()

    def keyPressEvent(self, event):
        # Arrow keys always navigate the grid (dialog has focus)
        if event.key() in (Qt.Key_Left, Qt.Key_Right, Qt.Key_Up, Qt.Key_Down,
                           Qt.Key_Return, Qt.Key_Enter, Qt.Key_Escape):
            cols = 4
            current_idx = self.selected_option
            max_idx = len(self.categories[self.current_category_index].symbols) - 1

            if event.key() == Qt.Key_Left:
                new_idx = max(0, current_idx - 1)
                self.symbol_clicked(new_idx)
                event.accept()
            elif event.key() == Qt.Key_Right:
                new_idx = min(max_idx, current_idx + 1)
                self.symbol_clicked(new_idx)
                event.accept()
            elif event.key() == Qt.Key_Up:
                new_idx = max(0, current_idx - cols)
                self.symbol_clicked(new_idx)
                event.accept()
            elif event.key() == Qt.Key_Down:
                new_idx = min(max_idx, current_idx + cols)
                self.symbol_clicked(new_idx)
                event.accept()
            elif event.key() in (Qt.Key_Return, Qt.Key_Enter):
                self.accept()
                event.accept()
            elif event.key() == Qt.Key_Escape:
                self.reject()
                event.accept()
            else:
                super().keyPressEvent(event)
        else:
            super().keyPressEvent(event)

    def add_new_symbol(self):
        """Open the add custom symbol dialog and append the new symbol."""
        dialog = AddCustomSymbolDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            symbol_data = dialog.get_symbol_data()
            if symbol_data:
                self.custom_symbols.append(symbol_data)
                self.save_custom_symbols()
                self.refresh_custom_category()

    def refresh_custom_category(self):
        custom_idx = len(self.categories) - 1
        self.categories[custom_idx].symbols = [s['display'] for s in self.custom_symbols]
        self.categories[custom_idx].descriptions = [s.get('description', s['type']) for s in self.custom_symbols]
        self.categories[custom_idx].selected_index = 0

        if self.current_category_index == custom_idx:
            self.selected_option = 0
            self.update_category_display()

    def export_custom_symbols(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Custom Symbols", "",
            "JSON Files (*.json);;All Files (*)"
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(self.custom_symbols, f, indent=2)
                QMessageBox.information(self, "Success", f"Custom symbols exported to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not export: {e}")

    def import_custom_symbols(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Import Custom Symbols", "",
            "JSON Files (*.json);;All Files (*)"
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    imported = json.load(f)

                if isinstance(imported, list):
                    reply = QMessageBox.question(
                        self, "Merge Symbols",
                        f"Found {len(imported)} symbols. Merge with existing?",
                        QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
                    )

                    if reply == QMessageBox.Yes:
                        self.custom_symbols.extend(imported)
                    elif reply == QMessageBox.No:
                        self.custom_symbols = imported
                    else:
                        return

                    self.save_custom_symbols()
                    self.refresh_custom_category()
                    QMessageBox.information(self, "Success", f"Imported {len(imported)} symbols")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not import: {e}")

    def delete_selected_symbol(self):
        if self.current_category_index != len(self.categories) - 1:
            return
        if not self.custom_symbols:
            return

        reply = QMessageBox.question(
            self, "Delete Symbol",
            f"Are you sure you want to delete the symbol '{self.categories[-1].symbols[self.selected_option]}'?",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            del self.custom_symbols[self.selected_option]
            self.save_custom_symbols()
            if self.selected_option >= len(self.custom_symbols):
                self.selected_option = max(0, len(self.custom_symbols) - 1)
            self.refresh_custom_category()

    def get_selected_symbol_info(self):
        category = self.categories[self.current_category_index]
        symbol_display = category.symbols[self.selected_option]

        if self.current_category_index == len(self.categories) - 1:
            symbol_data = self.custom_symbols[self.selected_option].copy()
            symbol_data['category'] = 'custom'
            return symbol_data

        return {
            'category': category.name.lower(),
            'display': symbol_display,
            'type': 'builtin',
            'index': self.selected_option
        }
    
class CommentDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.comment_text = ""
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Add Comment")
        self.setGeometry(300, 300, 500, 200)
        
        layout = QVBoxLayout(self)
        
        instructions = QLabel("Enter your comment (will be formatted as ((comment))):")
        instructions.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions)
        
        self.comment_edit = QLineEdit()
        self.comment_edit.setPlaceholderText("Enter comment here")
        layout.addWidget(self.comment_edit)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        self.comment_edit.setFocus()
        
    def get_comment(self):
        return f"(({self.comment_edit.text()}))"

class EditDialog(QDialog):
    def __init__(self, current_text, parent=None):
        super().__init__(parent)
        self.current_text = current_text
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Edit Segment Content")
        self.setGeometry(300, 300, 600, 150)
        
        layout = QVBoxLayout(self)
        
        instructions = QLabel("Edit the segment content (Enter to confirm, Escape to cancel):")
        instructions.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions)
        
        self.text_edit = QLineEdit()
        self.text_edit.setText(self.current_text)
        self.text_edit.setStyleSheet("""
            QLineEdit {
                font-family: monospace;
                font-size: 14px;
                padding: 8px;
                border: 2px solid #ccc;
                border-radius: 5px;
            }
            QLineEdit:focus {
                border: 2px solid #4a90e2;
            }
        """)
        
        self.text_edit.returnPressed.connect(self.accept)
        
        layout.addWidget(self.text_edit)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        for button in button_box.buttons():
            button.setFocusPolicy(Qt.NoFocus)
            
        layout.addWidget(button_box)
        
        self.text_edit.setFocus()
        self.text_edit.selectAll()
        
    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Escape:
            self.reject()
            event.accept()
        else:
            super().keyPressEvent(event)
    
    def get_text(self):
        return self.text_edit.text()

class SettingsDialog(QDialog):
    def __init__(self, current_font, current_theme, cjk_mode, parent=None):
        super().__init__(parent)
        self.selected_font = current_font
        self.current_theme = current_theme
        self.cjk_mode = cjk_mode
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Settings")
        self.setGeometry(100, 100, 400, 200)
        
        layout = QVBoxLayout(self)
        
        # Font selection
        font_layout = QHBoxLayout()
        font_layout.addWidget(QLabel("Text Display Font:"))
        self.font_button = QPushButton(f"{self.selected_font.family()} {self.selected_font.pointSize()}pt")
        self.font_button.clicked.connect(self.select_font)
        font_layout.addWidget(self.font_button)
        font_layout.addStretch()
        layout.addLayout(font_layout)
        
        # Theme selection
        theme_layout = QHBoxLayout()
        theme_layout.addWidget(QLabel("Viewer Theme:"))
        self.theme_combo = QComboBox()
        self.theme_combo.addItems(["Light", "Dark"])
        self.theme_combo.setCurrentText(self.current_theme.capitalize())
        theme_layout.addWidget(self.theme_combo)
        theme_layout.addStretch()
        layout.addLayout(theme_layout)
        
        # CJK optimization checkbox
        cjk_layout = QHBoxLayout()
        self.cjk_checkbox = QCheckBox("Optimize for CJK (double spaces for overlap indentation)")
        self.cjk_checkbox.setChecked(self.cjk_mode)
        cjk_layout.addWidget(self.cjk_checkbox)
        layout.addLayout(cjk_layout)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
    def select_font(self):
        font, ok = QFontDialog.getFont(self.selected_font, self)
        if ok:
            self.selected_font = font
            self.font_button.setText(f"{font.family()} {font.pointSize()}pt")
    
    def get_font(self):
        return self.selected_font
    
    def get_theme(self):
        return self.theme_combo.currentText().lower()
    
    def get_cjk_mode(self):
        return self.cjk_checkbox.isChecked()

class ProjectMemoDialog(QDialog):
    def __init__(self, project_name="", project_memo="", parent=None):
        super().__init__(parent)
        self.project_name = project_name
        self.project_memo = project_memo
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Project Memo")
        self.setGeometry(300, 300, 500, 400)
        
        layout = QVBoxLayout(self)
        
        # Project name
        layout.addWidget(QLabel("Project Name:"))
        self.name_edit = QLineEdit(self.project_name)
        self.name_edit.setPlaceholderText("Enter project name")
        layout.addWidget(self.name_edit)
        
        # Project memo
        layout.addWidget(QLabel("Project Memo:"))
        self.memo_edit = QPlainTextEdit()
        self.memo_edit.setPlainText(self.project_memo)
        self.memo_edit.setPlaceholderText("Enter project notes or description")
        layout.addWidget(self.memo_edit)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
    def get_project_info(self):
        return {
            'name': self.name_edit.text(),
            'memo': self.memo_edit.toPlainText()
        }

class JsonImportDialog(QDialog):
    def __init__(self, has_tokens=False, parent=None):
        super().__init__(parent)
        self.import_option = "one_block"  # one_block, tokens, auto_segment
        self.has_tokens = has_tokens
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("JSON Import Options")
        self.setGeometry(300, 300, 500, 250)
        
        layout = QVBoxLayout(self)
        
        instructions = QLabel("How would you like to import this JSON file?")
        instructions.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions)
        
        self.one_block_radio = QRadioButton("Import as one continuous block")
        self.one_block_radio.setChecked(True)
        layout.addWidget(self.one_block_radio)
        
        if self.has_tokens:
            self.tokens_radio = QRadioButton("Import each token as separate block")
            layout.addWidget(self.tokens_radio)
            
            self.auto_radio = QRadioButton("Auto-segment based on pause detection")
            layout.addWidget(self.auto_radio)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
    def get_import_option(self):
        if self.has_tokens:
            if self.one_block_radio.isChecked():
                return "one_block"
            elif self.tokens_radio.isChecked():
                return "tokens"
            else:
                return "auto_segment"
        else:
            return "one_block"

class UnassignedSegmentsDialog(QDialog):
    def __init__(self, unassigned_count, parent=None):
        super().__init__(parent)
        self.unassigned_count = unassigned_count
        self.selected_option = "skip"  # skip, no_label, unknown
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Unassigned Segments")
        self.setGeometry(300, 300, 450, 200)
        
        layout = QVBoxLayout(self)
        
        instructions = QLabel(f"Found {self.unassigned_count} unassigned segment(s). How should they be handled?")
        instructions.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions)
        
        # Radio buttons for options
        self.skip_radio = QRadioButton("Do not include in SRT file")
        self.skip_radio.setChecked(True)
        layout.addWidget(self.skip_radio)
        
        self.no_label_radio = QRadioButton("Include without speaker label")
        layout.addWidget(self.no_label_radio)
        
        self.unknown_radio = QRadioButton("Label as 'Unknown:'")
        layout.addWidget(self.unknown_radio)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
    def get_selected_option(self):
        if self.skip_radio.isChecked():
            return "skip"
        elif self.no_label_radio.isChecked():
            return "no_label"
        else:
            return "unknown"

class ExportPreviewDialog(QDialog):
    def __init__(self, parent=None, has_timestamps=True,
                 timestamp_style="curly", custom_pattern="{HH:MM:SS}",
                 project_info=None, audio_path=None):
        super().__init__(parent)
        self.main_window = parent
        self.include_timestamps = has_timestamps
        self.current_include_timestamps = has_timestamps
        self.export_format = "html"
        self.transcript_convention = "gat2"
        self.project_info = project_info or {}
        self.audio_path = audio_path
        self.timestamp_style = timestamp_style
        self.custom_timestamp_pattern = custom_pattern
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Export Preview")
        self.setGeometry(100, 100, 850, 750)

        layout = QVBoxLayout(self)

        # ----- Format selection -----
        format_layout = QHBoxLayout()
        format_layout.addWidget(QLabel("Export Format:"))

        self.html_radio = QRadioButton("HTML (.html)")
        self.html_radio.setChecked(True)
        format_layout.addWidget(self.html_radio)

        self.docx_radio = QRadioButton("Word Document (.docx)")
        format_layout.addWidget(self.docx_radio)

        self.txt_radio = QRadioButton("Plain Text (.txt)")
        format_layout.addWidget(self.txt_radio)

        self.srt_radio = QRadioButton("Subtitle File (.srt)")
        format_layout.addWidget(self.srt_radio)
        format_layout.addStretch()

        if not self.include_timestamps:
            self.srt_radio.setEnabled(False)
            self.srt_radio.setToolTip("SRT export requires timestamp information. Original file does not contain timestamps.")

        # ----- Convention selection -----
        convention_layout = QHBoxLayout()
        convention_layout.addWidget(QLabel("Transcript Convention:"))

        self.convention_combo = QComboBox()
        self.convention_combo.addItems([
            "GAT2 (Conversation Analysis)",
            "Dresing & Pehl (Semantic Transcription)",
            "TiQ (Talk in Qualitative Research)"
        ])
        convention_layout.addWidget(self.convention_combo)
        convention_layout.addStretch()

        # ----- Line wrapping group -----
        self.wrap_group = QGroupBox("Line Wrapping")
        wrap_layout = QHBoxLayout()

        self.wrap_check = QCheckBox("Wrap lines at:")
        self.wrap_spin = QSpinBox()
        self.wrap_spin.setRange(30, 200)
        self.wrap_spin.setValue(80)
        self.wrap_spin.setSuffix(" characters")
        self.wrap_spin.setEnabled(False)

        self.character_wrap_check = QCheckBox("Force character‑based wrapping")
        self.character_wrap_check.setEnabled(False)

        wrap_layout.addWidget(self.wrap_check)
        wrap_layout.addWidget(self.wrap_spin)
        wrap_layout.addWidget(self.character_wrap_check)
        wrap_layout.addStretch()
        self.wrap_group.setLayout(wrap_layout)

        # ----- Options group -----
        options_group = QGroupBox("Export Options")
        options_layout = QVBoxLayout()

        # First, create all option checkboxes
        self.timestamp_check = QCheckBox("Include timestamps")
        self.timestamp_check.setChecked(self.include_timestamps)
        self.timestamp_check.setEnabled(self.include_timestamps)

        self.diarization_check = QCheckBox("Include diarization (speaker labels)")
        self.diarization_check.setChecked(True)

        self.title_check = QCheckBox("Include project title")
        self.title_check.setChecked(True)

        self.memo_check = QCheckBox("Include project memo")
        self.memo_check.setChecked(True)

        self.audio_check = QCheckBox("Include audio file path")
        self.audio_check.setChecked(True)

        # ----- Timestamp format widgets (placed inside a container widget) -----
        self.ts_format_widget = QWidget()
        ts_format_layout = QHBoxLayout(self.ts_format_widget)
        ts_format_layout.setContentsMargins(0, 0, 0, 0)

        self.ts_curly_radio = QRadioButton("curly brackets")
        self.ts_hash_radio = QRadioButton("hash")
        self.ts_bracket_radio = QRadioButton("brackets")
        self.ts_custom_radio = QRadioButton("custom")

        ts_format_layout.addWidget(self.ts_curly_radio)
        ts_format_layout.addWidget(self.ts_hash_radio)
        ts_format_layout.addWidget(self.ts_bracket_radio)
        ts_format_layout.addWidget(self.ts_custom_radio)

        self.ts_custom_edit = QLineEdit()
        self.ts_custom_edit.setPlaceholderText("e.g. <HH:MM:SS-xx>")
        self.ts_custom_edit.setEnabled(False)
        ts_format_layout.addWidget(self.ts_custom_edit)
        ts_format_layout.addStretch()

        # ----- Assemble the timestamp line (checkbox + format widget) -----
        ts_line_layout = QHBoxLayout()
        ts_line_layout.addWidget(self.timestamp_check)
        ts_line_layout.addWidget(self.ts_format_widget)
        ts_line_layout.addStretch()

        # Add everything to options layout
        options_layout.addLayout(ts_line_layout)
        options_layout.addWidget(self.diarization_check)
        options_layout.addWidget(self.title_check)
        options_layout.addWidget(self.memo_check)
        options_layout.addWidget(self.audio_check)

        options_group.setLayout(options_layout)

        # ----- Preview area -----
        preview_label = QLabel("Preview:")
        preview_label.setFont(QFont("Arial", 12, QFont.Bold))

        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)

        # ----- Buttons -----
        button_box = QDialogButtonBox(QDialogButtonBox.Save | QDialogButtonBox.Cancel)

        # ----- Assemble main layout -----
        layout.addLayout(format_layout)
        layout.addLayout(convention_layout)
        layout.addWidget(self.wrap_group)
        layout.addWidget(options_group)
        layout.addWidget(preview_label)
        layout.addWidget(self.preview_text)
        layout.addWidget(button_box)

        # ----- Connect signals (all widgets now exist) -----
        self.html_radio.toggled.connect(self.on_format_changed)
        self.docx_radio.toggled.connect(self.on_format_changed)
        self.txt_radio.toggled.connect(self.on_format_changed)
        self.srt_radio.toggled.connect(self.on_format_changed)

        self.convention_combo.currentTextChanged.connect(self.on_convention_changed)

        self.wrap_check.toggled.connect(self.on_wrap_toggled)
        self.wrap_check.toggled.connect(self.update_preview)
        self.wrap_spin.valueChanged.connect(self.update_preview)
        self.character_wrap_check.toggled.connect(self.update_preview)

        self.timestamp_check.toggled.connect(self.on_timestamp_changed)

        self.ts_curly_radio.toggled.connect(self.on_timestamp_format_changed)
        self.ts_hash_radio.toggled.connect(self.on_timestamp_format_changed)
        self.ts_bracket_radio.toggled.connect(self.on_timestamp_format_changed)
        self.ts_custom_radio.toggled.connect(self.on_timestamp_format_changed)
        self.ts_custom_edit.textChanged.connect(self.update_preview)

        self.diarization_check.toggled.connect(self.update_preview)
        self.title_check.toggled.connect(self.update_preview)
        self.memo_check.toggled.connect(self.update_preview)
        self.audio_check.toggled.connect(self.update_preview)

        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)

        # ----- Set initial states -----
        # Timestamp style from saved settings
        if self.timestamp_style == "curly":
            self.ts_curly_radio.setChecked(True)
        elif self.timestamp_style == "hash":
            self.ts_hash_radio.setChecked(True)
        elif self.timestamp_style == "bracket":
            self.ts_bracket_radio.setChecked(True)
        else:  # custom
            self.ts_custom_radio.setChecked(True)
            self.ts_custom_edit.setText(self.custom_timestamp_pattern)

        # Initial format update
        self.on_format_changed()

    def on_wrap_toggled(self, checked):
        self.wrap_spin.setEnabled(checked)
        self.character_wrap_check.setEnabled(checked)

    def on_format_changed(self):
        if self.html_radio.isChecked():
            self.export_format = "html"
        elif self.docx_radio.isChecked():
            self.export_format = "docx"
        elif self.txt_radio.isChecked():
            self.export_format = "txt"
        else:
            self.export_format = "srt"

        is_srt = (self.export_format == "srt")

        # Project info options
        self.title_check.setEnabled(not is_srt)
        self.memo_check.setEnabled(not is_srt)
        self.audio_check.setEnabled(not is_srt)

        # Timestamp controls
        self.timestamp_check.setEnabled(not is_srt and self.include_timestamps)

        # Diarization
        if is_srt:
            self.diarization_check.setEnabled(True)
            self.diarization_check.setChecked(True)
        else:
            self.diarization_check.setEnabled(False)
            self.diarization_check.setChecked(True)

        # Timestamp format widget: disabled for SRT, otherwise depends on timestamp checkbox
        if is_srt:
            self.ts_format_widget.setEnabled(False)
        else:
            self.ts_format_widget.setEnabled(self.timestamp_check.isChecked())

        # Line wrapping group: disabled for SRT, otherwise enabled (individual widgets controlled by convention)
        self.wrap_group.setEnabled(not is_srt)

        self.update_preview()

    def on_convention_changed(self, convention_text):
        if "Dresing" in convention_text:
            self.transcript_convention = "dresing_pehl"
            self.wrap_check.setChecked(False)
            self.wrap_check.setEnabled(False)
            self.wrap_spin.setEnabled(False)
            self.set_timestamp_style("hash")
        elif "TiQ" in convention_text:
            self.transcript_convention = "tiq"
            self.wrap_check.setChecked(True)
            self.wrap_check.setEnabled(False)
            self.wrap_spin.setEnabled(True)
            self.set_timestamp_style("hash")
        else:
            self.transcript_convention = "gat2"
            self.wrap_check.setEnabled(True)
            self.wrap_spin.setEnabled(self.wrap_check.isChecked())
            self.set_timestamp_style("curly")
            
        self.on_wrap_toggled(self.wrap_check.isChecked())

        if self.export_format != "srt":
            self.diarization_check.setEnabled(False)
            self.diarization_check.setChecked(True)

        self.update_preview()

    def set_timestamp_style(self, style):
        """Helper to set the radio buttons based on style name."""
        if style == "curly":
            self.ts_curly_radio.setChecked(True)
        elif style == "hash":
            self.ts_hash_radio.setChecked(True)
        elif style == "bracket":
            self.ts_bracket_radio.setChecked(True)
        else:
            self.ts_custom_radio.setChecked(True)
            # keep existing custom text

    def on_timestamp_format_changed(self):
        self.ts_custom_edit.setEnabled(self.ts_custom_radio.isChecked())
        self.update_preview()

    def on_timestamp_changed(self, checked):
        self.current_include_timestamps = checked
        # Enable/disable the format widget only if not SRT
        if self.export_format != "srt":
            self.ts_format_widget.setEnabled(checked)
        self.ts_custom_edit.setEnabled(checked and self.ts_custom_radio.isChecked() and self.export_format != "srt")
        QTimer.singleShot(100, self.update_preview)

    def update_preview(self):
        main = self.main_window
        if not main:
            return

        # Determine current timestamp style and pattern
        if self.ts_curly_radio.isChecked():
            ts_style = "curly"
            custom = None
        elif self.ts_hash_radio.isChecked():
            ts_style = "hash"
            custom = None
        elif self.ts_bracket_radio.isChecked():
            ts_style = "bracket"
            custom = None
        else:
            ts_style = "custom"
            custom = self.ts_custom_edit.text()

        if self.export_format == "srt":
            srt_text = main.generate_srt_text(
                include_diarization=self.diarization_check.isChecked(),
                unassigned_handling="skip"
            )
            self.preview_text.setPlainText(srt_text)
            return

        if self.export_format == "docx":
            self.preview_text.setPlainText("Preview not available for this format. The exported file will contain the full transcript with selected options.")
            return

        # Generate transcript text
        transcript_text = main.generate_transcript_text(
            include_timestamps=self.current_include_timestamps,
            timestamp_style=ts_style,
            custom_pattern=custom,
            convention=self.transcript_convention,
            include_diarization=self.diarization_check.isChecked(),
            wrap_enabled=self.wrap_check.isChecked(),
            wrap_length=self.wrap_spin.value(),
            character_wrap=self.character_wrap_check.isChecked()
        )

        # Build header
        header_lines = []
        if self.title_check.isChecked() and main.project_name:
            if self.export_format == "html":
                header_lines.append(f"<h1>{main.escape_html(main.project_name)}</h1>")
            else:
                header_lines.append(main.project_name)
                header_lines.append("=" * len(main.project_name))
                header_lines.append("")
        if self.memo_check.isChecked() and main.project_memo:
            if self.export_format == "html":
                header_lines.append(f'<p class="headerstyle"><strong>Project Memo:</strong> {main.escape_html(main.project_memo)}</p>')
            else:
                header_lines.append(f"Project Memo: {main.project_memo}")
                header_lines.append("")
        if self.audio_check.isChecked() and main.audio_file_path:
            audio_name = Path(main.audio_file_path).name
            if self.export_format == "html":
                header_lines.append(f'<p class="headerstyle"><strong>Audio File:</strong> {main.escape_html(audio_name)}</p>')
            else:
                header_lines.append(f"Audio File: {audio_name}")
                header_lines.append("")

        header = "\n".join(header_lines) + "\n" if header_lines else ""

        if self.export_format == "html":
            font_family = "'Courier New', monospace"
            if self.transcript_convention == "dresing_pehl":
                font_family = "'Times New Roman', serif"

            full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
body {{
    font-family: {font_family};
    font-size: 10pt;
    line-height: 1.2;
    margin: 0;
    padding: 10px;
    white-space: pre-wrap;
}}
h1 {{
    font-family: Arial, sans-serif;
    color: #333;
    padding-bottom: 10px;
    margin-top: 0;
}}
.headerstyle {{
    font-family: Arial, sans-serif;
    color: #333;
}}
</style>
</head>
<body>
{header}<br>{main.escape_html(transcript_text)}
</body>
</html>"""
            self.preview_text.setHtml(full_html)
        else:  # TXT
            self.preview_text.setPlainText(header + transcript_text)

    def get_export_settings(self):
        if self.ts_curly_radio.isChecked():
            ts_style = "curly"
        elif self.ts_hash_radio.isChecked():
            ts_style = "hash"
        elif self.ts_bracket_radio.isChecked():
            ts_style = "bracket"
        else:
            ts_style = "custom"

        custom_pattern = self.ts_custom_edit.text() if ts_style == "custom" else ""

        return {
            'format': self.export_format,
            'convention': self.transcript_convention,
            'include_timestamps': self.current_include_timestamps,
            'include_diarization': self.diarization_check.isChecked(),
            'include_title': self.title_check.isChecked(),
            'include_memo': self.memo_check.isChecked(),
            'include_audio': self.audio_check.isChecked(),
            'wrap_enabled': self.wrap_check.isChecked(),
            'wrap_length': self.wrap_spin.value(),
            'character_wrap': self.character_wrap_check.isChecked(),
            'timestamp_style': ts_style,
            'custom_timestamp_pattern': custom_pattern
        }
    
class SearchDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.current_match = 0
        self.total_matches = 0
        self.match_positions = []  # List of (block_index, start_pos, end_pos)
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Search")
        self.setGeometry(300, 200, 500, 150)
        
        layout = QVBoxLayout(self)
        
        # Search input
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("Search:"))
        
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Enter search term")
        self.search_input.textChanged.connect(self.perform_search)
        search_layout.addWidget(self.search_input)
        
        self.case_sensitive_check = QCheckBox("Case sensitive")
        self.case_sensitive_check.toggled.connect(self.perform_search)
        search_layout.addWidget(self.case_sensitive_check)
        
        layout.addLayout(search_layout)
        
        # Match information and navigation
        nav_layout = QHBoxLayout()
        
        self.match_label = QLabel("0 matches")
        nav_layout.addWidget(self.match_label)
        
        self.prev_button = QPushButton("◀ Previous")
        self.prev_button.clicked.connect(self.previous_match)
        self.prev_button.setEnabled(False)
        nav_layout.addWidget(self.prev_button)
        
        self.next_button = QPushButton("Next ▶")
        self.next_button.clicked.connect(self.next_match)
        self.next_button.setEnabled(False)
        nav_layout.addWidget(self.next_button)
        
        layout.addLayout(nav_layout)
        
        # Close button
        button_box = QDialogButtonBox(QDialogButtonBox.Close)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        self.setFocus()
        self.search_input.setFocus()
        
    def perform_search(self):
        search_term = self.search_input.text()
        case_sensitive = self.case_sensitive_check.isChecked()
        
        if not search_term:
            self.match_label.setText("0 matches")
            self.prev_button.setEnabled(False)
            self.next_button.setEnabled(False)
            self.match_positions = []
            self.total_matches = 0
            if self.parent:
                self.parent.clear_search_highlights()
            return
            
        self.match_positions = []
        
        # Search through all blocks
        for block_idx, block in enumerate(self.parent.srt_blocks):
            text = block['text']
            if not text:
                continue
                
            if case_sensitive:
                search_in = text
            else:
                search_in = text.lower()
                search_term_lower = search_term.lower()
            
            start_pos = 0
            while True:
                if case_sensitive:
                    pos = search_in.find(search_term, start_pos)
                else:
                    pos = search_in.find(search_term_lower, start_pos)
                    
                if pos == -1:
                    break
                    
                self.match_positions.append((block_idx, pos, pos + len(search_term)))
                start_pos = pos + 1
        
        self.total_matches = len(self.match_positions)
        self.current_match = 0
        
        if self.total_matches > 0:
            self.match_label.setText(f"{self.current_match + 1}/{self.total_matches} matches")
            self.prev_button.setEnabled(True)
            self.next_button.setEnabled(True)
            self.highlight_current_match()
        else:
            self.match_label.setText("0 matches")
            self.prev_button.setEnabled(False)
            self.next_button.setEnabled(False)
            
    def highlight_current_match(self):
        if not self.match_positions or self.current_match >= len(self.match_positions):
            return
            
        block_idx, start_pos, end_pos = self.match_positions[self.current_match]
        
        if self.parent:
            self.parent.highlight_search_match(block_idx, start_pos, end_pos)
            
    def next_match(self):
        if self.total_matches == 0:
            return
            
        self.current_match = (self.current_match + 1) % self.total_matches
        self.match_label.setText(f"{self.current_match + 1}/{self.total_matches} matches")
        self.highlight_current_match()
        
    def previous_match(self):
        if self.total_matches == 0:
            return
            
        self.current_match = (self.current_match - 1) % self.total_matches
        self.match_label.setText(f"{self.current_match + 1}/{self.total_matches} matches")
        self.highlight_current_match()
        
    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Escape:
            self.reject()
            event.accept()
        elif event.key() == Qt.Key_F3 and event.modifiers() & Qt.ShiftModifier:
            self.previous_match()
            event.accept()
        elif event.key() == Qt.Key_F3:
            self.next_match()
            event.accept()
        elif event.key() == Qt.Key_G and event.modifiers() == Qt.ControlModifier:
            self.next_match()
            event.accept()
        else:
            super().keyPressEvent(event)

class JumpToTimeDialog(QDialog):
    def __init__(self, max_duration_ms, parent=None):
        super().__init__(parent)
        self.max_duration_ms = max_duration_ms
        self.target_time_ms = 0
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Jump to Time")
        self.setGeometry(300, 300, 400, 200)
        
        layout = QVBoxLayout(self)
        
        instructions = QLabel("Enter time to jump to (format: MM:SS or HH:MM:SS):")
        instructions.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions)
        
        self.time_edit = QLineEdit()
        self.time_edit.setPlaceholderText("e.g., 1:30 or 0:01:30")
        layout.addWidget(self.time_edit)
        
        # Current max time display
        max_minutes = self.max_duration_ms // 60000
        max_seconds = (self.max_duration_ms % 60000) // 1000
        max_hours = max_minutes // 60
        max_minutes = max_minutes % 60
        
        max_time_label = QLabel(f"Maximum time: {max_hours:02d}:{max_minutes:02d}:{max_seconds:02d}")
        layout.addWidget(max_time_label)
        
        self.error_label = QLabel("")
        self.error_label.setStyleSheet("color: red;")
        layout.addWidget(self.error_label)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.validate_and_accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
        
        self.time_edit.setFocus()
        
    def validate_and_accept(self):
        time_str = self.time_edit.text().strip()
        if not time_str:
            self.error_label.setText("Please enter a time")
            return
            
        # Parse time format (MM:SS or HH:MM:SS)
        parts = time_str.split(':')
        if len(parts) == 2:
            # MM:SS format
            try:
                minutes = int(parts[0])
                seconds = int(parts[1])
                hours = 0
            except ValueError:
                self.error_label.setText("Invalid time format")
                return
        elif len(parts) == 3:
            # HH:MM:SS format
            try:
                hours = int(parts[0])
                minutes = int(parts[1])
                seconds = int(parts[2])
            except ValueError:
                self.error_label.setText("Invalid time format")
                return
        else:
            self.error_label.setText("Use MM:SS or HH:MM:SS format")
            return
            
        # Validate time values
        if minutes >= 60 or seconds >= 60:
            self.error_label.setText("Minutes and seconds must be less than 60")
            return
            
        self.target_time_ms = (hours * 3600 + minutes * 60 + seconds) * 1000
        
        if self.target_time_ms > self.max_duration_ms:
            self.error_label.setText("Time exceeds audio duration")
            return
            
        self.accept()
        
    def get_target_time(self):
        return self.target_time_ms
    
class EnhancedPlacementDialog(QDialog):
    """Enhanced placement dialog that handles different symbol types"""
    
    def __init__(self, current_text, symbol_info, parent=None):
        super().__init__(parent)
        self.current_text = current_text
        self.symbol_info = symbol_info
        self.placement_position = 0
        self.create_new_line = False
        self.selected_text_start = 0
        self.selected_text_end = 0
        self.has_selection = False
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle(f"Insert {self.symbol_info.get('category', 'Symbol').title()}")
        self.setGeometry(300, 300, 700, 400)
        
        layout = QVBoxLayout(self)
        
        # Instructions based on symbol type
        instructions = self.get_instructions()
        instructions_label = QLabel(instructions)
        instructions_label.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions_label)
        
        # Text display with selection capability
        self.text_display = QTextEdit()
        self.text_display.setPlainText(self.current_text)
        self.text_display.setMaximumHeight(150)
        self.text_display.setStyleSheet("font-family: monospace; font-size: 14px;")
        layout.addWidget(self.text_display)
        
        # Selection info
        self.selection_label = QLabel("No text selected")
        layout.addWidget(self.selection_label)
        
        # Options
        self.option_label = QLabel("")
        self.update_option_label()
        layout.addWidget(self.option_label)
        
        # Button box
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        for button in button_box.buttons():
            button.setFocusPolicy(Qt.NoFocus)
            
        layout.addWidget(button_box)
        
        # Connect signals
        self.text_display.selectionChanged.connect(self.on_selection_changed)
        
        self.setFocusPolicy(Qt.StrongFocus)
        self.setFocus()
        self.text_display.setFocus()
        
    def get_instructions(self):
        """Get instructions based on symbol type"""
        category = self.symbol_info.get('category', '').lower()
        
        if 'dresing' in category:
            return "Dresing & Pehl: Use ← → to position, select text with mouse, Enter to confirm, N for new line"
        elif 'tiq' in category:
            return "TiQ: Use ← → to position, select text with mouse, Enter to confirm, N for new line"
        elif 'custom' in category:
            symbol_type = self.symbol_info.get('type', 'simple')
            if symbol_type in ['wrapper', 'comment', 'comment_reach']:
                return f"Select the text to wrap with {self.symbol_info.get('display', 'symbol')}"
            else:
                return "Use ← → to position, Enter to confirm, N for new line"
        else:  # GAT2
            return "Use ← → to position, select text with mouse, Enter to confirm, N for new line"
    
    def on_selection_changed(self):
        """Handle text selection changes"""
        cursor = self.text_display.textCursor()
        if cursor.hasSelection():
            self.has_selection = True
            self.selected_text_start = cursor.selectionStart()
            self.selected_text_end = cursor.selectionEnd()
            selected_text = cursor.selectedText()
            self.selection_label.setText(f"Selected: '{selected_text}'")
        else:
            self.has_selection = False
            self.selection_label.setText("No text selected")
    
    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Left and not event.modifiers():
            # Move cursor left
            cursor = self.text_display.textCursor()
            cursor.movePosition(cursor.Left)
            self.text_display.setTextCursor(cursor)
            event.accept()
        elif event.key() == Qt.Key_Right and not event.modifiers():
            # Move cursor right
            cursor = self.text_display.textCursor()
            cursor.movePosition(cursor.Right)
            self.text_display.setTextCursor(cursor)
            event.accept()
        elif event.key() == Qt.Key_N:
            self.create_new_line = not self.create_new_line
            self.update_option_label()
            event.accept()
        elif event.key() in (Qt.Key_Return, Qt.Key_Enter):
            if event.modifiers() & Qt.ControlModifier:
                # Ctrl+Enter to accept
                self.accept()
                event.accept()
            else:
                # Regular Enter passes to text edit for new line
                super().keyPressEvent(event)
        else:
            super().keyPressEvent(event)
    
    def update_option_label(self):
        if self.create_new_line:
            self.option_label.setText("Option: Create new line (Press N for inline)")
        else:
            self.option_label.setText("Option: Insert in current line (Press N for new line)")
    
    def get_placement_result(self):
        """Get the result of placement based on symbol type"""
        symbol_type = self.symbol_info.get('type', 'builtin')
        display = self.symbol_info.get('display', '')
        
        # New line option (only allowed for simple symbols)
        if self.create_new_line:
            return {
                'action': 'new_line',
                'text': display
            }
        
        current_text = self.text_display.toPlainText()
        
        if self.has_selection:
            selected_text = current_text[self.selected_text_start:self.selected_text_end]
            
            # Handle wrapper symbols (custom or built-in wrappers)
            if symbol_type in ['wrapper', 'comment', 'comment_reach']:
                left = self.symbol_info.get('left', '')
                right = self.symbol_info.get('right', '')
                # For built-in wrappers without explicit left/right, derive from display
                if not left and display == '@(   )@':
                    left = '@('
                    right = ')@'
                elif not left and display == '°   °':
                    left = '°'
                    right = '°'
                elif not left and display == '//   //':
                    left = '//'
                    right = '//'
                new_text = (current_text[:self.selected_text_start] + left + selected_text + right +
                           current_text[self.selected_text_end:])
                return {'action': 'replace', 'text': new_text}
            else:
                # For simple symbols, ignore selection and insert at cursor position
                cursor = self.text_display.textCursor()
                pos = cursor.position()
                new_text = current_text[:pos] + " " + display + " " + current_text[pos:]
                return {'action': 'replace', 'text': new_text}
        else:
            # No selection: insert at cursor
            cursor = self.text_display.textCursor()
            pos = cursor.position()
            new_text = current_text[:pos] + " " + display + " " + current_text[pos:]
            return {'action': 'replace', 'text': new_text}
    
class PlacementDialog(QDialog):
    def __init__(self, current_text, symbol, parent=None):
        super().__init__(parent)
        self.current_text = current_text
        self.symbol = symbol
        self.placement_position = 0
        self.create_new_line = False
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Place Symbol")
        self.setGeometry(300, 300, 600, 300)
        
        layout = QVBoxLayout(self)
        
        instructions = QLabel("Use ← → arrows to position, Enter to confirm, N for new line:")
        instructions.setStyleSheet("font-weight: bold;")
        layout.addWidget(instructions)
        
        self.text_display = QTextEdit()
        self.text_display.setReadOnly(True)
        self.text_display.setMaximumHeight(100)
        self.text_display.setStyleSheet("font-family: monospace; font-size: 14px;")
        layout.addWidget(self.text_display)
        
        self.option_label = QLabel("Placement: Insert in current line (Press N to create new line)")
        layout.addWidget(self.option_label)
        
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        for button in button_box.buttons():
            button.setFocusPolicy(Qt.NoFocus)
            
        layout.addWidget(button_box)
        
        self.update_display()
        self.setFocusPolicy(Qt.StrongFocus)
        self.setFocus()
        
    def keyPressEvent(self, event):
        if event.key() == Qt.Key_Left:
            self.placement_position = max(0, self.placement_position - 1)
            self.update_display()
            event.accept()
        elif event.key() == Qt.Key_Right:
            self.placement_position = min(len(self.current_text), self.placement_position + 1)
            self.update_display()
            event.accept()
        elif event.key() == Qt.Key_N:
            self.create_new_line = not self.create_new_line
            self.update_display()
            event.accept()
        elif event.key() in (Qt.Key_Return, Qt.Key_Enter):
            self.accept()
            event.accept()
        else:
            super().keyPressEvent(event)
    
    def update_display(self):
        if self.create_new_line:
            html_content = f"""
            <div style="font-family: monospace; font-size: 14px; padding: 10px;">
                <span style="background-color: #e0e0e0; padding: 5px; border-radius: 3px;">{self.current_text}</span><br>
                <span style="background-color: #c8f7c8; padding: 5px; border-radius: 3px;">{self.symbol}</span>
            </div>
            """
            self.option_label.setText("Placement: Create new line with symbol (Press N for inline)")
        else:
            before_text = self.current_text[:self.placement_position]
            after_text = self.current_text[self.placement_position:]
            html_content = f"""
            <div style="font-family: monospace; font-size: 14px; padding: 10px;">
                <span style="background-color: #e0e0e0; padding: 5px; border-radius: 3px;">{before_text}</span>
                <span style="background-color: #c8f7c8; padding: 5px; border-radius: 3px;">{self.symbol}</span>
                <span style="background-color: #e0e0e0; padding: 5px; border-radius: 3px;">{after_text}</span>
            </div>
            """
            self.option_label.setText("Placement: Insert in current line (Press N to create new line)")
        
        self.text_display.setHtml(html_content)

class SpeedKnob(QWidget):
    valueChanged = pyqtSignal(float)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.value = 1.0  # Default 1.0x speed
        self.min_value = 0.5
        self.max_value = 2.0
        self.step = 0.1
        self.is_dragging = False
        self.last_mouse_pos = None
        self.setMinimumSize(60, 60)
        
    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        
        # Draw knob background
        center = QPoint(self.width() // 2, self.height() // 2)
        radius = min(self.width(), self.height()) // 2 - 5
        
        # Outer circle
        painter.setBrush(QColor(240, 240, 240))
        painter.setPen(QPen(QColor(180, 180, 180), 2))
        painter.drawEllipse(center, radius, radius)
        
        # Value indicator
        angle = 142 + (self.value - self.min_value) / (self.max_value - self.min_value) * 270
        angle_rad = math.radians(angle)
        
        indicator_length = radius - 5
        end_x = center.x() + indicator_length * math.cos(angle_rad)
        end_y = center.y() + indicator_length * math.sin(angle_rad)
        
        painter.setPen(QPen(QColor(0, 120, 215), 2))
        painter.drawLine(center, QPoint(int(end_x), int(end_y)))
        
        # Draw center dot
        painter.setBrush(QColor(0, 120, 215))
        painter.setPen(Qt.NoPen)
        painter.drawEllipse(center, 4, 4)
        
        # Draw value text
        painter.setPen(QColor(50, 50, 50))
        font = painter.font()
        font.setPointSize(9)
        painter.setFont(font)
        value_text = f"{self.value:.1f}x"
        text_rect = QRect(center.x() - 30, center.y() - -2, 60, 20)
        painter.drawText(text_rect, Qt.AlignCenter, value_text)
        
        # Draw min/max labels
        font.setPointSize(7)
        painter.setFont(font)
        painter.drawText(center.x() - 40, center.y() - radius + 50, "0.5x")
        painter.drawText(center.x() + 27, center.y() - radius + 50, "2.0x")
        
    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.is_dragging = True
            self.last_mouse_pos = event.pos()
            event.accept()
            
    def set_value_direct(self, value):
        """Set value directly"""
        new_value = max(self.min_value, min(self.max_value, value))
        if abs(new_value - self.value) > 0.01:
            self.value = new_value
            self.update()
            self.valueChanged.emit(self.value)
            
    def mouseMoveEvent(self, event):
        if self.is_dragging and self.last_mouse_pos:
            delta_y = self.last_mouse_pos.y() - event.y()
            delta_x = event.x() - self.last_mouse_pos.x()
            delta = delta_y + delta_x
            
            new_value = self.value + (delta * self.step / 30)
            new_value = max(self.min_value, min(self.max_value, new_value))
            
            if new_value != self.value:
                self.value = round(new_value / self.step) * self.step
                self.value = max(self.min_value, min(self.max_value, self.value))
                self.update()
                self.valueChanged.emit(self.value)
                
            self.last_mouse_pos = event.pos()
            event.accept()
            
    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.is_dragging = False
            self.last_mouse_pos = None
            event.accept()
            
    def wheelEvent(self, event):
        delta = event.angleDelta().y()
        if delta > 0:
            self.set_value_direct(self.value + self.step)
        else:
            self.set_value_direct(self.value - self.step)
        event.accept()

class SRTEditor(QMainWindow):
    # Constant for placeholder character (visible in viewer)
    INDENT_PLACEHOLDER = '␣'  # U+2423 OPEN BOX

    def __init__(self, splash=None):
        super().__init__()
        self.splash = splash
        self.srt_blocks = []
        self.current_block_index = 0
        self.speakers = ["A", "B", "C", "D"]
        self.speaker_color_palette = [
            QColor(220, 240, 255),  # Light blue
            QColor(255, 220, 220),  # Light red
            QColor(220, 255, 220),  # Light green
            QColor(255, 255, 200),  # Light yellow
            QColor(230, 200, 255),  # Light purple
            QColor(255, 200, 150),  # Light orange
            QColor(200, 230, 230),  # Light cyan
            QColor(255, 210, 230)   # Light pink
        ]
    
        # Initialize speaker colors from palette
        self.speaker_colors = self.speaker_color_palette[:4]  # First 4 colors for initial speakers

        self.context_blocks = 5
        self.current_file_path = None
        self.file_has_timestamps = True
        self.timestamp_style="curly"
        self.custom_timestamp_pattern="{HH:MM:SS}"
        self.audio_file_path = None
        self.project_name = ""
        self.project_memo = ""
        self.text_display_font = QFont("Arial", 12)
        self.has_unsaved_changes = False
        self.current_theme = "light"
        self.playback_speed = 1.0
        self.segment_sync_buffer = 0
        self.original_audio_duration = 0
        self.last_symbol_category = 0
        self.cjk_mode = False
        
        # VLC audio player
        self.audio_player = None
        self.is_playing = False
        self.auto_sync_enabled = False
        self.auto_pause_enabled = False
        
        # Timer for UI updates
        self.ui_update_timer = QTimer()
        self.ui_update_timer.timeout.connect(self.update_ui)
        self.ui_update_timer.start(50)
        
        
        # Check for VLC at startup
#         self.vlc_available = self.check_vlc_available()      
#         if not self.vlc_available:
#             QTimer.singleShot(100, self.show_vlc_warning)
            
        
        self.update_splash("Creating user interface...")
        self.init_ui()
    
    def update_splash(self, message):
        """Update splash screen message if splash exists."""
        if self.splash:
            self.splash.showMessage(message, Qt.AlignBottom | Qt.AlignCenter, Qt.black)
            QApplication.processEvents()  # ensure UI updates
    
        
    def init_ui(self):
        self.setWindowTitle("CapsGAT 1.4 - Subtitle-to-Transcript Workstation")
        self.setGeometry(100, 100, 1400, 900)
        self.setWindowIcon(QIcon(resource_path('images/logo.ico')))
        
        # Create menu bar
        self.update_splash("Creating menu bar...")
        self.create_menu_bar()
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QHBoxLayout(central_widget)
        
        # Left panel - Context display
        left_panel = QVBoxLayout()
        
        # Current block info display
        self.current_info_label = QLabel("No block selected")
        self.current_info_label.setStyleSheet("""
            QLabel {
                background-color: #f0f0f0;
                padding: 10px;
                border: 2px solid #ccc;
                border-radius: 5px;
                font-weight: bold;
            }
        """)
        left_panel.addWidget(self.current_info_label)
        
        # Main content display label
        context_label = QLabel("Transcript Content:")
        context_label.setFont(QFont("Arial", 10, QFont.Bold))
        left_panel.addWidget(context_label)
        
        # Main text display
        self.text_display = QTextEdit()
        self.text_display.setReadOnly(True)
        self.text_display.setFont(self.text_display_font)
        self.text_display.setStyleSheet("""
            QTextEdit {
                background-color: #fafafa;
                border: 2px solid #ddd;
                border-radius: 5px;
                padding: 10px;
            }
        """)
        left_panel.addWidget(self.text_display)
        
        # Navigation buttons
        nav_layout = QHBoxLayout()
        
        self.btn_prev = QPushButton("← Previous (P)")
        self.btn_prev.clicked.connect(self.previous_block)
        
        self.lbl_current = QLabel("Current: -/-")
        
        self.btn_next = QPushButton("Next (N) →")
        self.btn_next.clicked.connect(self.next_block)
        
        self.btn_split = QPushButton("Split (Space)")
        self.btn_split.clicked.connect(self.split_current_block)
        
        self.btn_merge = QPushButton("Merge (Del)")
        self.btn_merge.clicked.connect(self.merge_with_next)
        
        self.btn_edit = QPushButton("Edit (E)")
        self.btn_edit.clicked.connect(self.edit_current_block)
        
        self.btn_unassign = QPushButton("Unassign (U)")
        self.btn_unassign.clicked.connect(self.unassign_current)

        self.btn_symbols = QPushButton("Symbols (*)")
        self.btn_symbols.clicked.connect(self.open_pause_dialog)

        nav_layout.addWidget(self.btn_prev)
        nav_layout.addWidget(self.lbl_current)
        nav_layout.addWidget(self.btn_next)
        nav_layout.addWidget(self.btn_split)
        nav_layout.addWidget(self.btn_merge)
        nav_layout.addWidget(self.btn_edit)
        nav_layout.addWidget(self.btn_unassign)
        nav_layout.addWidget(self.btn_symbols)
        
        left_panel.addLayout(nav_layout)
        
        # Right panel - Controls
        right_panel = QVBoxLayout()

        # Speaker assignment header with + and - buttons
        speaker_header = QHBoxLayout()
        speaker_label = QLabel("Assign Speaker:")
        speaker_label.setFont(QFont("Arial", 12, QFont.Bold))
        speaker_header.addWidget(speaker_label)

        # Add spacer to push buttons to the right
        speaker_header.addStretch()

        # Circular minus button
        self.btn_remove_speaker = QPushButton("−")  # Minus sign
        self.btn_remove_speaker.setFixedSize(28, 28)
        self.btn_remove_speaker.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 2px solid #ccc;
                border-radius: 14px;
                font-size: 16px;
                font-weight: bold;
                padding-bottom: 2px;
                color: #333;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                border-color: #999;
            }
            QPushButton:pressed {
                background-color: #d0d0d0;
            }
            QPushButton:disabled {
                background-color: #f8f8f8;
                border-color: #ddd;
                color: #aaa;
            }
        """)
        self.btn_remove_speaker.clicked.connect(self.decrease_speaker_count)
        speaker_header.addWidget(self.btn_remove_speaker)

        # Speaker count display
        self.speaker_count_label = QLabel("4")
        self.speaker_count_label.setFixedSize(30, 28)
        self.speaker_count_label.setAlignment(Qt.AlignCenter)
        self.speaker_count_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                font-weight: bold;
                color: #333;
            }
        """)
        speaker_header.addWidget(self.speaker_count_label)

        # Circular plus button
        self.btn_add_speaker = QPushButton("+")
        self.btn_add_speaker.setFixedSize(28, 28)
        self.btn_add_speaker.setStyleSheet("""
            QPushButton {
                background-color: #f0f0f0;
                border: 2px solid #ccc;
                border-radius: 14px;
                font-size: 16px;
                font-weight: bold;
                padding-bottom: 2px;
                color: #333;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                border-color: #999;
            }
            QPushButton:pressed {
                background-color: #d0d0d0;
            }
            QPushButton:disabled {
                background-color: #f8f8f8;
                border-color: #ddd;
                color: #aaa;
            }
        """)
        self.btn_add_speaker.clicked.connect(self.increase_speaker_count)
        speaker_header.addWidget(self.btn_add_speaker)

        # Add a tiny bit of spacing after buttons
        speaker_header.addSpacing(23)

        right_panel.addLayout(speaker_header)

        # Speaker container
        self.speaker_container = QWidget()
        self.speaker_layout = QVBoxLayout(self.speaker_container)
        self.create_speaker_widgets()
        right_panel.addWidget(self.speaker_container)
        

        # Audio Controls Label
        audio_label = QLabel("Audio Controls:")
        audio_label.setFont(QFont("Arial", 12, QFont.Bold))
        right_panel.addWidget(audio_label)

        # Audio controls group
        audio_group = QGroupBox()
        audio_layout = QVBoxLayout()

        # Audio file info
        self.audio_info_label = QLabel("No audio loaded")
        self.audio_info_label.setStyleSheet("""
            QLabel {
                background-color: #f0f0f0; 
                padding: 5px; 
                border-radius: 3px;
                font-size: 11px;
            }
        """)
        self.audio_info_label.setMinimumWidth(170)
        audio_layout.addWidget(self.audio_info_label)

        # Audio progress bar
        self.audio_progress = QSlider(Qt.Horizontal)
        self.audio_progress.setEnabled(False)
        self.audio_progress.sliderMoved.connect(self.seek_audio)
        audio_layout.addWidget(self.audio_progress)

        # Time display
        time_layout = QHBoxLayout()
        self.time_label = QLabel("00:00 / 00:00")
        self.time_label.setAlignment(Qt.AlignCenter)
        
        self.btn_jump_to = QPushButton("Jump (Ctrl+J)")
        self.btn_jump_to.clicked.connect(self.jump_to_time)
        self.btn_jump_to.setEnabled(False)
        
        self.btn_play_segment = QPushButton("Play from segment (Shift+Enter)")
        self.btn_play_segment.clicked.connect(self.play_from_current_segment)
        self.btn_play_segment.setEnabled(False)  # initially disabled until audio loaded
        self.btn_play_segment.setToolTip("Play audio from the start of the current segment")
        
        time_layout.addWidget(self.time_label)
        time_layout.addWidget(self.btn_jump_to)
        time_layout.addWidget(self.btn_play_segment)
        audio_layout.addLayout(time_layout)

        # Audio controls
        audio_controls_layout = QHBoxLayout()
        
        audio_button_font = QFont("Segoe UI Symbol")
        


        self.btn_load_audio = QPushButton("Load Audio")
        self.btn_load_audio.clicked.connect(self.load_audio_file)

        self.btn_rewind = QPushButton("⏪ (PgUp)")
        self.btn_rewind.clicked.connect(self.rewind_audio)
        self.btn_rewind.setFont(audio_button_font)
        self.btn_rewind.setEnabled(False)

        self.btn_play = QPushButton("▶ (End)")
        self.btn_play.clicked.connect(self.toggle_playback)
        self.btn_play.setFont(audio_button_font)
        self.btn_play.setEnabled(False)

        self.btn_forward = QPushButton("⏩ (PgDn)")
        self.btn_forward.clicked.connect(self.forward_audio)
        self.btn_forward.setFont(audio_button_font)
        self.btn_forward.setEnabled(False)

        audio_controls_layout.addWidget(self.btn_load_audio)
        audio_controls_layout.addWidget(self.btn_rewind)
        audio_controls_layout.addWidget(self.btn_play)
        audio_controls_layout.addWidget(self.btn_forward)
        audio_controls_layout.addStretch()

        audio_layout.addLayout(audio_controls_layout)

        # Auto-sync and Autopause checkboxes
        sync_layout = QHBoxLayout()
        self.auto_sync_check = QCheckBox("Auto-sync to audio")
        self.auto_sync_check.setEnabled(False)
        self.auto_sync_check.setChecked(False)
        self.auto_sync_check.toggled.connect(self.toggle_auto_sync)

        self.auto_pause_check = QCheckBox("Autopause during editing")
        self.auto_pause_check.setEnabled(False)
        self.auto_pause_check.setChecked(False)
        self.auto_pause_check.toggled.connect(self.toggle_auto_pause)

        sync_layout.addWidget(self.auto_sync_check)
        sync_layout.addWidget(self.auto_pause_check)
        sync_layout.addStretch()

        audio_layout.addLayout(sync_layout)
        
        # Speed control
        speed_layout = QHBoxLayout()
        speed_layout.addWidget(QLabel("Playback Speed:"))
        
        self.speed_slower_btn = QPushButton("-")
        self.speed_slower_btn.clicked.connect(lambda: self.speed_knob.set_value_direct(max(0.5, self.playback_speed - 0.1)))
        self.speed_slower_btn.setFixedWidth(30)
        speed_layout.addWidget(self.speed_slower_btn)
        
        self.speed_knob = SpeedKnob()
        self.speed_knob.valueChanged.connect(self.change_playback_speed)
        speed_layout.addWidget(self.speed_knob)
                        
        self.speed_normal_btn = QPushButton("Reset")
        self.speed_normal_btn.clicked.connect(lambda: self.speed_knob.set_value_direct(1.0))
        self.speed_normal_btn.setFixedWidth(50)
        
        self.speed_faster_btn = QPushButton("+")
        self.speed_faster_btn.clicked.connect(lambda: self.speed_knob.set_value_direct(min(2.0, self.playback_speed + 0.1)))
        self.speed_faster_btn.setFixedWidth(30)
        
        speed_layout.addWidget(self.speed_faster_btn)
        speed_layout.addWidget(self.speed_normal_btn)
        
        # Disable speed controls if VLC is not available
        self.speed_knob.setEnabled(False)
        self.speed_slower_btn.setEnabled(False)
        self.speed_normal_btn.setEnabled(False)
        self.speed_faster_btn.setEnabled(False)
        self.speed_knob.setToolTip("Load audio with VLC to enable speed control")
                
        audio_layout.addLayout(speed_layout)
        
        audio_group.setLayout(audio_layout)
        right_panel.addWidget(audio_group)
        
        # Unassigned blocks with counter
        self.unassigned_blocks_label = QLabel("Unassigned Segments (0/0):")
        self.unassigned_blocks_label.setFont(QFont("Arial", 12, QFont.Bold))
        right_panel.addWidget(self.unassigned_blocks_label)

        self.unassigned_list = QListWidget()
        self.unassigned_list.itemDoubleClicked.connect(self.jump_to_block)
        right_panel.addWidget(self.unassigned_list, 1)  # Give it stretch factor 1 to expand

        # Button container (will stay at bottom)
        button_container = QWidget()
        button_layout = QVBoxLayout(button_container)
        button_layout.setContentsMargins(0, 0, 0, 0)
        button_layout.setSpacing(5)

        self.btn_quick_export = QPushButton(" Export Transcript (Ctrl+Enter)")
        self.btn_quick_export.clicked.connect(self.export_transcript)
        svg_data = b'''<svg width="48" height="48" viewBox="0 0 600 500" xmlns="http://www.w3.org/2000/svg">
          <path d="m171,183l-94,0l0,227l397,0l0,-50" fill="none" stroke="#ffffff" stroke-width="13" transform="matrix(0.951576, 0, 0, 0.792082, 22.341, 67.1476)"/>
          <path d="m442.89,147.05l0.11,-76.05l114,118l-111,110l0,-69l-104.18,5.12c-29.52,4.98 -61.24,13.64 -106.76,43.52c-20.93,14.7 -31.06,17.82 -55.47,47c2.24,-13 10.61,-52.7 30.2,-83.17c35.09,-46.28 65.11,-63.01 74.11,-67.01c41.62,-30.5 165,-26.41 159,-28.41z" fill="none" stroke="#ffffff" stroke-width="13"/>
        </svg>'''
        pixmap = QPixmap()
        pixmap.loadFromData(svg_data)
        icon = QIcon(pixmap)
        self.btn_quick_export.setIcon(icon)
        self.btn_quick_export.setIconSize(QSize(48, 48))
        self.btn_quick_export.setStyleSheet("""
            QPushButton {
                background-color: #124607;
                padding: 5px 9px;
                font-size: 14px;
                font-weight: bold;
                color: white;
                border: 0px;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #71906a;
            }
        """)
        button_layout.addWidget(self.btn_quick_export)
        
        right_panel.addWidget(button_container)  # No stretch, stays at bottom
        
        layout.addLayout(left_panel, 4)
        layout.addLayout(right_panel, 1)
        
        self.update_splash("Setting up shortcuts...")
        self.setup_shortcuts()
               
    def format_timestamp(self, seconds, style="curly", custom_pattern=None):
        """
        Format a timestamp (seconds) according to the chosen style.
        style: "hash", "bracket", "custom"
        custom_pattern: a string containing placeholders HH, MM, SS, xx
        Returns a string (no surrounding whitespace).
        """
        if seconds is None:
            return ""

        # Compute components
        h = int(seconds // 3600)
        m = int((seconds % 3600) // 60)
        s = int(seconds % 60)
        tenths = int((seconds - int(seconds)) * 10)   # one digit after decimal
        # If you want milliseconds (three digits), use: ms = int((seconds - int(seconds)) * 1000)

        if style == "curly":
            return f"{{{h:02d}:{m:02d}:{s:02d}}}"
        elif style == "hash":
            return f"#{h:02d}:{m:02d}:{s:02d}-{tenths}#"
        elif style == "bracket":
            return f"[{h:02d}:{m:02d}:{s:02d}]"
        elif style == "custom" and custom_pattern:
            # Simple placeholder replacement (case‑insensitive)
            result = custom_pattern
            result = result.replace("HH", f"{h:02d}")
            result = result.replace("MM", f"{m:02d}")
            result = result.replace("SS", f"{s:02d}")
            result = result.replace("xx", f"{tenths}")
            # If you support milliseconds, also replace "mmm"
            return result
        else:
            return ""
        
    def get_timestamp_width(self, style, custom_pattern=None):
        """Return the width (in characters) of a timestamp for the given style."""
        if style =="curly":
            return 10
        elif style == "hash":
            return 12   # #HH:MM:SS-x# always 12 chars
        elif style == "bracket":
            return 10   # [HH:MM:SS] always 10 chars
        else:  # custom
            # Use a sample time (1:23:45.6) to determine length
            sample = self.format_timestamp(5025.6, style, custom_pattern)
            return len(sample)
    
    def count_leading_spaces(self, s):
        """Return number of leading space characters."""
        return len(s) - len(s.lstrip(' '))
    
    def create_menu_bar(self):
        menubar = self.menuBar()
        
        # File menu
        file_menu = menubar.addMenu('File')
        
        new_project_action = QAction('New Project', self)
        new_project_action.setShortcut('Ctrl+N')
        new_project_action.triggered.connect(self.new_project)
        file_menu.addAction(new_project_action)
        
        open_project_action = QAction('Open Project...', self)
        open_project_action.setShortcut('Ctrl+O')
        open_project_action.triggered.connect(self.load_project)
        file_menu.addAction(open_project_action)
        
        save_project_action = QAction('Save Project', self)
        save_project_action.setShortcut('Ctrl+S')
        save_project_action.triggered.connect(self.save_project)
        file_menu.addAction(save_project_action)
        
        save_as_action = QAction('Save Project As...', self)
        save_as_action.triggered.connect(lambda: self.save_project(force_save_as=True))
        file_menu.addAction(save_as_action)
        
        file_menu.addSeparator()
        
        # Import submenu
        import_menu = file_menu.addMenu('Import')
        
        import_subtitles_action = QAction('Subtitles...', self)
        import_subtitles_action.triggered.connect(self.import_subtitles)
        import_menu.addAction(import_subtitles_action)
        
        import_audio_action = QAction('Audio File...', self)
        import_audio_action.triggered.connect(self.load_audio_file)
        import_menu.addAction(import_audio_action)
        
        # Export
        export_action = QAction('Export...', self)
        export_action.setShortcut('Ctrl+Return')
        export_action.triggered.connect(self.export_transcript)
        file_menu.addAction(export_action)
        
        file_menu.addSeparator()
        
        exit_action = QAction('Exit', self)
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        
        # Edit menu
        edit_menu = menubar.addMenu('Edit')
        
        search_action = QAction('Search...', self)
        search_action.triggered.connect(self.open_search_dialog)
        edit_menu.addAction(search_action)
        
        settings_action = QAction('Settings...', self)
        settings_action.triggered.connect(self.open_settings)
        edit_menu.addAction(settings_action)
        
        project_memo_action = QAction('Project Memo...', self)
        project_memo_action.triggered.connect(self.open_project_memo)
        edit_menu.addAction(project_memo_action)
        
        custom_menu = edit_menu.addMenu('Custom Symbols')
    
        manage_custom_action = QAction('Manage Custom Symbols...', self)
        manage_custom_action.triggered.connect(self.manage_custom_symbols)
        custom_menu.addAction(manage_custom_action)
          
        custom_menu.addSeparator()
        
        export_custom_action = QAction('Export Custom Symbols...', self)
        export_custom_action.triggered.connect(self.export_custom_symbols)
        custom_menu.addAction(export_custom_action)
        
        import_custom_action = QAction('Import Custom Symbols...', self)
        import_custom_action.triggered.connect(self.import_custom_symbols)
        custom_menu.addAction(import_custom_action)
        self.manage_custom_action = manage_custom_action
        self.export_custom_action = export_custom_action
        self.import_custom_action = import_custom_action
        # Help menu
        help_menu = menubar.addMenu('Help')
        
        shortcuts_action = QAction('Shortcuts', self)
        shortcuts_action.setShortcut('F1')
        shortcuts_action.triggered.connect(self.show_shortcuts)
        help_menu.addAction(shortcuts_action)
        
        manual_action = QAction('Online Manual', self)
        manual_action.setShortcut('Ctrl+F1')
        manual_action.triggered.connect(self.open_manual)
        help_menu.addAction(manual_action)
        
        about_action = QAction('About CapsGAT', self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
        
    def update_menu_state(self):
        """Enable/disable custom symbol menu items based on transcript presence."""
        enabled = bool(self.srt_blocks)
        if hasattr(self, 'manage_custom_action'):
            self.manage_custom_action.setEnabled(enabled)
            self.export_custom_action.setEnabled(enabled)
            self.import_custom_action.setEnabled(enabled)
        
    def manage_custom_symbols(self):
        """Open symbol dialog and switch to Custom tab."""
        if not self.srt_blocks:
            QMessageBox.information(self, "No Transcript",
                "Please load a transcript first to manage custom symbols.")
            return
        dialog = EnhancedSymbolDialog(self)
        # Switch to the last category (Custom)
        dialog.switch_category(len(dialog.categories) - 1)
        if dialog.exec_() == QDialog.Accepted:
            symbol_info = dialog.get_selected_symbol_info()
            self.handle_symbol_insertion(symbol_info)

    def export_custom_symbols(self):
        """Export custom symbols to file"""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Custom Symbols", "",
            "JSON Files (*.json);;All Files (*)"
        )
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(EnhancedSymbolDialog.custom_symbols, f, indent=2)
                QMessageBox.information(self, "Success", f"Custom symbols exported to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not export: {e}")

    def import_custom_symbols(self):
        """Import custom symbols from file"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Import Custom Symbols", "",
            "JSON Files (*.json);;All Files (*)"
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    imported = json.load(f)
                
                if isinstance(imported, list):
                    reply = QMessageBox.question(
                        self, "Merge Symbols",
                        f"Found {len(imported)} symbols. Merge with existing?",
                        QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
                    )
                    
                    if reply == QMessageBox.Yes:
                        EnhancedSymbolDialog.custom_symbols.extend(imported)
                    elif reply == QMessageBox.No:
                        EnhancedSymbolDialog.custom_symbols = imported
                    else:
                        return
                    
                    # Save to file
                    with open(EnhancedSymbolDialog.custom_symbols_file, 'w', encoding='utf-8') as f:
                        json.dump(EnhancedSymbolDialog.custom_symbols, f, indent=2)
                    
                    QMessageBox.information(self, "Success", f"Imported {len(imported)} symbols")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not import: {e}")
        
    def increase_speaker_count(self):
        """Increase number of speakers by 1"""
        current = len(self.speakers)
        if current < 8:  # Max 8 speakers
            self.update_speaker_count(current + 1)
            self.speaker_count_label.setText(str(current + 1))
            # Update button states
            self.update_speaker_buttons()

    def decrease_speaker_count(self):
        """Decrease number of speakers by 1"""
        current = len(self.speakers)
        if current <= 2:  # Min 2 speakers
            return
        
        # Check if the last speaker (highest index) has any assigned segments
        last_speaker_idx = current - 1
        assigned_count = self.count_blocks_for_speaker(last_speaker_idx)
        
        if assigned_count > 0:
            # Show warning dialog
            speaker_name = self.speakers[last_speaker_idx]
            reply = QMessageBox.warning(
                self,
                "Remove Speaker",
                f"Speaker {speaker_name} is assigned to {assigned_count} segment(s).\n\n"
                f"Removing this speaker will unassign all these segments.\n\n"
                f"Do you want to continue?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.No
            )
            
            if reply != QMessageBox.Yes:
                return  # User canceled
            
            # Unassign all blocks for this speaker
            for block in self.srt_blocks:
                if block.get('speaker') == last_speaker_idx:
                    block['speaker'] = None
                    block['is_turn_start'] = True
            
            self.update_display()
            self.mark_unsaved_changes()
        
        # Now remove the speaker
        self.update_speaker_count(current - 1)

    def update_speaker_buttons(self):
        """Enable/disable speaker buttons based on current count"""
        current = len(self.speakers)
        self.btn_add_speaker.setEnabled(current < 8)
        self.btn_remove_speaker.setEnabled(current > 2)
           
    def setup_shortcuts(self):
        # Clear any existing shortcuts first (to avoid duplicates)
        for shortcut in self.findChildren(QShortcut):
            shortcut.disconnect()
            shortcut.deleteLater()
        
        # Dynamic speaker shortcuts based on current count
        self.speaker_shortcuts = []  # Store shortcuts to prevent garbage collection
        for i in range(len(self.speakers)):
            shortcut = QShortcut(QKeySequence(str(i+1)), self)
            # Use a default argument to capture i, and accept the checked parameter
            shortcut.activated.connect(lambda checked=False, idx=i: self.assign_speaker(idx))
            self.speaker_shortcuts.append(shortcut)
        
        # Navigation shortcuts
        QShortcut(QKeySequence("N"), self).activated.connect(self.next_block)
        QShortcut(QKeySequence("P"), self).activated.connect(self.previous_block)
        QShortcut(QKeySequence("Right"), self).activated.connect(self.next_block)
        QShortcut(QKeySequence("Left"), self).activated.connect(self.previous_block)
        QShortcut(QKeySequence("Space"), self).activated.connect(self.split_current_block)
        QShortcut(QKeySequence("Delete"), self).activated.connect(self.merge_with_next)
        QShortcut(QKeySequence("E"), self).activated.connect(self.edit_current_block)
        QShortcut(QKeySequence("F2"), self).activated.connect(self.edit_current_block)
        QShortcut(QKeySequence("*"), self).activated.connect(self.open_pause_dialog)
        QShortcut(QKeySequence("U"), self).activated.connect(self.unassign_current)
        QShortcut(QKeySequence("Return"), self).activated.connect(self.insert_empty_line)
        QShortcut(QKeySequence("."), self).activated.connect(lambda: self.handle_pause("(.)"))
        QShortcut(QKeySequence("H"), self).activated.connect(lambda: self.handle_pause("°h"))
        QShortcut(QKeySequence("Shift+H"), self).activated.connect(lambda: self.handle_pause("h°"))
        QShortcut(QKeySequence("Shift+Return"), self).activated.connect(self.play_from_current_segment)
        QShortcut(QKeySequence("Shift+Enter"), self).activated.connect(self.play_from_current_segment)  # for numpad Enter
        
        QShortcut(QKeySequence("PgUp"), self).activated.connect(self.rewind_audio)
        QShortcut(QKeySequence("End"), self).activated.connect(self.toggle_playback)
        QShortcut(QKeySequence("PgDown"), self).activated.connect(self.forward_audio)
        QShortcut(QKeySequence("Ctrl+J"), self).activated.connect(self.jump_to_time)
               
        QShortcut(QKeySequence("+"), self).activated.connect(lambda: self.speed_knob.set_value_direct(min(2.0, self.playback_speed + 0.1)))
        QShortcut(QKeySequence("-"), self).activated.connect(lambda: self.speed_knob.set_value_direct(max(0.5, self.playback_speed - 0.1)))
        QShortcut(QKeySequence("0"), self).activated.connect(lambda: self.speed_knob.set_value_direct(1.0))
        
        QShortcut(QKeySequence("Ctrl+F"), self).activated.connect(self.open_search_dialog)
        QShortcut(QKeySequence("F3"), self).activated.connect(self.find_next)
        QShortcut(QKeySequence("Shift+F3"), self).activated.connect(self.find_previous)
        
        self.update_splash("Loading user interface...")
    
    def update_ui(self):
        """Update UI elements based on current state"""
        if self.audio_player:
            current_time = self.audio_player.get_position()
            duration = self.audio_player.duration
            
            if duration > 0:
                # Update progress slider
                self.audio_progress.setValue(int(current_time / duration * 100))
                
                # Update time label
                current_str = f"{int(current_time // 60):02d}:{int(current_time % 60):02d}"
                duration_str = f"{int(duration // 60):02d}:{int(duration % 60):02d}"
                speed_str = f" ({self.playback_speed:.1f}x)" if abs(self.playback_speed - 1.0) > 0.01 else ""
                self.time_label.setText(f"{current_str} / {duration_str}")
                
                # Auto-sync if enabled
                if self.auto_sync_enabled and self.srt_blocks and self.file_has_timestamps:
                    self.auto_sync_with_audio(current_time)
    
    def update_speed_controls_state(self):
        """Enable speed controls only if the current player is a VlcAudioPlayer."""
        if self.audio_player and isinstance(self.audio_player, VlcAudioPlayer):
            self.speed_knob.setEnabled(True)
            self.speed_slower_btn.setEnabled(True)
            self.speed_normal_btn.setEnabled(True)
            self.speed_faster_btn.setEnabled(True)
            self.speed_knob.setToolTip("Adjust playback speed")
        else:
            self.speed_knob.setEnabled(False)
            self.speed_slower_btn.setEnabled(False)
            self.speed_normal_btn.setEnabled(False)
            self.speed_faster_btn.setEnabled(False)
            self.speed_knob.setToolTip("Speed control requires VLC media player")
    
    def load_audio_file(self):
        """Load an audio file using appropriate player"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Load Audio File", "", 
            "Audio Files (*.mp3 *.wav *.ogg *.m4a *.flac *.aac *.wma);;All Files (*)"
        )
        
        if not file_path:
            return

        # Check for subtitle files in the same directory
        audio_dir = Path(file_path).parent
        subtitle_files = list(audio_dir.glob("*.srt")) + list(audio_dir.glob("*.json")) + \
                         list(audio_dir.glob("*.txt")) + list(audio_dir.glob("*.tsv"))
        
        if subtitle_files:
            reply = QMessageBox.question(
                self,
                "Subtitle File Found",
                f"Found {len(subtitle_files)} subtitle file(s). Import one?",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel
            )
            
            if reply == QMessageBox.Yes:
                file_list = [str(f.name) for f in subtitle_files]
                file_name, ok = QInputDialog.getItem(
                    self, "Select Subtitle File", "Choose file:", file_list, 0, False
                )
                if ok and file_name:
                    if self.check_unsaved_changes():
                        subtitle_path = audio_dir / file_name
                        self.current_file_path = str(subtitle_path)
                        self.load_file_from_path(str(subtitle_path))
            elif reply == QMessageBox.Cancel:
                return
        
        # Stop any existing audio
        if self.audio_player:
            self.audio_player.cleanup()
            self.audio_player = None
        
        # Reset speed display
        self.playback_speed = 1.0
        self.speed_knob.value = 1.0
        self.speed_knob.update()
        self.speed_normal_btn.setText("1.0x")
        
        self.update_splash("Loading audio player...")
        
        # Try to create VLC player first
        try:
            self.audio_player = VlcAudioPlayer()
            player_name = "VLC"
            vlc_ok = True
        except Exception as e:
            logger.warning(f"VLC not available: {e}")
            # Fallback to simple player
            if HAS_PYAUDIO:
                self.audio_player = SimpleAudioPlayer()
                player_name = "Simple (fallback)"
                vlc_ok = False
            else:
                QMessageBox.warning(self, "No Audio Backend", 
                    "Neither VLC nor PyAudio is available.\n\n"
                    "Please install VLC from https://www.videolan.org/vlc/")
                return
        
        # Connect signals
        self.audio_player.playback_started.connect(self.on_playback_started)
        self.audio_player.playback_stopped.connect(self.on_playback_stopped)
        self.audio_player.position_changed.connect(self.on_position_changed)
        
        if hasattr(self.audio_player, 'playback_paused'):
            self.audio_player.playback_paused.connect(self.on_playback_paused)
        if hasattr(self.audio_player, 'end_reached'):
            self.audio_player.end_reached.connect(self.on_playback_ended)
        
        # Load the file
        if self.audio_player.load_file(file_path):
            self.audio_file_path = file_path
            audio_name = Path(file_path).name
            
            # Update status label
            if vlc_ok:
                status = f"Audio loaded: {audio_name}"
            else:
                status = f"Audio loaded: {audio_name} (⚠ VLC not found)"
            self.audio_info_label.setText(status)
            
            # Enable basic controls
            self.btn_play.setEnabled(True)
            self.btn_rewind.setEnabled(True)
            self.btn_forward.setEnabled(True)
            self.btn_jump_to.setEnabled(True)
            self.btn_play_segment.setEnabled(True)
            self.auto_sync_check.setEnabled(self.file_has_timestamps)
            self.auto_pause_check.setEnabled(True)
            self.audio_progress.setEnabled(True)
            
            # Update speed controls based on player type
            self.update_speed_controls_state()
            
            logger.info(f"Audio loaded with {player_name} player: {audio_name}")
        else:
            QMessageBox.critical(self, "Error", f"Failed to load audio file")
            self.audio_player = None
            self.audio_info_label.setText("No audio loaded")
    
    def toggle_playback(self):
        """Toggle play/pause"""
        if not self.audio_player:
            return
        
        if self.audio_player.is_playing:
            self.audio_player.pause()
            self.btn_play.setText("▶ (End)")
            self.is_playing = False
        else:
            # For fallback player, ensure thread is running
            if isinstance(self.audio_player, SimpleAudioPlayer) and not self.audio_player.isRunning():
                self.audio_player.start()
            
            self.audio_player.play()
            self.btn_play.setText("⏸ (End)")
            self.is_playing = True
            
    def rewind_audio(self):
        """Rewind by 5 seconds"""
        if self.audio_player:
            current_pos = self.audio_player.get_position()
            new_pos = max(0, current_pos - 5)
            self.audio_player.seek(new_pos)
    
    def forward_audio(self):
        """Fast forward by 5 seconds"""
        if self.audio_player:
            current_pos = self.audio_player.get_position()
            duration = self.audio_player.duration
            new_pos = min(duration, current_pos + 5)
            self.audio_player.seek(new_pos)
    
    def seek_audio(self, position_percent):
        """Seek to position"""
        if self.audio_player:
            position_seconds = position_percent / 100.0 * self.audio_player.duration
            self.audio_player.seek(position_seconds)
    
    def jump_to_time(self):
        """Jump to specific time"""
        if self.audio_player and self.audio_player.duration > 0:
            max_duration = self.audio_player.duration
            dialog = JumpToTimeDialog(int(max_duration * 1000), self)
            if dialog.exec_() == QDialog.Accepted:
                target_time_ms = dialog.get_target_time()
                self.audio_player.seek(target_time_ms / 1000.0)
                
    def play_from_current_segment(self):
        """Seek audio to the start of the current transcript block and play."""
        if not self.audio_player:
            QMessageBox.information(self, "No Audio", "No audio file loaded.")
            return

        if not self.srt_blocks:
            return

        block = self.srt_blocks[self.current_block_index]
        if not block.get('start_time'):
            QMessageBox.information(self, "No Timestamp",
                "The current segment does not have a start time.")
            return

        # Convert start_time to seconds
        seconds = self.time_to_seconds(block['start_time'])
        self.audio_player.seek(seconds)

        # If already playing, continue; otherwise start playing
        if not self.audio_player.is_playing:
            self.audio_player.play()
            self.btn_play.setText("⏸ (End)")
            self.is_playing = True
    
    def change_playback_speed(self, new_speed):
        """Change playback speed (only works with VLC)"""
        # Debug: see if we get here
        print(f"change_playback_speed called with {new_speed}, player: {self.audio_player}")

        if not self.audio_player:
            QMessageBox.information(self, "No Audio",
                "No audio file loaded.")
            return

        if not isinstance(self.audio_player, VlcAudioPlayer):
            QMessageBox.information(self, "Speed Control Not Available",
                "Playback speed control requires VLC media player.\n\n"
                "Please install VLC from https://www.videolan.org/vlc/")
            return

        new_speed = max(0.5, min(2.0, new_speed))

        # Don't process if speed hasn't changed
        if abs(new_speed - self.playback_speed) < 0.01:
            return

        self.playback_speed = new_speed
        self.audio_player.set_speed(new_speed)
    
    def on_playback_started(self):
        """Handle playback started"""
        self.is_playing = True
        logger.info("Playback started")
    
    def on_playback_paused(self):
        """Handle playback paused"""
        self.is_playing = False
        logger.info("Playback paused")
    
    def on_playback_stopped(self):
        """Handle playback stopped"""
        self.is_playing = False
        self.btn_play.setText("▶ (End)")
        logger.info("Playback stopped")
    
    def on_playback_ended(self):
        """Handle playback ended"""
        self.is_playing = False
        self.btn_play.setText("▶ (End)")
        logger.info("Playback ended")
    
    def on_position_changed(self, position):
        """Handle position change"""
        # This is handled in update_ui
        pass
    
    def auto_sync_with_audio(self, current_time):
        """Auto-sync transcript with audio position"""
        if not self.srt_blocks or not self.file_has_timestamps:
            return
        
        # Find block containing current time
        current_time_ms = current_time * 1000  # Convert to milliseconds
        buffer_ms = 100  # Small buffer for better sync
        
        # First check current block
        if 0 <= self.current_block_index < len(self.srt_blocks):
            block = self.srt_blocks[self.current_block_index]
            if block.get('start_time') and block.get('end_time'):
                start_ms = self.time_to_ms(block['start_time'])
                end_ms = self.time_to_ms(block['end_time'])
                
                if start_ms - buffer_ms <= current_time_ms <= end_ms + buffer_ms:
                    return  # Still in current block
        
        # Search for matching block
        for i, block in enumerate(self.srt_blocks):
            if block.get('start_time') and block.get('end_time'):
                start_ms = self.time_to_ms(block['start_time'])
                end_ms = self.time_to_ms(block['end_time'])
                
                if start_ms - buffer_ms <= current_time_ms <= end_ms + buffer_ms:
                    if i != self.current_block_index:
                        self.current_block_index = i
                        self.update_display()
                    break
    
    def time_to_ms(self, time_str):
        """Convert time string to milliseconds"""
        if not time_str:
            return 0
        
        if ',' in time_str:
            time_part, ms_part = time_str.split(',')
            ms = int(ms_part)
        else:
            time_part = time_str
            ms = 0
        
        parts = time_part.split(':')
        if len(parts) == 3:
            hours = int(parts[0])
            minutes = int(parts[1])
            seconds = int(parts[2])
        elif len(parts) == 2:
            hours = 0
            minutes = int(parts[0])
            seconds = int(parts[1])
        else:
            return 0
        
        return (hours * 3600 + minutes * 60 + seconds) * 1000 + ms
    
    def toggle_auto_sync(self, checked):
        """Toggle auto-sync"""
        self.auto_sync_enabled = checked
        logger.info(f"Auto-sync {'enabled' if checked else 'disabled'}")
    
    def toggle_auto_pause(self, checked):
        """Toggle auto-pause"""
        self.auto_pause_enabled = checked
        logger.info(f"Auto-pause {'enabled' if checked else 'disabled'}")
              

    def wrap_text(self, text, max_width, character_wrap=False, first_line_only_indent=True):
        """
        Wrap text to max_width characters, preserving internal spaces.
        If character_wrap is True, break at exact character positions.
        Otherwise, break at word boundaries without collapsing spaces.
        """
        if not text or max_width <= 0:
            return [text]

        if character_wrap:
            # Simple character‑based wrap
            return [text[i:i+max_width] for i in range(0, len(text), max_width)]

        # Word‑based wrap that preserves spaces
        # Use a regex that captures runs of spaces as separate tokens
        tokens = re.split(r'(\s+)', text)
        # tokens are alternating: word, space, word, space, ... (empty strings are filtered out)

        lines = []
        current_line = ''

        for token in tokens:
            if not token:
                continue

            # If adding this token would exceed max_width
            if len(current_line + token) > max_width:
                if current_line:
                    lines.append(current_line.rstrip())   # remove trailing space if any
                    current_line = ''

                # If the token itself is longer than max_width, break it into chunks
                if len(token) > max_width:
                    for i in range(0, len(token), max_width):
                        chunk = token[i:i+max_width]
                        if chunk:
                            lines.append(chunk)
                    continue

            current_line += token

        if current_line:
            lines.append(current_line.rstrip())

        return lines

    def escape_html(self, text):
        """Escape HTML special characters to prevent interpretation as HTML tags"""
        if not text:
            return text
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#39;'))
    
    def new_project(self):
        """Create new project"""
        if self.check_unsaved_changes():
            self.srt_blocks = []
            self.current_block_index = 0
            self.current_file_path = None
            self.project_name = ""
            self.project_memo = ""
            self.has_unsaved_changes = False
            
            self.speakers = ["A", "B", "C", "D"]
            self.speaker_count_label.setText("4")
            
            # Stop audio
            if self.audio_player:
                self.audio_player.cleanup()
                self.audio_player = None
            
            self.audio_file_path = None
            self.audio_info_label.setText("No audio loaded")
            self.btn_play.setEnabled(False)
            self.btn_rewind.setEnabled(False)
            self.btn_forward.setEnabled(False)
            self.btn_jump_to.setEnabled(False)
            self.auto_sync_check.setEnabled(False)
            self.auto_pause_check.setEnabled(False)
            self.auto_sync_check.setChecked(False)
            self.auto_pause_check.setChecked(False)
            self.audio_progress.setEnabled(False)
            self.time_label.setText("00:00 / 00:00")
            self.audio_progress.setValue(0)
            
            # Reset speed to 1.0
            self.playback_speed = 1.0
            self.speed_knob.value = 1.0
            self.speed_knob.update()
            
            self.clear_search_highlights()
            self.unassigned_list.clear()
            self.unassigned_blocks_label.setText("Unassigned Segments (0/0):")
            self.create_speaker_widgets()
            self.setup_shortcuts()
            
            self.update_display()
            self.clear_unsaved_changes()
            self.update_menu_state()
            
    def check_unsaved_changes(self):
        """Check if there are unsaved changes and prompt to save"""
        if self.has_unsaved_changes:
            reply = QMessageBox.question(
                self, 
                "Unsaved Changes", 
                "You have unsaved changes. Would you like to save them?",
                QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel
            )
            
            if reply == QMessageBox.Save:
                # Use Save As to ensure we're saving as a project file, not overwriting the original transcript
                return self.save_project(force_save_as=True)
            elif reply == QMessageBox.Discard:
                return True
            else:  # Cancel
                return False
        return True

    def open_search_dialog(self):
        """Open the search dialog"""
        if not hasattr(self, 'search_dialog') or not self.search_dialog:
            self.search_dialog = SearchDialog(self)
            self.search_dialog.show()
        else:
            self.search_dialog.show()
            self.search_dialog.activateWindow()
            self.search_dialog.search_input.setFocus()
            
    def highlight_search_match(self, block_idx, start_pos, end_pos):
        """Highlight a specific search match and navigate to it"""
        # Navigate to the block
        self.current_block_index = block_idx
        self.update_display()
        
        # Highlight the text within the block
        cursor = self.text_display.textCursor()
        cursor.select(cursor.Document)
        
        # Clear previous extra selections
        self.text_display.setExtraSelections([])
        
        # Find the position of this block in the displayed text
        start_idx = max(0, self.current_block_index - self.context_blocks)
        display_block_pos = (block_idx - start_idx) * 2
        
        # Calculate the position in the displayed text
        cursor.movePosition(cursor.Start)
        for i in range(display_block_pos):
            cursor.movePosition(cursor.Down)
            
        # Move to the beginning of the block text (after prefix)
        cursor.movePosition(cursor.Down)
        cursor.movePosition(cursor.StartOfLine)
        
        # Skip the prefix (">> " or "   ")
        cursor.movePosition(cursor.Right, cursor.MoveAnchor, 3)
        
        # Move to the start position within the text
        cursor.movePosition(cursor.Right, cursor.MoveAnchor, start_pos)
        cursor.movePosition(cursor.Right, cursor.KeepAnchor, end_pos - start_pos)
        
        # Create and apply highlight
        selection = QTextEdit.ExtraSelection()
        selection.cursor = cursor
        selection.format.setBackground(QColor(255, 255, 0))
        selection.format.setForeground(QColor(0, 0, 0))
        
        self.text_display.setExtraSelections([selection])
        self.text_display.setTextCursor(cursor)
        self.text_display.ensureCursorVisible()
        
    def clear_search_highlights(self):
        """Clear all search highlights"""
        self.text_display.setExtraSelections([])


    def show_shortcuts(self):
        """Show keyboard shortcuts dialog"""
        shortcuts_text = """
Keyboard Shortcuts:

Navigation:
• P / Left Arrow: Previous block
• N / Right Arrow: Next block
• 1-4: Assign speakers A-D
• U: Unassign current block

Editing:
• Space: Split current block
• Delete: Merge with next block
• E/F2: Edit block content
• Enter: Insert empty line

GAT2 Symbols:
• *: Open symbols dialog
• .: Insert micropause (with placement)
• h: Insert short inhale (with placement)
• H: Insert short exhale (with placement)

Audio Controls:
• End: Play/Pause audio
• PgUp: Rewind 5 seconds
• PgDn: Fast forward 5 seconds
• Ctrl+J: Jump to Time
• Shift+Enter: Play from current segment
• -: Lower Playback Speed
• +: Speed up Playback

Search Functions:
• Ctrl+F: Open Search Dialog
• F3: Find Next
• Shift+F3: Find Previous

File Operations:
• Ctrl+N: New Project
• Ctrl+O: Open Project
• Ctrl+S: Save Project
• Ctrl+Return: Export Transcript

Help:
• F1: Show Shortcuts (this dialog)
• Ctrl+F1: Open Online Manual (requires internet)
"""
        QMessageBox.information(self, "Keyboard Shortcuts", shortcuts_text)
        
    def open_manual(self):
        """Open online manual in browser"""
        webbrowser.open("https://github.com/anouarg88/CapsGAT/wiki")
        
    def show_about(self):
        """Show about dialog"""
        about_text = """
<b style="font-size: 16px;">CapsGAT 1.4</b><br><br>

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.
<br><br>
This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.
<br><br>
You should have received a copy of the GNU General Public License along with this program.  If not, see 
<a href="https://www.gnu.org/licenses/">https://www.gnu.org/licenses/</a>.
<br><br>
(c) 2026 Anouâr Gadermann
<br>
Engineered wtih DeepSeek V3.2

"""
        QMessageBox.about(self, "About CapsGAT", about_text)
        
    def mark_unsaved_changes(self):
        """Mark that there are unsaved changes"""
        self.has_unsaved_changes = True
        base_title = "CapsGAT 1.4 - Subtitle-to-Transcript Workstation"
        if self.project_name:
            self.setWindowTitle(f"{base_title} - {self.project_name} *")
        else:
            self.setWindowTitle(f"{base_title} *")
            
    def clear_unsaved_changes(self):
        """Clear unsaved changes marker"""
        self.has_unsaved_changes = False
        base_title = "CapsGAT 1.4 - Subtitle-to-Transcript Workstation"
        if self.project_name:
            self.setWindowTitle(f"{base_title} - {self.project_name}")
        else:
            self.setWindowTitle(base_title)
            
    def find_next(self):
        """Find next occurrence (for F3 shortcut)"""
        if hasattr(self, 'search_dialog') and self.search_dialog and self.search_dialog.isVisible():
            self.search_dialog.next_match()
            
    def find_previous(self):
        """Find previous occurrence (for Shift+F3 shortcut)"""
        if hasattr(self, 'search_dialog') and self.search_dialog and self.search_dialog.isVisible():
            self.search_dialog.previous_match()
        
    def import_subtitles(self):
        """Import subtitles via menu"""
        self.load_file()
        
    def save_project_as(self):
        """Save project with new filename"""
        self.save_project(force_save_as=True)
        
    def open_settings(self):
        dialog = SettingsDialog(self.text_display_font, self.current_theme, self.cjk_mode, self)
        if dialog.exec_() == QDialog.Accepted:
            self.text_display_font = dialog.get_font()
            self.text_display.setFont(self.text_display_font)
            theme = dialog.get_theme()
            self.apply_viewer_theme(theme)
            # Update cjk_mode
            self.cjk_mode = dialog.get_cjk_mode()
            # If CJK mode is enabled, suggest a suitable monospaced CJK font
#             if self.cjk_mode:
#                 # Try to set a monospaced CJK font; if not available, keep current
#                 cjk_font = QFont("Noto Sans Mono CJK", self.text_display_font.pointSize())
#                 if cjk_font.exactMatch():
#                     self.text_display_font = cjk_font
#                     self.text_display.setFont(self.text_display_font)
#                 else:
#                     # Fallback to any monospace
#                     monospace = QFont("Courier New", self.text_display_font.pointSize())
#                     if monospace.exactMatch():
#                         self.text_display_font = monospace
#                         self.text_display.setFont(self.text_display_font)
            self.update_display()   # refresh display to apply indentation changes
            self.mark_unsaved_changes()
            
    def open_project_memo(self):
        """Open project memo dialog"""
        dialog = ProjectMemoDialog(self.project_name, self.project_memo, self)
        if dialog.exec_() == QDialog.Accepted:
            project_info = dialog.get_project_info()
            self.project_name = project_info['name']
            self.project_memo = project_info['memo']
            self.mark_unsaved_changes()
            self.clear_unsaved_changes()
        
    def ms_to_time(self, ms):
        """Convert milliseconds to SRT time format"""
        hours = int(ms // 3600000)
        ms %= 3600000
        minutes = int(ms // 60000)
        ms %= 60000
        seconds = int(ms // 1000)
        milliseconds = int(ms % 1000)
        
        return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"
     
    def load_file_from_path(self, file_path):
        """Load a subtitle file from given path"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.srt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.srt_blocks = self.parse_srt(content)
                self.update_menu_state()
                self.file_has_timestamps = True
                
                
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.srt_blocks = self.parse_text(content)
                self.update_menu_state()
                self.file_has_timestamps = False
                
            elif file_extension == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = json.load(f)
                self.srt_blocks = self.parse_json(content)
                self.update_menu_state()
                self.file_has_timestamps = any(block.get('start_time') for block in self.srt_blocks)
                
            elif file_extension == '.tsv':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.srt_blocks = self.parse_tsv(content)
                self.update_menu_state()
                self.file_has_timestamps = True
            
            self.current_block_index = 0
            self.current_file_path = file_path
            
            # Only enable auto-sync checkbox and play from segment button if we have timestamps AND audio is loaded
            self.auto_sync_check.setEnabled(self.file_has_timestamps and self.audio_file_path is not None)
            self.btn_play_segment.setEnabled(self.audio_file_path is not None and self.file_has_timestamps)
            # Auto-pause is always enabled when audio is loaded
            self.auto_pause_check.setEnabled(self.audio_file_path is not None)
            
            # Ensure checkboxes reflect the actual state
            self.auto_sync_check.setChecked(self.auto_sync_enabled)
            self.auto_pause_check.setChecked(self.auto_pause_enabled)
            
            self.update_display()
            self.mark_unsaved_changes()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not load file: {str(e)}")

    def create_speaker_widgets(self):
        for i in reversed(range(self.speaker_layout.count())): 
            self.speaker_layout.itemAt(i).widget().setParent(None)
        
        self.speaker_layout.setSpacing(5)
        self.speaker_layout.setContentsMargins(0, 0, 0, 0)
        
        self.speaker_widgets = []
        for i, speaker in enumerate(self.speakers):
            speaker_widget = QWidget()
            speaker_layout = QHBoxLayout(speaker_widget)
            speaker_layout.setSpacing(10)
            speaker_layout.setContentsMargins(5, 2, 5, 2)
            
            color_label = QLabel("■")
            color_label.setStyleSheet(f"color: {self.speaker_colors[i].name()}; font-size: 20px;")
            
            speaker_name_edit = QLineEdit(speaker)
            speaker_name_edit.textChanged.connect(lambda text, idx=i: self.rename_speaker(idx, text))
            speaker_name_edit.setFixedWidth(120)
            
            speaker_btn = QPushButton(f"{i+1}. Assign")
            speaker_btn.clicked.connect(lambda checked, idx=i: self.assign_speaker(idx))
            
            if self.current_theme == "dark":
                speaker_btn.setStyleSheet(f"""
                    QPushButton {{ 
                        background-color: {self.speaker_colors[i].name()}; 
                        color: white;
                        border: 2px solid #676767;
                        padding: 3px 5px;
                        font-weight: bold;
                        min-width: 100px;
                    }}
                    QPushButton:hover {{
                        background-color: {self.speaker_colors[i].lighter(120).name()};
                        color: white;
                    }}
                """)
            else:
                speaker_btn.setStyleSheet(f"""
                    QPushButton {{ 
                        background-color: {self.speaker_colors[i].name()}; 
                        border: 2px solid darkgray;
                        padding: 3px 5px;
                        font-weight: bold;
                        min-width: 100px;
                    }}
                    QPushButton:hover {{
                        background-color: {self.speaker_colors[i].lighter(120).name()};
                    }}
                """)
            
            speaker_layout.addWidget(color_label)
            speaker_layout.addWidget(QLabel("Name:"))
            speaker_layout.addWidget(speaker_name_edit)
            speaker_layout.addStretch(1)
            speaker_layout.addWidget(speaker_btn)
            
            centered_widget = QWidget()
            centered_widget.setMinimumHeight(40)
            centered_layout = QHBoxLayout(centered_widget)
            centered_layout.setSpacing(0)
            centered_layout.setContentsMargins(0, 0, 0, 0)
            centered_layout.addStretch(1)
            centered_layout.addWidget(speaker_widget)
            centered_layout.addStretch(1)
            
            self.speaker_layout.addWidget(centered_widget)
            self.speaker_widgets.append({
                'name_edit': speaker_name_edit,
                'button': speaker_btn
            })
        
    def rename_speaker(self, speaker_idx, new_name):
        if speaker_idx < len(self.speakers):
            self.speakers[speaker_idx] = new_name
            self.update_display()
            self.mark_unsaved_changes()
    def count_blocks_for_speaker(self, speaker_idx):
        """Count how many blocks are assigned to a specific speaker"""
        count = 0
        for block in self.srt_blocks:
            if block.get('speaker') == speaker_idx:
                count += 1
        return count
            
    def update_speaker_count(self, count):
        while len(self.speakers) > count:
            self.speakers.pop()
            # Don't pop colors - we'll rebuild from palette
        
        while len(self.speakers) < count:
            new_idx = len(self.speakers)
            self.speakers.append(chr(65 + new_idx))
            # Colors will be handled by rebuild
        
        # Rebuild colors from palette based on current speaker count
        self.speaker_colors = []
        for i in range(len(self.speakers)):
            if i < len(self.speaker_color_palette):
                # Use color from fixed palette
                self.speaker_colors.append(self.speaker_color_palette[i])
            else:
                # Fallback for more than 8 speakers (though we limit to 8 in UI)
                self.speaker_colors.append(QColor(200, 200, 200))
        
        self.create_speaker_widgets()
        self.setup_shortcuts()  # Recreate shortcuts for new count
        self.update_display()
        self.mark_unsaved_changes()
        
        # Update the count display and button states
        self.speaker_count_label.setText(str(count))
        self.update_speaker_buttons()
    
    def load_file(self):
        if not self.check_unsaved_changes():
            return
            
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", 
            "All Supported Files (*.srt *.txt *.json *.tsv);;SRT Files (*.srt);;Text Files (*.txt);;JSON Files (*.json);;TSV Files (*.tsv)"
        )
        if file_path:
            self.load_file_from_path(file_path)
    
    def parse_srt(self, content):
        blocks = []
        srt_blocks = content.strip().split('\n\n')
        
        for block in srt_blocks:
            lines = block.strip().split('\n')
            if len(lines) >= 3:
                try:
                    index = int(lines[0].strip())
                    time_line = lines[1].strip()
                    # Fixed regex to properly capture milliseconds
                    time_match = re.match(r'(\d{2}):(\d{2}):(\d{2}),(\d{3}) --> (\d{2}):(\d{2}):(\d{2}),(\d{3})', time_line)
                    
                    if time_match:
                        text = '\n'.join(lines[2:]).strip()
                        # Include milliseconds in the time string
                        block_data = {
                            'index': index,
                            'start_time': f"{time_match.group(1)}:{time_match.group(2)}:{time_match.group(3)},{time_match.group(4)}",
                            'start_ms': int(time_match.group(4)),
                            'end_time': f"{time_match.group(5)}:{time_match.group(6)}:{time_match.group(7)},{time_match.group(8)}",
                            'end_ms': int(time_match.group(8)),
                            'text': text,
                            'raw_text': text,   # NEW: store raw text separately
                            'speaker': None,
                            'is_turn_start': True,
                        }
                        blocks.append(block_data)
                except ValueError:
                    continue
        
        return blocks
    
    def parse_text(self, content):
        blocks = []
        lines = content.strip().split('\n')
        
        for i, line in enumerate(lines):
            if line.strip():
                block_data = {
                    'index': i + 1,
                    'start_time': '',
                    'end_time': '',
                    'text': line.strip(),
                    'raw_text': line.strip(),
                    'speaker': None,
                    'is_turn_start': True,
                }
                blocks.append(block_data)
        
        return blocks
    
    def parse_tsv(self, content):
        """Parse TSV file with start, end, and text columns"""
        blocks = []
        lines = content.strip().split('\n')
        
        for i, line in enumerate(lines):
            if i == 0:  # Skip header
                continue
                
            parts = line.split('\t')
            if len(parts) >= 3:
                start_ms = int(parts[0])
                end_ms = int(parts[1])
                text = parts[2]
                
                start_time = self.ms_to_srt_time(start_ms)
                end_time = self.ms_to_srt_time(end_ms)
                
                block_data = {
                    'index': i,
                    'start_time': start_time,
                    'end_time': end_time,
                    'text': text,
                    'raw_text': text,
                    'speaker': None,
                    'is_turn_start': True,
                }
                blocks.append(block_data)
        
        return blocks
    
    def ms_to_srt_time(self, ms):
        """Convert milliseconds to SRT time format (HH:MM:SS,mmm)"""
        hours = ms // 3600000
        ms %= 3600000
        minutes = ms // 60000
        ms %= 60000
        seconds = ms // 1000
        milliseconds = ms % 1000
        
        return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"
    
    def auto_segment_tokens(self, tokens, timestamps):
        """Auto-segment tokens based on pause detection"""
        if len(tokens) != len(timestamps) or len(tokens) < 2:
            return [{'text': ''.join(tokens), 'start_time': '', 'end_time': ''}]
        
        gaps = []
        for i in range(1, len(timestamps)):
            gap = timestamps[i] - timestamps[i-1]
            gaps.append(gap)
        
        avg_gap = sum(gaps) / len(gaps)
        threshold = avg_gap * 2.5
        
        segments = []
        current_segment = []
        current_start = timestamps[0]
        
        for i, (token, timestamp) in enumerate(zip(tokens, timestamps)):
            current_segment.append(token)
            
            if i < len(timestamps) - 1:
                gap = timestamps[i+1] - timestamp
                if gap > threshold:
                    segment_text = ''.join(current_segment)
                    segments.append({
                        'text': segment_text,
                        'start_time': self.seconds_to_srt_time(current_start),
                        'end_time': self.seconds_to_srt_time(timestamp + gap/2)
                    })
                    current_segment = []
                    if i < len(timestamps) - 1:
                        current_start = timestamps[i+1]
        
        if current_segment:
            segment_text = ''.join(current_segment)
            segments.append({
                'text': segment_text,
                'start_time': self.seconds_to_srt_time(current_start),
                'end_time': self.seconds_to_srt_time(timestamps[-1] + avg_gap)
            })
        
        return segments
    
    def seconds_to_srt_time(self, seconds):
        """Convert seconds to SRT time format"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = seconds % 60
        milliseconds = int((secs - int(secs)) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{int(secs):02d},{milliseconds:03d}"
    
    def parse_json(self, content):
        blocks = []
        
        try:
            if isinstance(content, dict) and 'tokens' in content and 'timestamps' in content:
                tokens = content['tokens']
                timestamps = content['timestamps']
                
                dialog = JsonImportDialog(has_tokens=True, parent=self)
                if dialog.exec_() == QDialog.Accepted:
                    option = dialog.get_import_option()
                    
                    if option == "one_block":
                        text = ''.join(tokens) if tokens else ""
                        block_data = {
                            'index': 1,
                            'start_time': '',
                            'end_time': '',
                            'text': text,
                            'raw_text': text,
                            'speaker': None,
                            'is_turn_start': True,
                        }
                        blocks.append(block_data)
                        
                    elif option == "tokens":
                        for i, (token, timestamp) in enumerate(zip(tokens, timestamps)):
                            block_data = {
                                'index': i + 1,
                                'start_time': self.seconds_to_srt_time(timestamp),
                                'end_time': '',
                                'text': token,
                                'raw_text': token,
                                'speaker': None,
                                'is_turn_start': True,
                            }
                            blocks.append(block_data)
                            
                    elif option == "auto_segment":
                        segments = self.auto_segment_tokens(tokens, timestamps)
                        for i, segment in enumerate(segments):
                            block_data = {
                                'index': i + 1,
                                'start_time': segment['start_time'],
                                'end_time': segment['end_time'],
                                'text': segment['text'],
                                'raw_text': segment['text'],
                                'speaker': None,
                                'is_turn_start': True,
                            }
                            blocks.append(block_data)
                else:
                    return []
                
            elif isinstance(content, dict) and 'segments' in content:
                segments = content['segments']
                for i, segment in enumerate(segments):
                    block_data = {
                        'index': i + 1,
                        'start_time': self.seconds_to_srt_time(segment.get('start', 0)),
                        'end_time': self.seconds_to_srt_time(segment.get('end', 0)),
                        'text': segment.get('text', '').strip(),
                        'raw_text': segment.get('text', '').strip(),
                        'speaker': None,
                        'is_turn_start': True,
                    }
                    blocks.append(block_data)
                    
            elif isinstance(content, dict) and 'text' in content:
                block_data = {
                    'index': 1,
                    'start_time': '',
                    'end_time': '',
                    'text': content['text'].strip(),
                    'raw_text': content['text'].strip(),
                    'speaker': None,
                    'is_turn_start': True,
                }
                blocks.append(block_data)
                
            elif isinstance(content, list):
                for i, item in enumerate(content):
                    if isinstance(item, dict):
                        block_data = {
                            'index': i + 1,
                            'start_time': item.get('start_time', ''),
                            'end_time': item.get('end_time', ''),
                            'text': item.get('text', ''),
                            'raw_text': item.get('text', ''),
                            'speaker': None,
                            'is_turn_start': True,
                        }
                        blocks.append(block_data)
            elif isinstance(content, dict):
                transcript_data = content.get('transcript', content.get('blocks', []))
                if isinstance(transcript_data, list):
                    for i, item in enumerate(transcript_data):
                        if isinstance(item, dict):
                            block_data = {
                                'index': i + 1,
                                'start_time': item.get('start_time', ''),
                                'end_time': item.get('end_time', ''),
                                'text': item.get('text', ''),
                                'raw_text': item.get('text', ''),
                                'speaker': None,
                                'is_turn_start': True,
                            }
                            blocks.append(block_data)
            
            if not blocks:
                QMessageBox.warning(self, "JSON Format", 
                                   "The JSON file doesn't contain recognizable transcript data.")
                
        except Exception as e:
            QMessageBox.critical(self, "JSON Error", 
                               f"Could not parse JSON file: {str(e)}")
        
        return blocks
    
    def save_project(self, force_save_as=False):
        if not self.srt_blocks:
            return
            
        file_path = None
        
        if not force_save_as and self.current_file_path:
            if self.current_file_path.endswith('.capsgat'):
                file_path = self.current_file_path
            else:
                force_save_as = True
        
        if force_save_as or not file_path:
            default_name = ""
            if self.project_name:
                default_name = self.project_name.replace(" ", "_") + ".capsgat"
            elif self.current_file_path:
                original_stem = Path(self.current_file_path).stem
                default_name = f"{original_stem}.capsgat"
            else:
                default_name = "transcript_project.capsgat"
                
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Save Project As", 
                default_name,
                "CapsGAT Project Files (*.capsgat)"
            )
            
            if not file_path:
                return
                
            if not file_path.endswith('.capsgat'):
                file_path += '.capsgat'
        
        if file_path:
            try:
                project_data = {
                    'srt_blocks': self.srt_blocks,
                    'current_block_index': self.current_block_index,
                    'speakers': self.speakers,
                    'source_file': self.current_file_path,
                    'file_has_timestamps': self.file_has_timestamps,
                    'audio_file_path': self.audio_file_path,
                    'project_name': self.project_name,
                    'project_memo': self.project_memo,
                    'text_display_font': {
                        'family': self.text_display_font.family(),
                        'size': self.text_display_font.pointSize()
                    },
                    'viewer_theme': self.current_theme,
                    'playback_speed': self.playback_speed,
                    'cjk_mode': self.cjk_mode,
                    'timestamp_style': self.timestamp_style,
                    'custom_timestamp_pattern': self.custom_timestamp_pattern
                }
                
                with open(file_path, 'w', encoding='utf-8') as f:
                    json.dump(project_data, f, indent=2, ensure_ascii=False)
                    
                self.current_file_path = file_path
                self.clear_unsaved_changes()
                QMessageBox.information(self, "Success", f"Project saved to {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not save project: {str(e)}")
    
    def load_project(self):
        if not self.check_unsaved_changes():
            return
            
        file_path, _ = QFileDialog.getOpenFileName(self, "Load Project", "", "CapsGAT Project Files (*.capsgat)")
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    project_data = json.load(f)
                
                self.srt_blocks = project_data['srt_blocks']
                # Migrate old projects: ensure raw_text exists and remove any 'overlap' field
                for block in self.srt_blocks:
                    if 'raw_text' not in block:
                        block['raw_text'] = block['text']
                    # Remove any leftover overlap metadata
                    if 'overlap' in block:
                        del block['overlap']
                
                self.current_block_index = project_data['current_block_index']
                self.speakers = project_data['speakers']
                self.current_file_path = project_data.get('source_file', '')
                self.file_has_timestamps = project_data.get('file_has_timestamps', True)
                self.project_name = project_data.get('project_name', '')
                self.project_memo = project_data.get('project_memo', '')
                self.playback_speed = project_data.get('playback_speed', 1.0)
                self.cjk_mode = project_data.get('cjk_mode', False)
                self.timestamp_style = project_data.get('timestamp_style', 'curly')
                self.custom_timestamp_pattern = project_data.get('custom_timestamp_pattern', '{HH:MM:SS}')
                
                font_data = project_data.get('text_display_font')
                if font_data:
                    self.text_display_font = QFont(font_data['family'], font_data['size'])
                    self.text_display.setFont(self.text_display_font)
                
                audio_path = project_data.get('audio_file_path')
                if audio_path and Path(audio_path).exists():
                    self.audio_file_path = audio_path
                    # Get original duration
                    try:
                        info = sf.info(audio_path)
                        self.original_audio_duration = info.duration
                    except:
                        self.original_audio_duration = 0
                    
                    # Load audio at saved speed
                    self.load_audio_file_for_project(audio_path, self.playback_speed)
                    
                viewer_theme = project_data.get('viewer_theme', 'light')
                self.apply_viewer_theme(viewer_theme)
                
                # Update speed display
                self.speed_knob.value = self.playback_speed
                self.speed_knob.update()
                
                self.update_display()
                self.update_menu_state()
                self.clear_unsaved_changes()
                
                QMessageBox.information(self, "Success", f"Project loaded from {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not load project: {str(e)}")

    def load_audio_file_for_project(self, audio_path, speed):
        """Load audio file for a project (without showing file dialog)"""
        # Stop any existing audio
        if self.audio_player:
            self.audio_player.cleanup()
            self.audio_player = None

        # Try to create VLC player first
        try:
            self.audio_player = VlcAudioPlayer()
            vlc_ok = True
        except Exception:
            # Fallback to simple player
            if HAS_PYAUDIO:
                self.audio_player = SimpleAudioPlayer()
                vlc_ok = False
            else:
                logger.error("No audio backend available for project loading")
                return

        # Connect signals
        self.audio_player.playback_started.connect(self.on_playback_started)
        self.audio_player.playback_stopped.connect(self.on_playback_stopped)
        self.audio_player.position_changed.connect(self.on_position_changed)

        try:
            # Get original duration
            info = sf.info(audio_path)
            self.original_audio_duration = info.duration

            # Load audio file
            if not self.audio_player.load_file(audio_path):
                raise Exception("Failed to load audio file")

            self.audio_file_path = audio_path
            self.playback_speed = 1.0
            self.speed_knob.value = 1.0
            self.speed_knob.update()

            # Update UI
            audio_name = Path(audio_path).name
            self.audio_info_label.setText(f"Audio: {audio_name}")

            # Enable controls
            self.btn_play.setEnabled(True)
            self.btn_rewind.setEnabled(True)
            self.btn_forward.setEnabled(True)
            self.btn_jump_to.setEnabled(True)
            self.btn_play_segment.setEnabled(True)
            self.auto_sync_check.setEnabled(self.file_has_timestamps)
            self.auto_pause_check.setEnabled(True)
            self.audio_progress.setEnabled(True)

            self.update_speed_controls_state()

            logger.info(f"Audio loaded for project at normal speed: {audio_name}")

        except Exception as e:
            logger.error(f"Failed to load audio for project: {e}")
            QMessageBox.critical(self, "Error", f"Failed to load audio: {str(e)}")
            self.audio_player = None
            
    def apply_viewer_theme(self, theme):
        self.current_theme = theme
        if theme == "dark":
            self.text_display.setStyleSheet("""
                QTextEdit {
                    background-color: #2b2b2b;
                    color: #ffffff;
                    border: 2px solid #555;
                    border-radius: 5px;
                    padding: 10px;
                }
            """)
            self.current_info_label.setStyleSheet("""
                QLabel {
                    background-color: #3a3a3a;
                    color: #ffffff;
                    padding: 10px;
                    border: 2px solid #555;
                    border-radius: 5px;
                    font-weight: bold;
                }
            """)
            # Dark theme palette
            self.speaker_color_palette = [
                QColor(60, 80, 100),   # Dark blue
                QColor(100, 60, 60),   # Dark red
                QColor(60, 100, 60),   # Dark green
                QColor(100, 100, 60),  # Dark yellow
                QColor(80, 60, 100),   # Dark purple
                QColor(100, 70, 50),   # Dark orange
                QColor(50, 80, 80),    # Dark cyan
                QColor(100, 60, 80)    # Dark pink
            ]
        else:
            self.text_display.setStyleSheet("""
                QTextEdit {
                    background-color: #fafafa;
                    color: #000000;
                    border: 2px solid #ddd;
                    border-radius: 5px;
                    padding: 10px;
                }
            """)
            self.current_info_label.setStyleSheet("""
                QLabel {
                    background-color: #f0f0f0;
                    color: #000000;
                    padding: 10px;
                    border: 2px solid #ccc;
                    border-radius: 5px;
                    font-weight: bold;
                }
            """)
            # Light theme palette
            self.speaker_color_palette = [
                QColor(220, 240, 255),  # Light blue
                QColor(255, 220, 220),  # Light red
                QColor(220, 255, 220),  # Light green
                QColor(255, 255, 200),  # Light yellow
                QColor(230, 200, 255),  # Light purple
                QColor(255, 200, 150),  # Light orange
                QColor(200, 230, 230),  # Light cyan
                QColor(255, 210, 230)   # Light pink
            ]
        
        # Rebuild speaker colors from new palette
        self.speaker_colors = []
        for i in range(len(self.speakers)):
            if i < len(self.speaker_color_palette):
                self.speaker_colors.append(self.speaker_color_palette[i])
            else:
                self.speaker_colors.append(QColor(200, 200, 200))
        
        self.create_speaker_widgets()
        self.update_display()
    
    def replace_indent_placeholders(self, text, for_export=False):
        """Replace indent placeholders with spaces.
        If for_export is True and cjk_mode is True, replace each placeholder with two spaces.
        Otherwise replace with one space.
        """
        if for_export and self.cjk_mode:
            return text.replace(self.INDENT_PLACEHOLDER, '  ')
        else:
            return text.replace(self.INDENT_PLACEHOLDER, ' ')
    
    def update_display(self):
        if not self.srt_blocks:
            self.text_display.setPlainText("No content loaded")
            self.current_info_label.setText("No block selected")
            self.lbl_current.setText("Current: -/-")
            # Update unassigned counter
            self.unassigned_blocks_label.setText("Unassigned Segments (0/0):")
            return
        
        current_block = self.srt_blocks[self.current_block_index]
        speaker_name = self.speakers[current_block['speaker']] if current_block['speaker'] is not None else "UNASSIGNED"
        turn_indicator = " [TURN START]" if current_block.get('is_turn_start', True) else " [CONTINUATION]"
        self.current_info_label.setText(
            f"Block {current_block['index']} | Speaker: {speaker_name}{turn_indicator} | "
            f"Time: {current_block['start_time']} --> {current_block['end_time']}"
        )
        
        start_idx = max(0, self.current_block_index - self.context_blocks)
        end_idx = min(len(self.srt_blocks), self.current_block_index + self.context_blocks + 1)
        
        display_text = ""
        for i in range(start_idx, end_idx):
            block = self.srt_blocks[i]
            # Use raw_text directly (contains placeholders)
            display_text_block = block['raw_text']
            if i == self.current_block_index:
                display_text += f">> {display_text_block}\n\n"
            else:
                display_text += f"   {display_text_block}\n\n"
        
        self.text_display.setPlainText(display_text)
        
        self.colorize_display()
        
        self.lbl_current.setText(f"Current: {self.current_block_index + 1}/{len(self.srt_blocks)}")
        
        # Update unassigned list with counter
        self.unassigned_list.clear()
        unassigned_count = 0
        total_blocks = len(self.srt_blocks)
        
        for i, block in enumerate(self.srt_blocks):
            if block['speaker'] is None:
                unassigned_count += 1
                preview = block['text'][:50] + "..." if len(block['text']) > 50 else block['text']
                self.unassigned_list.addItem(f"{i+1}: {preview}")
        
        # Update label with counter
        self.unassigned_blocks_label.setText(f"Unassigned Segments ({unassigned_count}/{total_blocks}):")
    
    def colorize_display(self):
        cursor = self.text_display.textCursor()
        cursor.select(cursor.Document)
        
        format_normal = QTextCharFormat()
        cursor.setCharFormat(format_normal)
        
        start_idx = max(0, self.current_block_index - self.context_blocks)
        
        for i in range(start_idx, min(len(self.srt_blocks), self.current_block_index + self.context_blocks + 1)):
            block = self.srt_blocks[i]
            if block['speaker'] is not None and block['speaker'] < len(self.speaker_colors):
                color = self.speaker_colors[block['speaker']]
                
                block_pos = (i - start_idx) * 2
                
                cursor.movePosition(cursor.Start)
                for _ in range(block_pos):
                    cursor.movePosition(cursor.Down)
                
                cursor.movePosition(cursor.Down, cursor.KeepAnchor)
                
                block_format = QTextCharFormat()
                block_format.setBackground(color)
                cursor.setCharFormat(block_format)
        
        current_pos = (self.current_block_index - start_idx) * 2
        cursor.movePosition(cursor.Start)
        for _ in range(current_pos):
            cursor.movePosition(cursor.Down)
        
        cursor.movePosition(cursor.Down, cursor.KeepAnchor)
        current_format = QTextCharFormat()
        if self.current_theme == "dark":
            current_format.setBackground(QColor(120, 120, 200))
        else:
            current_format.setBackground(QColor(255, 240, 200))
        current_format.setFontWeight(QFont.Bold)
        cursor.setCharFormat(current_format)
        
        self.scroll_to_current_block()
    
    def scroll_to_current_block(self):
        cursor = self.text_display.textCursor()
        cursor.movePosition(cursor.Start)
        
        start_idx = max(0, self.current_block_index - self.context_blocks)
        blocks_before_current = self.current_block_index - start_idx
        
        for _ in range(blocks_before_current * 2):
            cursor.movePosition(cursor.Down)
        
        self.text_display.setTextCursor(cursor)
        self.text_display.ensureCursorVisible()
        
    def previous_block(self):
        if self.current_block_index > 0:
            self.current_block_index -= 1
            self.update_display()
            
    def next_block(self):
        if self.current_block_index < len(self.srt_blocks) - 1:
            self.current_block_index += 1
            self.update_display()
            
    def assign_speaker(self, speaker_idx):
        if not self.srt_blocks or speaker_idx >= len(self.speakers):
            return
            
        current_block = self.srt_blocks[self.current_block_index]
        current_block['speaker'] = speaker_idx
        
        is_turn_start = True
        if self.current_block_index > 0:
            prev_block = self.srt_blocks[self.current_block_index - 1]
            if prev_block['speaker'] == speaker_idx:
                is_turn_start = False
        
        current_block['is_turn_start'] = is_turn_start
        
        for i in range(self.current_block_index + 1, len(self.srt_blocks)):
            if self.srt_blocks[i]['speaker'] == speaker_idx:
                self.srt_blocks[i]['is_turn_start'] = False
            else:
                break
        
        # Update the unassigned list immediately
        self.update_unassigned_list()  # We'll create this helper method
        
        if not (self.is_playing and self.auto_sync_enabled):
            self.find_next_unassigned()
            
        self.mark_unsaved_changes()

    def update_unassigned_list(self):
        """Update just the unassigned segments list (without full display refresh)"""
        self.unassigned_list.clear()
        unassigned_count = 0
        total_blocks = len(self.srt_blocks)
        
        for i, block in enumerate(self.srt_blocks):
            if block['speaker'] is None:
                unassigned_count += 1
                preview = block['text'][:50] + "..." if len(block['text']) > 50 else block['text']
                self.unassigned_list.addItem(f"{i+1}: {preview}")
        
        # Update label with counter
        self.unassigned_blocks_label.setText(f"Unassigned Segments ({unassigned_count}/{total_blocks}):")
        
    def unassign_current(self):
        if not self.srt_blocks:
            return
            
        current_block = self.srt_blocks[self.current_block_index]
        current_block['speaker'] = None
        current_block['is_turn_start'] = True
        
        for i in range(self.current_block_index + 1, len(self.srt_blocks)):
            if i > 0 and self.srt_blocks[i]['speaker'] is not None:
                prev_speaker = self.srt_blocks[i-1]['speaker']
                current_speaker = self.srt_blocks[i]['speaker']
                self.srt_blocks[i]['is_turn_start'] = (prev_speaker != current_speaker)
        
        self.update_display()
        self.mark_unsaved_changes()
        
    def find_next_unassigned(self):
        start_index = self.current_block_index
        for i in range(1, len(self.srt_blocks) + 1):
            next_index = (start_index + i) % len(self.srt_blocks)
            if self.srt_blocks[next_index]['speaker'] is None:
                self.current_block_index = next_index
                self.update_display()
                return
        
        if self.current_block_index < len(self.srt_blocks) - 1:
            self.current_block_index += 1
            self.update_display()
            
    def split_current_block(self):
        if not self.srt_blocks:
            return
            
        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()
            
        current_block = self.srt_blocks[self.current_block_index]
        dialog = BlockSplitDialog(current_block['raw_text'], self)   # use raw_text
        
        if dialog.exec_() == QDialog.Accepted:
            split_pos = dialog.split_position
            if 0 < split_pos < len(current_block['raw_text']):
                text_before = current_block['raw_text'][:split_pos].strip()
                text_after = current_block['raw_text'][split_pos:].strip()
                
                if text_before and text_after:
                    if current_block.get('start_time') and current_block.get('end_time'):
                        original_end_time = current_block['end_time']
                        original_end_ms = self.time_to_ms(original_end_time)
                        
                        start_ms = self.time_to_ms(current_block['start_time'])
                        end_ms = original_end_ms
                        total_duration = end_ms - start_ms
                        
                        total_chars = len(text_before) + len(text_after)
                        before_proportion = len(text_before) / total_chars

                        split_ms = start_ms + int(total_duration * before_proportion)
                        split_ms = max(start_ms + 100, min(end_ms - 100, split_ms))

                        current_block['raw_text'] = text_before
                        current_block['text'] = text_before   # also update text for compatibility
                        current_block['end_time'] = self.ms_to_time(split_ms)

                        new_block = current_block.copy()
                        new_block['raw_text'] = text_after
                        new_block['text'] = text_after
                        new_block['index'] = max(block['index'] for block in self.srt_blocks) + 1
                        new_block['start_time'] = self.ms_to_time(split_ms)
                        new_block['end_time'] = original_end_time
                        new_block['speaker'] = None
                        new_block['is_turn_start'] = False

                        if current_block['speaker'] is not None:
                            new_block['speaker'] = current_block['speaker']
                            new_block['is_turn_start'] = False
                    else:
                        current_block['raw_text'] = text_before
                        current_block['text'] = text_before

                        new_block = current_block.copy()
                        new_block['raw_text'] = text_after
                        new_block['text'] = text_after
                        new_block['index'] = max(block['index'] for block in self.srt_blocks) + 1
                        new_block['speaker'] = None
                        new_block['is_turn_start'] = False

                        if current_block['speaker'] is not None:
                            new_block['speaker'] = current_block['speaker']
                            new_block['is_turn_start'] = False

                    self.srt_blocks.insert(self.current_block_index + 1, new_block)
                    self.update_display()
                    self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()

    def merge_with_next(self):
        if self.current_block_index >= len(self.srt_blocks) - 1:
            return

        current_block = self.srt_blocks[self.current_block_index]
        next_block = self.srt_blocks[self.current_block_index + 1]

        if current_block['speaker'] is None or current_block['speaker'] == next_block['speaker']:
            # Merge raw texts
            current_block['raw_text'] += " " + next_block['raw_text']
            current_block['text'] = current_block['raw_text']   # sync text
            current_block['end_time'] = next_block['end_time']

            if next_block.get('is_turn_start', False):
                current_block['is_turn_start'] = True

            del self.srt_blocks[self.current_block_index + 1]
            self.update_display()
            self.mark_unsaved_changes()

    def edit_current_block(self):
        if not self.srt_blocks:
            return

        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        current_block = self.srt_blocks[self.current_block_index]
        dialog = EditDialog(current_block['raw_text'], self)   # use raw_text

        if dialog.exec_() == QDialog.Accepted:
            new_text = dialog.get_text()
            current_block['raw_text'] = new_text
            current_block['text'] = new_text   # keep text in sync
            self.update_display()
            self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()

    def open_symbol_dialog(self):
        """Open the enhanced symbol dialog"""
        if not self.srt_blocks:
            return

        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        dialog = EnhancedSymbolDialog(self, initial_category=self.last_symbol_category)
        if dialog.exec_() == QDialog.Accepted:
            symbol_info = dialog.get_selected_symbol_info()
            self.handle_symbol_insertion(symbol_info)
            # Store the category that was just used
            self.last_symbol_category = dialog.current_category_index

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()

    def handle_gat2_symbol(self, option_index):
        """Handle GAT2 symbols by index."""
        symbols = ["(.)", "(-)", "(--)", "(---)", "(_._)", "(())", "<<>>", "[ ]", 
                   "°h", "°hh", "°hhh", "h°", "hh°", "hhh°"]
        if option_index == 4:          # measured pause (_._)
            self.handle_measured_pause()
        elif option_index == 5:        # comment (())
            self.handle_comment()
        elif option_index == 6:        # action <<>>
            self.handle_action()
        elif option_index == 7:        # overlap [ ]
            self.handle_overlap()
        else:                           # all other symbols (pauses, breaths)
            symbol = symbols[option_index]
            self.handle_pause(symbol)

    def handle_symbol_insertion(self, symbol_info):
        """Handle insertion of any symbol type"""
        category = symbol_info.get('category', '').lower()
        
        if category == 'dresing && pehl':
            self.handle_dresing_pehl_symbol(symbol_info)
        elif category == 'tiq':
            self.handle_tiq_symbol(symbol_info)
        elif category == 'custom':
            self.handle_custom_symbol(symbol_info)
        else:  # GAT2
            self.handle_gat2_symbol(symbol_info.get('index', 0))

    def handle_dresing_pehl_symbol(self, symbol_info):
        """Handle Dresing & Pehl symbols"""
        display = symbol_info.get('display', '')
        current_block = self.srt_blocks[self.current_block_index]
        
        if display in ["(.)", "(..)", "(...)"]:
            # Simple pauses – use simple placement dialog
            dialog = PlacementDialog(current_block['raw_text'], display, self)
            if dialog.exec_() == QDialog.Accepted:
                if dialog.create_new_line:
                    self.create_new_block_with_symbol(display)
                else:
                    pos = dialog.placement_position
                    new_raw = (current_block['raw_text'][:pos] + " " + display +
                                           " " + current_block['raw_text'][pos:]).strip()
                    current_block['raw_text'] = new_raw
                    current_block['text'] = new_raw
                self.update_display()
                self.mark_unsaved_changes()
        
        elif display == "(_)":
            # Measured pause – input seconds, then simple placement
            seconds, ok = QInputDialog.getInt(
                self, "Measured Pause",
                "Enter pause length in seconds:",
                value=2, min=1, max=60
            )
            if ok:
                symbol = f"({seconds})"
                dialog = PlacementDialog(current_block['raw_text'], symbol, self)
                if dialog.exec_() == QDialog.Accepted:
                    if dialog.create_new_line:
                        self.create_new_block_with_symbol(symbol)
                    else:
                        pos = dialog.placement_position
                        new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                               " " + current_block['raw_text'][pos:]).strip()
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()
        
        elif display == "//":
            # Overlap – use dedicated overlap handler (similar to GAT2)
            self.handle_dresing_overlap()
        
        elif display == "(   )":
            # Comment – input text, then simple placement
            comment, ok = QInputDialog.getText(self, "Comment", "Enter comment:")
            if ok and comment:
                symbol = f"({comment})"
                dialog = PlacementDialog(current_block['raw_text'], symbol, self)
                if dialog.exec_() == QDialog.Accepted:
                    if dialog.create_new_line:
                        self.create_new_block_with_symbol(symbol)
                    else:
                        pos = dialog.placement_position
                        new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                               " " + current_block['raw_text'][pos:]).strip()
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()
          
        elif display == "⏱":
            # Pre‑fill timestamp with current block's start time (if available)
            default_time = ""
            if current_block.get('start_time'):
                # Convert "HH:MM:SS,mmm" to "HH:MM:SS" or "MM:SS"
                time_str = current_block['start_time']
                if ',' in time_str:
                    time_str = time_str.split(',')[0]  # remove milliseconds
                parts = time_str.split(':')
                if len(parts) == 3 and parts[0] == '00':
                    # drop hours if zero
                    default_time = f"{parts[1]}:{parts[2]}"
                else:
                    default_time = time_str
            
            timestamp, ok = QInputDialog.getText(
                self, "Insert Timestamp",
                "Enter timestamp (e.g., 01:23):",
                text=default_time
            )
            if ok and timestamp:
                symbol = f"#{timestamp}#"
                dialog = PlacementDialog(current_block['raw_text'], symbol, self)
                if dialog.exec_() == QDialog.Accepted:
                    if dialog.create_new_line:
                        self.create_new_block_with_symbol(symbol)
                    else:
                        pos = dialog.placement_position
                        new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                               " " + current_block['raw_text'][pos:]).strip()
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()
                    
    def handle_dresing_overlap(self):
        """Handle Dresing & Pehl overlap marker // (similar to GAT2 overlap)"""
        if not self.srt_blocks or self.current_block_index == 0:
            QMessageBox.information(self, "Overlap Feature",
                                   "Overlap requires at least two consecutive blocks.")
            return

        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        current_block = self.srt_blocks[self.current_block_index]
        prev_block = self.srt_blocks[self.current_block_index - 1]

        # Ensure raw_text exists
        if 'raw_text' not in current_block:
            current_block['raw_text'] = current_block['text']
        if 'raw_text' not in prev_block:
            prev_block['raw_text'] = prev_block['text']

        # Select overlapping text in current block
        dialog = TextSelectionDialog(current_block['raw_text'], self)
        dialog.setWindowTitle("Select Overlapping Text in Current Block")
        if dialog.exec_() == QDialog.Accepted:
            start_pos, end_pos, selected_text = dialog.get_selection()
            if selected_text:
                # Select overlapping text in previous block
                prev_dialog = TextSelectionDialog(prev_block['raw_text'], self)
                prev_dialog.setWindowTitle("Select Overlapping Text in Previous Block")
                if prev_dialog.exec_() == QDialog.Accepted:
                    prev_start, prev_end, prev_selected = prev_dialog.get_selection()
                    if prev_selected:
                        # For Dresing & Pehl, we just insert // markers without indentation.
                        prev_before = prev_block['raw_text'][:prev_start]
                        prev_after = prev_block['raw_text'][prev_end:]
                        prev_block['raw_text'] = f"{prev_before}//{prev_selected}//{prev_after}"
                        prev_block['text'] = prev_block['raw_text']

                        curr_before = current_block['raw_text'][:start_pos]
                        curr_after = current_block['raw_text'][end_pos:]
                        current_block['raw_text'] = f"{curr_before}//{selected_text}//{curr_after}"
                        current_block['text'] = current_block['raw_text']

                        self.update_display()
                        self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()
     
    def handle_tiq_overlap(self):
        if not self.srt_blocks or self.current_block_index == 0:
            QMessageBox.information(self, "Overlap Feature",
                                   "Overlap requires at least two consecutive blocks.")
            return

        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        current_block = self.srt_blocks[self.current_block_index]
        prev_block = self.srt_blocks[self.current_block_index - 1]

        if 'raw_text' not in current_block:
            current_block['raw_text'] = current_block['text']
        if 'raw_text' not in prev_block:
            prev_block['raw_text'] = prev_block['text']

        curr_dialog = TextSelectionDialog(current_block['raw_text'], self)
        curr_dialog.setWindowTitle("Select Overlapping Text in Later Block")
        if curr_dialog.exec_() == QDialog.Accepted:
            curr_start, curr_end, curr_selected = curr_dialog.get_selection()
            if curr_selected:
                prev_dialog = TextSelectionDialog(prev_block['raw_text'], self)
                prev_dialog.setWindowTitle("Select Overlapping Text in Earlier Block")
                if prev_dialog.exec_() == QDialog.Accepted:
                    prev_start, prev_end, prev_selected = prev_dialog.get_selection()
                    if prev_selected:
                        # For TiQ, we only mark the later block with indentation and the '└' marker
                        indent = prev_start  # characters before overlap in earlier block
                        # Insert placeholders for indentation (visible as ␣ in viewer)
                        curr_before = current_block['raw_text'][:curr_start]
                        curr_after = current_block['raw_text'][curr_end:]
                        # Insert placeholders (one per required space)
                        indent_placeholders = self.INDENT_PLACEHOLDER * indent
                        current_block['raw_text'] = f"{curr_before}{indent_placeholders}└{curr_selected}{curr_after}"
                        current_block['text'] = current_block['raw_text']
                        # Optionally, we could mark the earlier block with something, but TiQ typically doesn't
                        self.update_display()
                        self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()

    def handle_tiq_symbol(self, symbol_info):
        """Handle TiQ symbols"""
        display = symbol_info.get('display', '')
        current_block = self.srt_blocks[self.current_block_index]
        
        if display == "(.)":
            # Short pause – simple placement
            dialog = PlacementDialog(current_block['raw_text'], "(.)", self)
            if dialog.exec_() == QDialog.Accepted:
                if dialog.create_new_line:
                    self.create_new_block_with_symbol("(.)")
                else:
                    pos = dialog.placement_position
                    new_raw = (current_block['raw_text'][:pos] + " (.) " + current_block['raw_text'][pos:]).strip()
                    current_block['raw_text'] = new_raw
                    current_block['text'] = new_raw
                self.update_display()
                self.mark_unsaved_changes()
        
        elif display == "(_)":
            # Measured pause – input seconds, then simple placement
            seconds, ok = QInputDialog.getInt(
                self, "Measured Pause",
                "Enter pause length in seconds:",
                value=2, min=1, max=60
            )
            if ok:
                symbol = f"({seconds})"
                dialog = PlacementDialog(current_block['raw_text'], symbol, self)
                if dialog.exec_() == QDialog.Accepted:
                    if dialog.create_new_line:
                        self.create_new_block_with_symbol(symbol)
                    else:
                        pos = dialog.placement_position
                        new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                               " " + current_block['raw_text'][pos:]).strip()
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()
        
        elif display == "(())":
            # Comment – input text, then simple placement
            comment, ok = QInputDialog.getText(self, "Comment", "Enter comment:")
            if ok and comment:
                symbol = f"(({comment}))"
                dialog = PlacementDialog(current_block['raw_text'], symbol, self)
                if dialog.exec_() == QDialog.Accepted:
                    if dialog.create_new_line:
                        self.create_new_block_with_symbol(symbol)
                    else:
                        pos = dialog.placement_position
                        new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                               " " + current_block['raw_text'][pos:]).strip()
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()
        
        elif display == "└":
                    self.handle_tiq_overlap()
                    self.update_display()
                    self.mark_unsaved_changes()
        
        elif display == "@(.)@":
            # Short laughter – simple placement
            dialog = PlacementDialog(current_block['raw_text'], "@(.)@", self)
            if dialog.exec_() == QDialog.Accepted:
                if dialog.create_new_line:
                    self.create_new_block_with_symbol("@(.)@")
                else:
                    pos = dialog.placement_position
                    new_raw = (current_block['raw_text'][:pos] + " @(.)@ " + current_block['raw_text'][pos:]).strip()
                    current_block['raw_text'] = new_raw
                    current_block['text'] = new_raw
                self.update_display()
                self.mark_unsaved_changes()
        
        elif display == "@(_)@":
            # Laughing seconds – input, then simple placement
            seconds, ok = QInputDialog.getInt(
                self, "Laughing Duration",
                "Enter laughter duration in seconds:",
                value=2, min=1, max=10
            )
            if ok:
                symbol = f"@({seconds}s)@"
                dialog = PlacementDialog(current_block['raw_text'], symbol, self)
                if dialog.exec_() == QDialog.Accepted:
                    if dialog.create_new_line:
                        self.create_new_block_with_symbol(symbol)
                    else:
                        pos = dialog.placement_position
                        new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                               " " + current_block['raw_text'][pos:]).strip()
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()
        
        elif display in ["@(   )@", "°   °", "//   //"]:
            was_playing = False
            if self.auto_pause_enabled and self.is_playing:
                was_playing = True
                if self.audio_player:
                    self.audio_player.pause()
            
            # Determine left/right markers
            if display == "@(   )@":
                left, right = "@(", ")@"
            elif display == "°   °":
                left, right = "°", "°"
            else:  # "//   //"
                left, right = "//", "//"
            
            dialog = TextSelectionDialog(current_block['raw_text'], self)
            dialog.setWindowTitle(f"Select text for {display}")
            if dialog.exec_() == QDialog.Accepted:
                start_pos, end_pos, selected_text = dialog.get_selection()
                if selected_text:
                    new_raw = (current_block['raw_text'][:start_pos] + 
                                            left + selected_text + right + 
                                            current_block['raw_text'][end_pos:])
                    current_block['raw_text'] = new_raw
                    current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()
            
            if was_playing and self.auto_pause_enabled:
                if self.audio_player:
                    self.audio_player.play()
                

    def handle_custom_symbol(self, symbol_info):
        """Handle insertion of custom symbols according to their type"""
        symbol_type = symbol_info.get('type', 'simple')
        current_block = self.srt_blocks[self.current_block_index]

        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        # --- Simple symbol: use PlacementDialog ---
        if symbol_type == 'simple':
            symbol = symbol_info.get('display', '')
            dialog = PlacementDialog(current_block['raw_text'], symbol, self)
            if dialog.exec_() == QDialog.Accepted:
                if dialog.create_new_line:
                    self.create_new_block_with_symbol(symbol)
                else:
                    pos = dialog.placement_position
                    new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                           " " + current_block['raw_text'][pos:]).strip()
                    current_block['raw_text'] = new_raw
                    current_block['text'] = new_raw
                self.update_display()
                self.mark_unsaved_changes()

        # --- Segment wrapper: select text, then wrap with left/right ---
        elif symbol_type == 'wrapper':
            left = symbol_info.get('left', '')
            right = symbol_info.get('right', '')
            dialog = TextSelectionDialog(current_block['raw_text'], self)
            dialog.setWindowTitle("Select text to wrap")
            if dialog.exec_() == QDialog.Accepted:
                start_pos, end_pos, selected_text = dialog.get_selection()
                if selected_text:
                    new_raw = (current_block['raw_text'][:start_pos] + 
                                            left + selected_text + right + 
                                            current_block['raw_text'][end_pos:])
                    current_block['raw_text'] = new_raw
                    current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()

        # --- Comment wrapper: input comment, then use PlacementDialog ---
        elif symbol_type == 'comment':
            left = symbol_info.get('left', '')
            right = symbol_info.get('right', '')
            comment, ok = QInputDialog.getText(self, "Comment", "Enter comment:")
            if ok and comment:
                symbol = left + comment + right
                # Use PlacementDialog to allow cursor positioning or new line
                dialog = PlacementDialog(current_block['raw_text'], symbol, self)
                if dialog.exec_() == QDialog.Accepted:
                    if dialog.create_new_line:
                        self.create_new_block_with_symbol(symbol)
                    else:
                        pos = dialog.placement_position
                        new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                               " " + current_block['raw_text'][pos:]).strip()
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()

        # --- Comment wrapper with reach: select text, input description, then wrap ---
        elif symbol_type == 'comment_reach':
            left = symbol_info.get('left', '')
            right_action = symbol_info.get('right', '')
            right_segment = symbol_info.get('segment_right', '')

            # Step 1: select the spoken text
            select_dialog = TextSelectionDialog(current_block['raw_text'], self)
            select_dialog.setWindowTitle("Select spoken text to annotate")
            if select_dialog.exec_() == QDialog.Accepted:
                start_pos, end_pos, selected_text = select_dialog.get_selection()
                if selected_text:
                    # Step 2: input the action/attitude description
                    description, ok = QInputDialog.getText(self, "Action Description",
                                                           f"Enter description for the action/attitude:")
                    if ok and description:
                        # Construct the wrapped result: left + description + right_action + selected_text + right_segment
                        wrapped = left + description + right_action + selected_text + right_segment
                        new_raw = (current_block['raw_text'][:start_pos] + 
                                                wrapped + 
                                                current_block['raw_text'][end_pos:])
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw
                        self.update_display()
                        self.mark_unsaved_changes()
                    
        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()
            
    def create_new_block_with_symbol(self, symbol_text):
        """Create a new block containing just the symbol"""
        new_block = {
            'index': max(block['index'] for block in self.srt_blocks) + 1,
            'start_time': '',
            'end_time': '',
            'text': symbol_text,
            'raw_text': symbol_text,
            'speaker': None,
            'is_turn_start': False,
            'is_pause': True,
        }
        self.srt_blocks.insert(self.current_block_index + 1, new_block)
        self.update_display()
        self.mark_unsaved_changes()

    # Update the existing open_pause_dialog method to call the new one
    def open_pause_dialog(self):
        """Legacy method - calls the new enhanced dialog"""
        self.open_symbol_dialog()

    def handle_measured_pause(self):
        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        time_ms, ok = QInputDialog.getInt(
            self,
            "Measured Pause",
            "Enter pause length in 100 milliseconds (e.g., 8 for 0.8 seconds):",
            value=5, min=1, max=50, step=1
        )

        if ok:
            seconds = time_ms / 10.0
            symbol = f"({seconds:.1f})"

            current_block = self.srt_blocks[self.current_block_index]
            dialog = PlacementDialog(current_block['raw_text'], symbol, self)

            if dialog.exec_() == QDialog.Accepted:
                if dialog.create_new_line:
                    self.create_new_block_with_symbol(symbol)
                else:
                    pos = dialog.placement_position
                    new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                           " " + current_block['raw_text'][pos:]).strip()
                    current_block['raw_text'] = new_raw
                    current_block['text'] = new_raw

                self.update_display()
                self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()

    def handle_pause(self, symbol):
        if not self.srt_blocks:
            return

        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        current_block = self.srt_blocks[self.current_block_index]
        dialog = PlacementDialog(current_block['raw_text'], symbol, self)

        if dialog.exec_() == QDialog.Accepted:
            if dialog.create_new_line:
                self.create_new_block_with_symbol(symbol)
            else:
                pos = dialog.placement_position
                new_raw = (current_block['raw_text'][:pos] + " " + symbol +
                                       " " + current_block['raw_text'][pos:]).strip()
                current_block['raw_text'] = new_raw
                current_block['text'] = new_raw

            self.update_display()
            self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()
                
    def handle_comment(self):
        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        dialog = CommentDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            comment = dialog.get_comment()
            if comment:
                current_block = self.srt_blocks[self.current_block_index]
                placement_dialog = PlacementDialog(current_block['raw_text'], comment, self)

                if placement_dialog.exec_() == QDialog.Accepted:
                    if placement_dialog.create_new_line:
                        self.create_new_block_with_symbol(comment)
                    else:
                        pos = placement_dialog.placement_position
                        new_raw = (current_block['raw_text'][:pos] + " " + comment +
                                               " " + current_block['raw_text'][pos:]).strip()
                        current_block['raw_text'] = new_raw
                        current_block['text'] = new_raw

                    self.update_display()
                    self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()

    def handle_action(self):
        if not self.srt_blocks:
            return

        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        current_block = self.srt_blocks[self.current_block_index]
        dialog = TextSelectionDialog(current_block['raw_text'], self)

        if dialog.exec_() == QDialog.Accepted:
            start_pos, end_pos, selected_text = dialog.get_selection()
            if selected_text:
                action_text, ok = QInputDialog.getText(self, "Action Description",
                                                     f"Describe the action for '{selected_text}':")
                if ok and action_text:
                    before_text = current_block['raw_text'][:start_pos]
                    after_text = current_block['raw_text'][end_pos:]
                    new_raw = f"{before_text}<<{action_text}> {selected_text}>{after_text}"
                    current_block['raw_text'] = new_raw
                    current_block['text'] = new_raw
                    self.update_display()
                    self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()

    def handle_overlap(self):
        """GAT2 overlap: both blocks get brackets, and later block gets indentation."""
        if not self.srt_blocks or self.current_block_index == 0:
            QMessageBox.information(self, "Overlap Feature",
                                   "Overlap requires at least two consecutive blocks.")
            return

        was_playing = False
        if self.auto_pause_enabled and self.is_playing:
            was_playing = True
            if self.audio_player:
                self.audio_player.pause()

        current_block = self.srt_blocks[self.current_block_index]
        prev_block = self.srt_blocks[self.current_block_index - 1]

        # Ensure raw_text exists
        if 'raw_text' not in current_block:
            current_block['raw_text'] = current_block['text']
        if 'raw_text' not in prev_block:
            prev_block['raw_text'] = prev_block['text']

        # Select overlapping text in current block
        curr_dialog = TextSelectionDialog(current_block['raw_text'], self)
        curr_dialog.setWindowTitle("Select Overlapping Text in Current Block")
        if curr_dialog.exec_() == QDialog.Accepted:
            curr_start, curr_end, curr_selected = curr_dialog.get_selection()
            if curr_selected:
                # Select overlapping text in previous block
                prev_dialog = TextSelectionDialog(prev_block['raw_text'], self)
                prev_dialog.setWindowTitle("Select Overlapping Text in Previous Block")
                if prev_dialog.exec_() == QDialog.Accepted:
                    prev_start, prev_end, prev_selected = prev_dialog.get_selection()
                    if prev_selected:
                        # Modify previous block: insert brackets around the overlapping text
                        prev_before = prev_block['raw_text'][:prev_start]
                        prev_after = prev_block['raw_text'][prev_end:]
                        prev_block['raw_text'] = f"{prev_before}[{prev_selected}]{prev_after}"
                        prev_block['text'] = prev_block['raw_text']

                        # Compute indentation length = number of characters before the overlap in the previous block
                        indent = len(prev_before)

                        # Modify current block: insert placeholders and brackets around the overlapping text
                        curr_before = current_block['raw_text'][:curr_start]
                        curr_after = current_block['raw_text'][curr_end:]
                        indent_placeholders = self.INDENT_PLACEHOLDER * indent
                        current_block['raw_text'] = f"{curr_before}{indent_placeholders}[{curr_selected}]{curr_after}"
                        current_block['text'] = current_block['raw_text']

                        self.update_display()
                        self.mark_unsaved_changes()

        if was_playing and self.auto_pause_enabled:
            if self.audio_player:
                self.audio_player.play()

    def insert_empty_line(self):
        if not self.srt_blocks:
            return

        new_block = {
            'index': max(block['index'] for block in self.srt_blocks) + 1,
            'start_time': '',
            'end_time': '',
            'text': '',
            'raw_text': '',
            'speaker': None,
            'is_turn_start': False,
            'is_empty': True,
        }
        self.srt_blocks.insert(self.current_block_index + 1, new_block)
        self.current_block_index += 1
        self.update_display()
        self.mark_unsaved_changes()

    def jump_to_block(self, item):
        text = item.text()
        match = re.match(r'(\d+):', text)
        if match:
            block_idx = int(match.group(1)) - 1
            if 0 <= block_idx < len(self.srt_blocks):
                self.current_block_index = block_idx
                self.update_display()

    def export_transcript(self):
        
        if not self.srt_blocks:
            return

        project_info = {
            'name': self.project_name,
            'memo': self.project_memo
        }

        preview_dialog = ExportPreviewDialog(
            self,
            has_timestamps=self.file_has_timestamps,
            timestamp_style=self.timestamp_style,
            custom_pattern=self.custom_timestamp_pattern,
            project_info=project_info,
            audio_path=self.audio_file_path
        )
        if preview_dialog.exec_() == QDialog.Accepted:
            settings = preview_dialog.get_export_settings()

            unassigned_handling = "skip"
            if settings['format'] == 'srt':
                unassigned_count = sum(1 for block in self.srt_blocks
                                     if block['speaker'] is None and not block.get('is_pause')
                                     and not block.get('is_comment') and not block.get('is_empty'))

                if unassigned_count > 0:
                    dialog = UnassignedSegmentsDialog(unassigned_count, self)
                    if dialog.exec_() == QDialog.Accepted:
                        unassigned_handling = dialog.get_selected_option()
                    else:
                        return
                else:
                    unassigned_handling = "skip"

            if settings['format'] == 'srt':
                transcript_text = self.generate_srt_text(
                    include_diarization=settings['include_diarization'],
                    unassigned_handling=unassigned_handling
                )
            else:
                transcript_text = self.generate_transcript_text(
                    include_timestamps=settings['include_timestamps'],
                    timestamp_style=settings.get('timestamp_style', 'hash'),
                    custom_pattern=settings.get('custom_timestamp_pattern', None),
                    convention=settings['convention'],
                    include_diarization=settings['include_diarization'],
                    wrap_enabled=settings['wrap_enabled'],
                    wrap_length=settings['wrap_length'],
                    character_wrap=settings['character_wrap']
                )

            self.final_export(transcript_text, settings, project_info, unassigned_handling)


    def generate_transcript_text(self, include_timestamps=True,
                                 timestamp_style="hash", custom_pattern=None,
                                 convention="gat2", include_diarization=True,
                                 wrap_enabled=False, wrap_length=80, character_wrap=False):
        if convention == "dresing_pehl":
            return self.generate_dresing_pehl_text(
                include_timestamps, timestamp_style, custom_pattern, include_diarization)
        elif convention == "tiq":
            return self.generate_tiq_text(
                include_timestamps, timestamp_style, custom_pattern, include_diarization,
                wrap_enabled, wrap_length, character_wrap)
        else:  # gat2
            return self.generate_gat2_text(
                include_timestamps, timestamp_style, custom_pattern, include_diarization,
                wrap_enabled, wrap_length, character_wrap)
        
    def generate_gat2_text(self, include_timestamps=True,
                           timestamp_style="curly", custom_pattern=None,
                           include_diarization=True,
                           wrap_enabled=False, wrap_length=80, character_wrap=False):
        if not self.srt_blocks:
            return ""

        # Filter blocks
        included_blocks = [b for b in self.srt_blocks
                           if b['speaker'] is not None or b.get('is_pause')
                           or b.get('is_comment') or b.get('is_empty')]
        if not included_blocks:
            return ""

        # Max speaker label length
        max_speaker_length = 2
        for b in included_blocks:
            if b['speaker'] is not None and b.get('is_turn_start', True):
                max_speaker_length = max(max_speaker_length, len(self.speakers[b['speaker']] + ":"))

        total_lines = len(included_blocks)
        line_digits = len(str(total_lines))
        
        ts_width = self.get_timestamp_width(timestamp_style, custom_pattern) if include_timestamps else 0
        timestamp_padding = " " * (ts_width + 3)

        output_lines = []

        for line_num, block in enumerate(included_blocks, start=1):
            # Build left part
            if include_timestamps and block.get('is_turn_start', True) and block['speaker'] is not None:
                if block.get('start_time'):
                    seconds = self.time_to_seconds(block['start_time'])
                    ts = self.format_timestamp(seconds, timestamp_style, custom_pattern)
                    timestamp = f"{ts}   "
                else:
                    timestamp = timestamp_padding
            else:
                timestamp = timestamp_padding

            line_number_part = f"{line_num:0{line_digits}d}   "

            if block.get('is_turn_start', True) and block['speaker'] is not None:
                speaker_part = self.speakers[block['speaker']] + ":"
                speaker_part = speaker_part.ljust(max_speaker_length) + "   "
            else:
                speaker_part = " " * (max_speaker_length + 3)

            left_part = timestamp + line_number_part + speaker_part

            # Replace placeholders for export
            text = self.replace_indent_placeholders(block['raw_text'], for_export=True)

            if wrap_enabled and wrap_length > 0:
                available_width = wrap_length - len(left_part)
                if available_width < 10:
                    available_width = 40
                lines = self.wrap_text(text, available_width, character_wrap, first_line_only_indent=True)
                for idx, line in enumerate(lines):
                    if idx == 0:
                        output_lines.append(left_part + line)
                    else:
                        output_lines.append(' ' * len(left_part) + line)
            else:
                output_lines.append(left_part + text)

        return '\n'.join(output_lines)

    def generate_dresing_pehl_text(self, include_timestamps=True,
                                   timestamp_style="hash", custom_pattern=None,
                                   include_diarization=True):
        """Generate Dresing & Pehl format transcript (sociological interviews)."""
        if not self.srt_blocks:
            return ""

        turns = []
        current_turn = None

        for block in self.srt_blocks:
            if block.get('is_pause') or block.get('is_comment') or block.get('is_empty'):
                continue
            if block['speaker'] is None:
                continue

            speaker_label = self.speakers[block['speaker']]

            if current_turn is None or current_turn['speaker'] != speaker_label or block.get('is_turn_start', True):
                if current_turn is not None:
                    turns.append(current_turn)

                current_turn = {
                    'speaker': speaker_label,
                    'blocks': [],
                    'start_time': block['start_time'] if include_timestamps else None,
                    'end_time': block['end_time'] if include_timestamps else None
                }

            current_turn['blocks'].append(block)
            if include_timestamps and block['end_time']:
                current_turn['end_time'] = block['end_time']

        if current_turn is not None:
            turns.append(current_turn)

        output_lines = []
        output_lines.append("")  # blank line at top

        for turn in turns:
            # Concatenate all block texts, replacing placeholders with spaces for export
            turn_text = " ".join(
                self.replace_indent_placeholders(b['raw_text'], for_export=True).strip()
                for b in turn['blocks']
            )

            if include_diarization:
                line = f"{turn['speaker']}: {turn_text}"
            else:
                line = turn_text

            if include_timestamps and turn['start_time']:
                seconds = self.time_to_seconds(turn['start_time'])
                ts = self.format_timestamp(seconds, timestamp_style, custom_pattern)
                line += f" {ts}"

            output_lines.append(line)
            output_lines.append("")  # blank line after each turn

        return '\n'.join(output_lines)
    
    def _group_into_turns(self, include_timestamps=False):
        """Group blocks by speaker turns, ignoring pause/comment blocks.
        Returns list of turns with keys: speaker, blocks, start_time, end_time."""
        turns = []
        current_turn = None
        for block in self.srt_blocks:
            if block.get('is_pause') or block.get('is_comment') or block.get('is_empty'):
                continue
            if block['speaker'] is None:
                continue

            speaker_label = self.speakers[block['speaker']]
            if current_turn is None or current_turn['speaker'] != speaker_label or block.get('is_turn_start', True):
                if current_turn is not None:
                    turns.append(current_turn)
                current_turn = {
                    'speaker': speaker_label,
                    'blocks': [],
                    'start_time': block['start_time'] if include_timestamps else None
                }
            current_turn['blocks'].append(block)
            if include_timestamps and block['end_time']:
                current_turn['end_time'] = block['end_time']
        if current_turn is not None:
            turns.append(current_turn)
        return turns
    
    def generate_tiq_text(self, include_timestamps=True,
                          timestamp_style="hash", custom_pattern=None,
                          include_diarization=True,
                          wrap_enabled=False, wrap_length=80, character_wrap=False):
        """Generate TiQ format transcript.
        - Each speaker turn is a single logical line (concatenated blocks).
        - After wrapping, each physical display line gets a line number.
        - Speaker label only on the first display line of a turn.
        - Timestamp only on the last display line of a turn.
        - Special blocks (pauses, comments) are separate logical lines.
        """
        if not self.srt_blocks:
            return ""

        # Filter blocks: keep assigned or special
        included_blocks = []
        for block in self.srt_blocks:
            if block['speaker'] is not None or block.get('is_pause') or block.get('is_comment') or block.get('is_empty'):
                included_blocks.append(block)
        if not included_blocks:
            return ""

        # Group speaker blocks into turns
        turns = self._group_into_turns(include_timestamps)  # returns list of dicts with 'speaker', 'blocks', 'start_time'

        # Separate special blocks (pauses, comments, empty) – they will be handled individually
        special_blocks = [b for b in included_blocks if b['speaker'] is None]

        # Determine max speaker label width (for padding)
        max_speaker_width = 0
        if include_diarization:
            for speaker in self.speakers:
                label = f"{speaker}: "
                max_speaker_width = max(max_speaker_width, len(label))
        else:
            max_speaker_width = 0

        # First pass: generate content lines (without line numbers)
        content_lines = []  # list of strings (actual text, with prefixes and timestamps)

        # Process turns
        for turn in turns:
            turn_text = " ".join(
                self.replace_indent_placeholders(b['raw_text'], for_export=True)
                for b in turn['blocks'] if b['text'].strip()
            )
            if not turn_text:
                continue

            # Prepare speaker prefix for first line
            if include_diarization:
                speaker_prefix = f"{turn['speaker']}: ".ljust(max_speaker_width)
            else:
                speaker_prefix = " " * max_speaker_width

            # Wrap the turn text (without timestamp)
            if wrap_enabled and wrap_length > 0:
                text_width = wrap_length - len(speaker_prefix) - 5
                if text_width < 10:
                    text_width = 40
                wrapped_lines = self.wrap_text(turn_text, text_width, character_wrap, first_line_only_indent=True)
            else:
                wrapped_lines = [turn_text]

            # Build display lines
            for idx, line in enumerate(wrapped_lines):
                if idx == 0:
                    display_line = speaker_prefix + line
                else:
                    display_line = " " * len(speaker_prefix) + line
                content_lines.append(display_line)

            # Add timestamp to the last line of this turn
            if include_timestamps and turn.get('start_time'):
                seconds = self.time_to_seconds(turn['start_time'])
                ts = self.format_timestamp(seconds, timestamp_style, custom_pattern)
                content_lines[-1] += " " + ts
                
        # Process special blocks (each gets its own line, no speaker label)
        for block in special_blocks:
            text = self.replace_indent_placeholders(block['raw_text'], for_export=True).strip()
            if not text:
                continue
            # No speaker label, just spaces for alignment
            speaker_padding = " " * max_speaker_width
            if wrap_enabled and wrap_length > 0:
                placeholder_width = 5
                text_width = wrap_length - placeholder_width - len(speaker_padding)
                if text_width < 10:
                    text_width = 40
                wrapped = self.wrap_text(text, text_width, character_wrap)
                for idx, line in enumerate(wrapped):
                    if idx == 0:
                        display_line = speaker_padding + line
                    else:
                        display_line = " " * len(speaker_padding) + line
                    content_lines.append(display_line)
            else:
                display_line = speaker_padding + text
                content_lines.append(display_line)

        # Now we have all content lines. Determine required line number width
        total_lines = len(content_lines)
        line_digits = len(str(total_lines))

        # Second pass: add line numbers to each content line
        output_lines = []
        for idx, line in enumerate(content_lines, start=1):
            line_num = f"{idx:0{line_digits}d}"
            output_lines.append(f"{line_num} {line}")

        return '\n'.join(output_lines)

    def time_to_seconds(self, time_str):
        """Convert time string to seconds with milliseconds as decimal."""
        if not time_str:
            return 0

        if ',' in time_str:
            time_part, ms_part = time_str.split(',')
            ms = int(ms_part)
        else:
            time_part = time_str
            ms = 0

        parts = time_part.split(':')
        if len(parts) == 3:
            hours = int(parts[0])
            minutes = int(parts[1])
            seconds = int(parts[2])
        elif len(parts) == 2:
            hours = 0
            minutes = int(parts[0])
            seconds = int(parts[1])
        else:
            return 0

        return hours * 3600 + minutes * 60 + seconds + ms / 1000.0

    def generate_srt_text(self, include_diarization=True, unassigned_handling="skip"):
        """Generate SRT format text with optional diarization."""
        if not self.file_has_timestamps:
            return "SRT export requires timestamp information. Original file does not contain timestamps.\n\nNote: SRT files require precise timing information for each subtitle."

        blocks_with_timestamps = self.estimate_missing_timestamps()

        srt_blocks = []
        subtitle_index = 1

        for block in blocks_with_timestamps:
            if block.get('is_pause') or block.get('is_comment') or block.get('is_empty'):
                continue

            if block['speaker'] is None:
                if unassigned_handling == "skip":
                    continue
                elif unassigned_handling == "no_label":
                    speaker_prefix = ""
                elif unassigned_handling == "unknown":
                    speaker_prefix = "Unknown: "
            else:
                speaker_prefix = ""
                if include_diarization:
                    speaker_prefix = f"{self.speakers[block['speaker']]}: "

            # For SRT, we don't need indentation; replace placeholders with spaces (1 space each)
            formatted = self.replace_indent_placeholders(block['raw_text'], for_export=True).lstrip()
            start_time = self.format_srt_time(block['start_time'])
            end_time = self.format_srt_time(block['end_time'])

            srt_block = f"{subtitle_index}\n{start_time} --> {end_time}\n{speaker_prefix}{formatted}\n"
            srt_blocks.append(srt_block)
            subtitle_index += 1

        return "\n".join(srt_blocks)

    def format_srt_time(self, time_str):
        """Convert time string to SRT format (HH:MM:SS,mmm)."""
        if not time_str:
            return "00:00:00,000"

        if ',' in time_str:
            return time_str
        elif ':' in time_str:
            parts = time_str.split(':')
            if len(parts) == 2:
                return f"00:{parts[0].zfill(2)}:{parts[1].zfill(2)},000"
            elif len(parts) == 3:
                time_part = parts[2]
                if ',' in time_part:
                    time_part, ms_part = time_part.split(',')
                    return f"{parts[0].zfill(2)}:{parts[1].zfill(2)}:{time_part.zfill(2)},{ms_part.zfill(3)}"
                else:
                    return f"{parts[0].zfill(2)}:{parts[1].zfill(2)}:{time_part.zfill(2)},000"

        return "00:00:00,000"

    def estimate_missing_timestamps(self):
        """Estimate timestamps for blocks that don't have them."""
        if not self.srt_blocks:
            return []

        blocks = self.srt_blocks.copy()

        timestamped_blocks = []
        for i, block in enumerate(blocks):
            if block.get('start_time') and block.get('end_time'):
                timestamped_blocks.append((i, block))

        if not timestamped_blocks:
            return blocks

        segments = []
        last_timestamped_idx = -1

        for i, block in timestamped_blocks:
            if last_timestamped_idx == -1:
                segments.append({
                    'start_idx': 0,
                    'end_idx': i,
                    'start_time': None,
                    'end_time': block['start_time'],
                    'total_chars': sum(len(b['text']) for b in blocks[0:i])
                })
            else:
                segments.append({
                    'start_idx': last_timestamped_idx + 1,
                    'end_idx': i,
                    'start_time': blocks[last_timestamped_idx]['end_time'],
                    'end_time': block['start_time'],
                    'total_chars': sum(len(b['text']) for b in blocks[last_timestamped_idx + 1:i])
                })
            last_timestamped_idx = i

        if last_timestamped_idx < len(blocks) - 1:
            last_block = timestamped_blocks[-1][1] if timestamped_blocks else None
            segments.append({
                'start_idx': last_timestamped_idx + 1,
                'end_idx': len(blocks) - 1,
                'start_time': last_block['end_time'] if last_block else None,
                'end_time': None,
                'total_chars': sum(len(b['text']) for b in blocks[last_timestamped_idx + 1:])
            })

        for segment in segments:
            if segment['start_time'] and segment['end_time'] and segment['total_chars'] > 0:
                start_ms = self.time_to_ms(segment['start_time'])
                end_ms = self.time_to_ms(segment['end_time'])
                total_duration = end_ms - start_ms

                current_time = start_ms
                for i in range(segment['start_idx'], segment['end_idx'] + 1):
                    block = blocks[i]
                    if not block.get('start_time') or not block.get('end_time'):
                        block_chars = len(block['text'])
                        if segment['total_chars'] > 0:
                            block_duration = (block_chars / segment['total_chars']) * total_duration
                        else:
                            block_duration = 1000

                        block_duration = max(100, block_duration)

                        block['start_time'] = self.ms_to_time(current_time)
                        block['end_time'] = self.ms_to_time(current_time + block_duration)
                        current_time += block_duration

        return blocks

    def time_to_ms(self, time_str):
        """Convert time string to milliseconds."""
        if not time_str:
            return 0

        if ',' in time_str:
            time_part, ms_part = time_str.split(',')
            ms = int(ms_part)
        else:
            time_part = time_str
            ms = 0

        parts = time_part.split(':')
        if len(parts) == 3:
            hours = int(parts[0])
            minutes = int(parts[1])
            seconds = int(parts[2])
        elif len(parts) == 2:
            hours = 0
            minutes = int(parts[0])
            seconds = int(parts[1])
        else:
            return 0

        return (hours * 3600 + minutes * 60 + seconds) * 1000 + ms

    def ms_to_time(self, ms):
        """Convert milliseconds to SRT time format."""
        hours = int(ms // 3600000)
        ms %= 3600000
        minutes = int(ms // 60000)
        ms %= 60000
        seconds = int(ms // 1000)
        milliseconds = int(ms % 1000)

        return f"{hours:02d}:{minutes:02d}:{seconds:02d},{milliseconds:03d}"

    def final_export(self, transcript_text, settings, project_info, unassigned_handling="skip"):
        format_extensions = {
            'html': '.html',
            'txt': '.txt',
            'srt': '.srt',
            'docx': '.docx'
        }

        file_ext = format_extensions.get(settings['format'], '.txt')

        default_name = ""
        if project_info.get('name'):
            default_name = project_info['name'].replace(" ", "_")
        else:
            default_name = "transcript"

        if settings['convention'] != 'gat2':
            default_name += f"_{settings['convention']}"

        default_name += file_ext

        file_path, _ = QFileDialog.getSaveFileName(
            self, f"Export Transcript", default_name,
            f"{settings['format'].upper()} Files (*{file_ext})"
        )

        if not file_path:
            return

        try:
            if settings['format'] == 'srt':
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(transcript_text)

            elif settings['format'] == 'docx':
                try:
                    import docx
                    from docx.shared import Pt
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    doc = docx.Document()

                    if settings.get('include_title', True) and project_info.get('name'):
                        title = doc.add_heading(project_info['name'], 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    if settings.get('include_memo', True) and project_info.get('memo'):
                        doc.add_paragraph(f"Project Memo: {project_info['memo']}")

                    if settings.get('include_audio', True) and self.audio_file_path:
                        doc.add_paragraph(f"Audio File: {Path(self.audio_file_path).name}")

                    doc.add_paragraph()

                    for line in transcript_text.split('\n'):
                        if line.strip():
                            p = doc.add_paragraph(line)
                            font_name = 'Courier New'
                            if settings['convention'] == 'dresing_pehl':
                                font_name = 'Times New Roman'
                            for run in p.runs:
                                run.font.name = font_name
                                run.font.size = Pt(10)
                        else:
                            doc.add_paragraph()

                    doc.save(file_path)

                except ImportError:
                    QMessageBox.warning(self, "DOCX Export",
                        "python-docx library not found. Please install it with: pip install python-docx\n\n"
                        "Exporting as plain text instead.")
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(transcript_text)

            elif settings['format'] == 'html':
                # Build header from settings
                header_lines = []
                if settings.get('include_title', True) and project_info.get('name'):
                    header_lines.append(f"<h1>{self.escape_html(project_info['name'])}</h1>")
                if settings.get('include_memo', True) and project_info.get('memo'):
                    header_lines.append(f"<p class=\"headerstyle\"><strong>Project Memo:</strong> {self.escape_html(project_info['memo'])}</p>")
                if settings.get('include_audio', True) and self.audio_file_path:
                    audio_name = Path(self.audio_file_path).name
                    header_lines.append(f"<p class=\"headerstyle\"><strong>Audio File:</strong> {self.escape_html(audio_name)}</p>")
                header = "\n".join(header_lines) + "\n" if header_lines else ""

                # Choose font based on convention
                font_family = "'Courier New', monospace"
                if settings['convention'] == "dresing_pehl":
                    font_family = "'Times New Roman', serif"

                html_content = f"""<!DOCTYPE html>
            <html>
            <head>
            <meta charset="UTF-8">
            <title>Transcript - {self.escape_html(project_info.get('name', 'Untitled'))}</title>
            <style>
            body {{
                font-family: {font_family};
                font-size: 10pt;
                line-height: 1.2;
                margin: 20px;
                white-space: pre-wrap;
            }}
            h1 {{
                font-family: Arial, sans-serif;
                color: #333;
                padding-bottom: 10px;
            }}
            
            .headerstyle
                {{
                font-family: Arial, sans-serif;
                color: #333;
                }}
            </style>
            </head>
            <body>
            {header}<br>{self.escape_html(transcript_text)}
            </body>
            </html>"""
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            else:
                # Plain text export
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(transcript_text)

            QMessageBox.information(self, "Success", f"Transcript exported to {file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Could not export file: {str(e)}")

    def closeEvent(self, event):
        """Clean up on close"""
        if self.audio_player:
            self.audio_player.cleanup()

        if self.check_unsaved_changes():
            event.accept()
        else:
            event.ignore()


def main():
    app = QApplication(sys.argv)

    splash = None
    splash_path = resource_path("images/splash.png")
    if os.path.exists(splash_path):
        splash_pix = QPixmap(splash_path)
        splash = QSplashScreen(splash_pix, Qt.WindowStaysOnTopHint)
        splash.show()
        app.processEvents()
        splash.showMessage("Initializing CapsGAT...", Qt.AlignBottom | Qt.AlignCenter, Qt.black)
        app.processEvents()
    else:
        logger.warning("Splash image not found, continuing without splash.")

    # Create main window, passing splash for progress updates
    editor = SRTEditor(splash)

    if splash:
        splash.finish(editor)
    editor.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()
