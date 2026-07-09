"""File export/writing functions for CapsQual.

Consolidates all file-writing logic (HTML, DOCX, TXT, SRT) into one module.
Generation functions live in generators.py.

This module contains no Qt imports — pure Python file I/O.
"""

import os
import re
from pathlib import Path

from generators import escape_html, convert_markup_to_html, strip_markup

__all__ = [
    'build_html_content',
    'write_html_file',
    'write_srt_file',
    'write_txt_file',
    'add_formatted_paragraph',
    'write_docx_file',
]


# ──────────────────────────────────────────────
#  HTML helpers
# ──────────────────────────────────────────────

def build_html_content(transcript_text, settings, project_info, audio_file_path=None):
    """Build a full HTML document string from transcript text.

    Parameters
    ----------
    transcript_text : str
        The formatted transcript text (plain text with line breaks).
    settings : dict
        Must include at least 'convention'. May also include
        'include_title', 'include_memo', 'include_audio'.
    project_info : dict
        May include 'name', 'memo'.
    audio_file_path : str or None
        Path to the audio file (used for filename in header).

    Returns
    -------
    str
        Complete HTML document.
    """
    escaped_text = escape_html(transcript_text)
    formatted = convert_markup_to_html(escaped_text)

    # Convert newlines to <p> tags for proper block-level rendering in QDA software
    # Replace all spaces with &nbsp; so QDA software preserves all formatting
    lines = formatted.split('\n')
    lines = [line.replace(' ', '&nbsp;') for line in lines]
    # Blank lines become <br> instead of <p></p> (which has zero height with p margin:0)
    formatted_lines = []
    for line in lines:
        if line:
            formatted_lines.append(f'<p>{line}</p>')
        else:
            formatted_lines.append('<br>')
    formatted = '\n    '.join(formatted_lines)

    header_lines = []
    if settings.get('include_title', True) and project_info.get('name'):
        header_lines.append(f"<h1>{escape_html(project_info['name'])}</h1>")
    if settings.get('include_memo', True) and project_info.get('memo'):
        header_lines.append(f"<p class=\"headerstyle\"><strong>Project Memo:</strong> {escape_html(project_info['memo'])}</p>")
    if settings.get('include_audio', True) and audio_file_path:
        audio_name = Path(audio_file_path).name
        header_lines.append(f"<p class=\"headerstyle\"><strong>Audio File:</strong> {escape_html(audio_name)}</p>")
    header = "\n".join(header_lines) + "\n" if header_lines else ""

    convention = settings.get('convention', 'gat2')
    if convention == "dresing_pehl":
        font_family = "'Times New Roman', serif"
    else:
        font_family = "'Courier New', monospace"

    html_content = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>Transcript - {escape_html(project_info.get('name', 'Untitled'))}</title>
<style>
body {{
    font-family: {font_family};
    font-size: 10pt;
    line-height: 1.2;
    margin: 20px;
}}
h1 {{
    font-family: Arial, sans-serif;
    color: #333;
    padding-bottom: 10px;
}}
.headerstyle {{
    font-family: Arial, sans-serif;
    color: #333;
}}
p {{
    margin: 0;
    padding: 0;
}}
</style>
</head>
<body>
{header}<br>
{formatted}
</body>
</html>"""
    return html_content


def write_html_file(html_content, file_path):
    """Write HTML content to a file."""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)


# ──────────────────────────────────────────────
#  SRT writer
# ──────────────────────────────────────────────

def write_srt_file(srt_text, file_path):
    """Write SRT transcript text to a file."""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(srt_text)


# ──────────────────────────────────────────────
#  TXT writer
# ──────────────────────────────────────────────

def write_txt_file(transcript_text, file_path):
    """Strip markup and write plain text to a file."""
    stripped = strip_markup(transcript_text)
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(stripped)


# ──────────────────────────────────────────────
#  DOCX helpers
# ──────────────────────────────────────────────

def add_formatted_paragraph(doc, text, style_name=None):
    """Add a paragraph to a python-docx document with appropriate runs
    (bold/italic/underline markers converted to run formatting)."""
    if not text:
        if style_name:
            doc.add_paragraph(style=style_name)
        else:
            doc.add_paragraph()
        return

    # Split on markers, but keep them as tokens
    pattern = re.compile(r'(#@[BIU]|#@/[BIU])')
    parts = []
    last_end = 0
    for m in pattern.finditer(text):
        start, end = m.span()
        if start > last_end:
            parts.append(('text', text[last_end:start]))
        parts.append(('marker', m.group()))
        last_end = end
    if last_end < len(text):
        parts.append(('text', text[last_end:]))

    if style_name:
        p = doc.add_paragraph(style=style_name)
    else:
        p = doc.add_paragraph()
    bold = italic = underline = False
    run_text = ""

    for typ, content in parts:
        if typ == 'text':
            run_text += content
        else:  # marker
            if run_text:
                run = p.add_run(run_text)
                run.bold = bold
                run.italic = italic
                run.underline = underline
                run_text = ""
            if content == '#@B':
                bold = True
            elif content == '#@/B':
                bold = False
            elif content == '#@I':
                italic = True
            elif content == '#@/I':
                italic = False
            elif content == '#@U':
                underline = True
            elif content == '#@/U':
                underline = False
    if run_text:
        run = p.add_run(run_text)
        run.bold = bold
        run.italic = italic
        run.underline = underline


def write_docx_file(transcript_text, settings, project_info, audio_file_path, file_path):
    """Build and save a DOCX document from transcript text.

    Returns True on success, False on failure (with fallback to TXT handled by caller).
    On ImportError (python-docx not installed), returns False without writing.
    """
    try:
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        return False

    doc = docx.Document()

    if settings.get('include_title', True) and project_info.get('name'):
        title = doc.add_heading(project_info['name'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if settings.get('include_memo', True) and project_info.get('memo'):
        doc.add_paragraph(f"Project Memo: {project_info['memo']}")

    if settings.get('include_audio', True) and audio_file_path:
        doc.add_paragraph(f"Audio File: {Path(audio_file_path).name}")

    doc.add_paragraph()

    styles = doc.styles
    style_name = "TranscriptBody"
    if style_name in [s.name for s in styles]:
        body_style = styles[style_name]
    else:
        body_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    if settings.get('convention') in ('gat2', 'tiq'):
        body_style.font.name = 'Courier New'
    body_style.font.size = Pt(10)
    para_fmt = body_style.paragraph_format
    para_fmt.line_spacing = 1
    para_fmt.space_after = Pt(0)
    para_fmt.space_before = Pt(0)

    for line in transcript_text.split('\n'):
        if line.strip():
            add_formatted_paragraph(doc, line, style_name=style_name)
        else:
            doc.add_paragraph(style=style_name)

    doc.save(file_path)
    return True

